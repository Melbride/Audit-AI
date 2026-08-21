# Core FastAPI imports for building the API, handling file uploads, and raising errors
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends, APIRouter, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.security import OAuth2PasswordBearer
from pydantic import BaseModel, field_validator
from typing import Optional, Literal, List

# Auth-related imports for password hashing and JWT tokens
from passlib.context import CryptContext
from jose import JWTError, jwt
from datetime import datetime, timedelta
from validators.field_validator import FieldValidator, FieldValidationError
from contact_validators import validate_kenyan_phone, validate_real_email, validate_password_strength
from engagement_validators import validate_engagement_creation, validate_engagement_update, validate_engagement_exists, validate_client_exists
from file_validators import validate_file_submission, validate_selected_sheets
from section_validators import validate_section_in_engagement_scope
from error_responses import register_error_handlers, error_detail
from engagement_notifications import notify_if_ready_for_final_analysis

# File-reading libraries for PDF and DOCX extraction
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import pdfplumber
from docx import Document

# General-purpose libraries used throughout the file
import pandas as pd
import json
import uuid
import os
import sys
import re
import shutil
import secrets
import hashlib
import time
import warnings
import io
from dotenv import load_dotenv

# AI detection helpers, used to identify columns and suggest file types
from detector import detect_columns_with_llm, build_detection_result, suggest_file_type, FILE_TYPE_CATEGORIES

# Header detection for Excel files with report-style metadata
from header_detector import detect_excel_header, detect_csv_header

# Shared JWT auth dependency (decodes token, restricts routes by role)
from auth import require_role, get_current_user

# Database helper functions for mappings, uploads, corrections, snapshots, and acknowledgments
from database import (
    init_db, get_db, get_connection, save_mapping, get_mapping, save_upload, get_upload, get_uploads,
    save_cleaning_acknowledgment, get_acknowledged_issue_ids,
    save_cleaning_corrections, get_cleaning_corrections,
    save_fingerprint, get_fingerprint,
    save_cleaning_snapshot, get_cleaning_snapshot,
    save_cleaned_registry, get_cleaned_files_for_client, get_cleaned_file_data,
    save_analysis, get_saved_analyses, get_saved_analysis, delete_saved_analysis,
    get_saved_analyses_for_engagement, get_saved_analyses_for_file,
    get_or_create_workspace, get_workspace_by_id, update_workspace_data, get_engagement_workspaces, get_user_workspaces,
    check_all_sections_completed, get_all_sections_data,
    get_workflow_stage, update_workflow_stage, mark_workflow_step_completed, initialize_workflow_stage,
     save_tb_validation_result, clear_tb_validation_result, normalize_mapped_to,
)

# Cleaning engine, Excel export, and Excel diff logic
from cleaner import clean_dataframe
from excel_export import build_cleaning_workbook
from excel_diff import diff_uploaded_against_snapshot
from dateutil import parser as dateutil_parser
from analyzer import calculate_breakdowns, calculate_monthly_trend, detect_anomalies, generate_ai_insights, generate_financial_ai_insights, determine_analysis_scope

from engines.accounting_validation.trial_balance_validator import validate_trial_balance
from engines.account_mapping.account_classifier import build_account_mapping_result
from database import save_account_mapping, get_account_mapping
from engines.financial_reporting.statement_generator import generate_financial_statements
from engines.financial_reporting.financial_ratios import calculate_financial_ratios
from engines.financial_reporting.financial_analytics import calculate_financial_analytics
from engines.financial_reporting.comparative_analytics import generate_comparative_analytics
from validators.date_validator import DateValidator, DateContext, DateValidationError
# Load environment variables from the .env file
load_dotenv()

# Resolve the backend directory and add it to the Python path so internal imports work correctly
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BACKEND_DIR = os.path.join(BASE_DIR, "backend")
if BACKEND_DIR not in sys.path:
    sys.path.append(BACKEND_DIR)

# Create the FastAPI app instance
app = FastAPI(title="AuditAI API Running!", debug=True)

# Allow requests from any origin, method, and header (open CORS policy)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Standardized error response shape (Section #9 of the validation doc) —
# wraps every HTTPException raised anywhere in the app into a consistent
# {success, error, message, details} body, while keeping the existing
# `detail` string field so already-shipped frontend pages keep working
# unchanged. See error_responses.py for the full rationale.
register_error_handlers(app)

# Initialize the database tables when the app starts up
@app.on_event("startup")
async def startup_event():
    init_db()

# Report review/approval workflow (Month 3): /api/reports/...
from report_routes import router as report_router
app.include_router(report_router)

# Report export system (Month 3): /api/reports/{id}/export, /api/reports/exports/{id}/download
from report_exports import router as report_export_router
app.include_router(report_export_router)

# Section milestones (preset checkpoints) and review entries (issues/highlights/redo)
from section_tracking import router as section_tracking_router
app.include_router(section_tracking_router)

# Financial statement starter template (Trial Balance + 3-statement skeleton)
# from statement_template import router as statement_template_router
# app.include_router(statement_template_router)

# Auth configuration: secret key, hashing algorithm, token lifetime, and password hashing context
SECRET_KEY = os.getenv("SECRET_KEY")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_HOURS = 8
pwd_context = CryptContext(schemes=["sha256_crypt"], deprecated="auto")

# Folder where uploaded files are stored, created automatically if missing
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# File extensions allowed for upload through the AI pipeline
ALLOWED_EXTENSIONS = {"csv", "xlsx", "xls", "pdf", "docx"}

# Internal columns the system adds for its own bookkeeping that should never be treated as real data and should never reach column detection or mapping
RESERVED_INTERNAL_COLUMNS = {"_row_id", "_is_duplicate"}

# Hash a plain text password for storage
def hash_password(password: str):
    return pwd_context.hash(password)

# Check a plain text password against a stored hash
def verify_password(plain: str, hashed: str):
    return pwd_context.verify(plain, hashed)

# Create a signed JWT access token that expires after ACCESS_TOKEN_EXPIRE_HOURS
def create_token(data: dict):
    expire = datetime.utcnow() + timedelta(hours=ACCESS_TOKEN_EXPIRE_HOURS)
    data.update({"exp": expire})
    return jwt.encode(data, SECRET_KEY, algorithm=ALGORITHM)

# Get the lowercase file extension from a filename
def get_extension(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

# Extract tables (or plain text if no tables found) from a PDF file into a dataframe
def extract_pdf(file_path: str):
    tables = []
    full_text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_tables = page.extract_tables()
            for table in page_tables:
                if table:
                    headers = table[0]
                    rows = table[1:]
                    df = pd.DataFrame(rows, columns=headers)
                    tables.append(df)
            full_text += page.extract_text() or ""
    if tables:
        return pd.concat(tables, ignore_index=True), "table"
    lines = [line.strip() for line in full_text.splitlines() if line.strip()]
    if lines:
        return pd.DataFrame({"raw_text": lines}), "text"
    return None, None

# Extract tables (or plain text if no tables found) from a DOCX file into a dataframe
def extract_docx(file_path: str):
    doc = Document(file_path)
    tables = []
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = []
        for row in table.rows[1:]:
            rows.append([cell.text.strip() for cell in row.cells])
        df = pd.DataFrame(rows, columns=headers)
        tables.append(df)
    if tables:
        return pd.concat(tables, ignore_index=True), "table"
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if lines:
        return pd.DataFrame({"raw_text": lines}), "text"
    return None, None

# Read a file from disk into a dataframe based on its extension, cleaning up
# empty columns, whitespace, and reserved internal columns along the way

# Read a file from disk into a dataframe based on its extension.
# Supports sheet selection, smart header detection, and stored header rows.
def read_file_to_df(
    save_path: str,
    ext: str,
    sheet_name: str = None,
    file_type: Optional[str] = None,
    header_row_index: Optional[int] = None
):
    metadata = {}

    if ext == "csv":
        if header_row_index is not None:
            df = pd.read_csv(
                save_path,
                header=header_row_index,
                dtype=str
            )
        else:
            header_row, df, metadata = detect_csv_header(
                save_path,
                file_type
            )

    elif ext in ["xlsx", "xls"]:
        if header_row_index is not None:
            df = pd.read_excel(
                save_path,
                sheet_name=sheet_name if sheet_name else 0,
                header=header_row_index,
                dtype=str
            )
        else:
            header_row, df, metadata = detect_excel_header(
                save_path,
                sheet_name,
                file_type
            )

    elif ext == "pdf":
        df, _ = extract_pdf(save_path)
        return df, {}

    elif ext == "docx":
        df, _ = extract_docx(save_path)
        return df, {}

    else:
        return None, {}

    if df is not None:
        df = df.dropna(axis=1, how="all")
        df = df.loc[:, ~(df == "").all()]

        df = df.map(
            lambda x: x.strip() if isinstance(x, str) else x
        )

        df = df.drop(
            columns=[
                c for c in df.columns
                if c in RESERVED_INTERNAL_COLUMNS
            ],
            errors="ignore"
        )

    return df, metadata

# Calculate the fraction of non-empty values in each column of a dataframe
def calculate_fill_rates(df: pd.DataFrame) -> dict:
    fill_rates = {}
    total = len(df)
    for col in df.columns:
        filled = df[col].replace("", float("nan")).dropna().count()
        fill_rates[col] = round(filled / total, 2) if total > 0 else 0.0
    return fill_rates

# Check that each column's declared field_type actually matches its real data. Prevents usersfrom mistakenly marking a text column as numeric/date
# Check that each column's declared field_type actually matches its real data. Prevents users from mistakenly marking a text column as numeric/date

def validate_field_types_against_data(df: pd.DataFrame, mapping: dict) -> list:
    problems = []
    def try_parse_date(val_str: str) -> bool:
        # Strip ordinal suffixes: 1st, 2nd, 3rd, 22nd, 28th ...
        cleaned = re.sub(r'(\d+)(st|nd|rd|th)\b', r'\1', val_str, flags=re.IGNORECASE).strip()
        # Bare numbers are not dates (e.g. '999', '42')
        if re.fullmatch(r'\d+', cleaned):
            return False
        # dateutil handles any date format without a hardcoded list
        try:
            dateutil_parser.parse(cleaned, default=pd.Timestamp('2000-01-01'))
            return True
        except (ValueError, OverflowError):
            return False

    for original_col, info in mapping.items():
        if not isinstance(info, dict):
            continue
        field_type = info.get("field_type")
        mapped_to = info.get("mapped_to")
        if original_col not in df.columns or field_type in (None, "unknown") or mapped_to in (None, "unknown"):
            continue
        values = df[original_col].dropna().replace("", None).dropna()
        if len(values) == 0:
            continue

        if field_type == "numeric":
            def strip_for_check(v):
                v = str(v).strip()
                # If it's already a plain number, no stripping needed
                try:
                    float(v)
                    return v
                except ValueError:
                    pass
                is_negative = v.startswith("(") and v.endswith(")")
                v = v.strip("()%")
                v = re.sub(r"[^\d.-]", "", v)
                return ("-" + v) if is_negative and not v.startswith("-") else v

            cleaned_values = values.apply(strip_for_check)
            convertible = pd.to_numeric(cleaned_values, errors="coerce")
            failure_rate = convertible.isna().mean()
            if failure_rate > 0.3:
                problems.append(
                    f"'{original_col}' is marked as Number but most of its values "
                    f"aren't numbers (e.g. '{values.iloc[0]}'). Please check the field type."
                )

        elif field_type == "date":
            failures = sum(1 for val in values if not try_parse_date(str(val).strip()))
            failure_rate = failures / len(values)
            if failure_rate > 0.3:
                problems.append(
                    f"'{original_col}' is marked as Date but most of its values "
                    f"don't look like dates (e.g. '{values.iloc[0]}'). Please check the field type."
                )
    return problems

# Compute a stable fingerprint for a set of column names, used to detect when a client uploads a file with the same schema as before
def compute_schema_fingerprint(columns: list) -> str:
    sorted_cols = sorted([col.lower().strip() for col in columns])
    fingerprint = hashlib.md5(json.dumps(sorted_cols).encode()).hexdigest()
    return fingerprint

# Helper function to get file-specific mapping using file_id
def get_file_specific_mapping(file_id: str, client_id: str, file_type: str) -> dict:
    """
    Get mapping for a specific file using its file_id.
    This ensures file-specific mapping isolation and prevents cross-file mapping pollution.
    """
    mapping = get_mapping(client_id, file_type, None, file_id)
    if not mapping:
        raise HTTPException(
            status_code=400, 
            detail="No saved mapping found for this specific file. Please detect the columns and confirm the mapping first."
        )
    
    return mapping

# Validate mapping for duplicate target fields before saving
def validate_mapping_before_save(mapping: dict) -> tuple:
    """
    Validate a mapping before saving to catch duplicate target fields.
    Returns (is_valid, error_message).
    """
    seen = {}
    for original_col, info in mapping.items():
        if not isinstance(info, dict):
            continue
        raw_mapped_to = str(info.get("mapped_to", "")).strip()
        mapped_to = normalize_mapped_to(raw_mapped_to)
        # "unknown" is allowed to repeat
        if mapped_to == ("unknown"):
            continue
        seen.setdefault(mapped_to, []).append(original_col)
    
    duplicates = {standard_name: cols for standard_name, cols in seen.items() if len(cols) > 1}
    if duplicates:
        details = "; ".join(
            f"'{standard_name}' is claimed by columns {cols}"
            for standard_name, cols in duplicates.items()
        )
        return False, f"Mapping conflict: two or more original columns are mapped to the same field. {details}. Please give each column a unique 'Mapped To' value."
    
    return True, None

# Find an uploaded file on disk by its file_id, trying each allowed extension
def locate_uploaded_file(file_id: str, preferred_ext=None):
    if preferred_ext and preferred_ext in ALLOWED_EXTENSIONS:
        path = os.path.join(UPLOAD_DIR, f"{file_id}.{preferred_ext}")
        if os.path.exists(path):
            return path, preferred_ext
    for extension in ALLOWED_EXTENSIONS:
        path = os.path.join(UPLOAD_DIR, f"{file_id}.{extension}")
        if os.path.exists(path):
            return path, extension
    return None, None

def remove_uploaded_file_variants(file_id: str, keep_ext: str | None = None):
    for extension in ALLOWED_EXTENSIONS:
        if keep_ext and extension == keep_ext:
            continue
        path = os.path.join(UPLOAD_DIR, f"{file_id}.{extension}")
        if os.path.exists(path):
            try:
                os.remove(path)
            except OSError:
                pass

# Strip the "[UNRESOLVED]" marker from a header label to get its real underlying name
def normalize_header_label(value) -> str:
    label = str(value or "").strip()
    if label.startswith("[UNRESOLVED]"):
        label = label.replace("[UNRESOLVED]", "", 1).strip()
    return label

# Build a stable hash that uniquely identifies one issue, used to track acknowledgments and avoid duplicate issue records
def issue_fingerprint(file_id: str, client_id: str, file_type: str, issue: dict) -> str:
    raw = "|".join([
        str(file_id),
        str(client_id),
        str(file_type),
        str(issue.get("row_index", "")),
        str(issue.get("column", "")),
        str(issue.get("original_value", "")),
        str(issue.get("issue", "")),
    ])
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()

# Attach a stable issue_id and a default "pending" decision to every issue in a report
def enrich_issues_with_ids(report: dict, file_id: str, client_id: str, file_type: str) -> dict:
    for issue in report.get("issues", []):
        issue["issue_id"] = issue_fingerprint(file_id, client_id, file_type, issue)
        issue["decision"] = "pending"
    return report

# Recalculate all summary counts on a report (flagged rows, clean rows, and
# issue counts by severity) based on whatever issues are currently in the list
def rebuild_report_counts(report: dict) -> dict:
    issues = report.get("issues", [])
    total_rows = report.get("total_rows", 0)

    row_issues = [
        i for i in issues
        if i.get("row_index") not in ("N/A", None)
    ]

    column_issues = [
        i for i in issues
        if i.get("row_index") in ("N/A", None)
    ]
    flagged_rows = len(set(
    i.get("row_index") for i in row_issues
    if i.get("row_index") not in (None, "N/A")
    ))
    report["flagged_rows"] = flagged_rows
    report["clean_rows"] = total_rows - flagged_rows
    report["row_issues"] = len(row_issues)
    report["column_issues"] = len(column_issues)
    report["total_issues"] = len(issues)
    report["high_issues"] = len([i for i in issues if i.get("severity") == "high"])
    report["medium_issues"] = len([i for i in issues if i.get("severity") == "medium"])
    report["info_issues"] = len([i for i in issues if i.get("severity") == "info"])
    return report

# Remove any issue the auditor has already acknowledged, then recompute the report counts
def filter_acknowledged_issues(report: dict, file_id: str, client_id: str, file_type: str) -> dict:
    acknowledged = get_acknowledged_issue_ids(file_id, client_id, file_type)
    if not acknowledged:
        return report
    report["issues"] = [
        issue for issue in report.get("issues", [])
        if issue.get("issue_id") not in acknowledged
    ]
    return rebuild_report_counts(report)

# Apply every saved correction (cell edits, row deletions, column deletions) to
# the source dataframe before cleaning runs, so corrections persist across cleaning cycles
def apply_saved_corrections(df: pd.DataFrame, mapping: dict, corrections: list) -> pd.DataFrame:
    if not corrections:
        return df
    # Build a lookup from standard mapped column names back to their original column names
    standard_to_original = {
        info.get("mapped_to"): original_col
        for original_col, info in mapping.items()
        if isinstance(info, dict) and info.get("mapped_to") not in ("", "unknown", None)
    }
    # Separate the corrections list into rows to drop, columns to drop, and plain value edits
    rows_to_drop = []
    columns_to_drop = []
    value_corrections = []
    for correction in corrections:
        col = correction.get("column_name", correction.get("column"))
        if col == "_row_deleted":
            rows_to_drop.append(int(correction["row_index"]))
        elif col and str(col).startswith("_column_deleted:"):
            original_col_name = col.split(":", 1)[1]
            columns_to_drop.append(original_col_name)
        else:
            value_corrections.append(correction)

    # Apply row deletions
    if rows_to_drop:
        df = df.drop(index=[r for r in rows_to_drop if r in df.index])

    # Apply column deletions, resolving standard names back to original column names where needed
    if columns_to_drop:
        resolved_columns_to_drop = []
        for col_name in columns_to_drop:
            if col_name in df.columns:
                resolved_columns_to_drop.append(col_name)
            else:
                original = standard_to_original.get(col_name)
                if original and original in df.columns:
                    resolved_columns_to_drop.append(original)
        df = df.drop(columns=resolved_columns_to_drop)

    # Apply plain cell value corrections
    for correction in value_corrections:
        row_index = int(correction["row_index"])
        issue_col = correction.get("column_name", correction.get("column"))
        source_col = standard_to_original.get(issue_col, issue_col)
        if source_col in df.columns and row_index in df.index:
            df.at[row_index, source_col] = correction["corrected_value"]
    return df

# Adapt a saved mapping to match whatever column names are present in a newly
# uploaded file, in case columns were renamed since the mapping was last saved
def adapt_mapping_to_uploaded_headers(mapping: dict, columns: list) -> dict:
    adapted = {}
    column_set = set(columns)
    for original_col, info in mapping.items():
        if not isinstance(info, dict):
            adapted[original_col] = info
            continue
        mapped_to = info.get("mapped_to")
        if original_col in column_set:
            adapted[original_col] = info
        elif mapped_to and mapped_to != "unknown" and mapped_to in column_set:
            adapted[mapped_to] = {**info, "mapped_to": mapped_to}
        else:
            adapted[original_col] = info
    return adapted

# Run a full cleaning cycle for a file: read it from disk, apply saved corrections,
# run the cleaning engine, filter out acknowledged issues, and compute final counts
def run_cleaning_cycle(file_id: str, client_id: str, file_type: str, mapping: dict, sheet_name: str = None):
    save_path, file_ext = locate_uploaded_file(file_id)
    if not save_path:
        raise HTTPException(status_code=404, detail="File not found. Please upload the file first.")
    
    # Get stored header configuration from uploads table
    upload_record = get_upload(file_id)
    stored_header_row = upload_record.get('header_row_index') if upload_record else None
    stored_sheet_name = upload_record.get('sheet_name') if upload_record else None
    
    try:

       df, metadata = read_file_to_df(
           save_path,
           file_ext,
           stored_sheet_name,
           file_type=file_type,
           header_row_index=stored_header_row
) 

       

       if df is None:
            raise HTTPException(status_code=400, detail="Could not read file.")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")

    corrections = get_cleaning_corrections(file_id, client_id, file_type)
    df = apply_saved_corrections(df, mapping, corrections)
    fill_rates = calculate_fill_rates(df)
    try:
        cleaned_df, report = clean_dataframe(df, mapping, fill_rates, file_type)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Cleaning failed: {str(e)}")
    report = enrich_issues_with_ids(report, file_id, client_id, file_type)
    report = filter_acknowledged_issues(report, file_id, client_id, file_type)
    report = rebuild_report_counts(report)
    # can_proceed is only true once every issue (high, medium, and info severity) has been resolved
    report["can_proceed"] = report["total_issues"] == 0
    # Save the current cleaned state so the auditor can retun to it and also download
    save_cleaned_registry(file_id, client_id, file_type, cleaned_df, report)
    # Initialize workflow stage when file is successfully cleaned
    if report["can_proceed"]:
        initialize_workflow_stage(file_id, client_id, file_type)
    return cleaned_df, report

# Function to build financial analysis context
def build_financial_analysis_context(cleaned_df: pd.DataFrame, mapping: dict, client_id: str, file_type: str) -> dict:
    is_ledger_file = file_type in ("trial_balance", "general_ledger")

    if is_ledger_file:
        breakdowns = {}
        monthly_trend = {}
        anomalies = []
    else:
        breakdowns = calculate_breakdowns(cleaned_df, mapping)
        monthly_trend = calculate_monthly_trend(cleaned_df, mapping)
        anomalies = detect_anomalies(monthly_trend)

    generic_scope = determine_analysis_scope(breakdowns, monthly_trend)

    context = {
        "analysis_scope": generic_scope,
        "analysis_basis": "generic_columns",
        "breakdowns": breakdowns,
        "monthly_trend": monthly_trend,
        "anomalies": anomalies,
        "financial_statements": None,
        "financial_ratios": None,
        "financial_analytics": None,
        "comparative_analytics": None,
        "generic_analysis": {
            "analysis_scope": generic_scope,
            "breakdowns": breakdowns,
            "monthly_trend": monthly_trend,
            "anomalies": anomalies,
        },
    }

    account_mapping = get_account_mapping(client_id, file_type)
    if not account_mapping:
        return context

    statements = generate_financial_statements(cleaned_df, mapping, account_mapping)
    if not statements.get("applicable"):
        return context

    ratios = calculate_financial_ratios(
        statements["income_statement"],
        statements["balance_sheet"],
    )
    statements["financial_ratios"] = ratios

    financial_analytics = calculate_financial_analytics(statements, ratios)
    comparative_analytics = generate_comparative_analytics(
        cleaned_df,
        mapping,
        account_mapping,
        grain="year",
    )
    context.update({
        "analysis_scope": "financial_statements",
        "analysis_basis": "classified_accounts",
        "financial_statements": statements,
        "financial_ratios": ratios,
        "financial_analytics": financial_analytics,
        "comparative_analytics": comparative_analytics,
    })
    return context

def get_engagement_approved_files(engagement_id: int) -> list:
    """
    For every in-scope, approved section in this engagement, return the
    file_id/sheet_name/file_type its latest approved submission points to.

    Returns a list of dicts:
      { section_id, section_name, file_id, sheet_name, file_type }

    Raises HTTPException(400) if any approved section's file_id no longer
    resolves to a real uploaded file — better to fail loudly here than
    silently skip a section's data out of a final analysis.
    """
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("""
        SELECT sec.section_id, sec.section_name, latest.file_id, latest.sheet_name
        FROM audit_sections sec
        INNER JOIN (
            SELECT section_id, file_id, sheet_name, status,
                   ROW_NUMBER() OVER (PARTITION BY section_id ORDER BY created_at DESC) AS rn
            FROM submissions
            WHERE engagement_id = %s
        ) latest ON latest.section_id = sec.section_id AND latest.rn = 1
        WHERE sec.engagement_id = %s AND sec.in_scope = 1 AND latest.status = 'Approved'
    """, (engagement_id, engagement_id))
    rows = cursor.fetchall()

    for row in rows:
        # Select both the raw extension (file_type) and the semantic type
        # (semantic_file_type) -- mappings are saved under the SEMANTIC
        # type (e.g. 'trial_balance'), not the file extension (e.g. 'csv'),
        # same distinction get_resume_state() already relies on elsewhere.
        cursor.execute("SELECT file_type, semantic_file_type FROM uploads WHERE file_id = %s", (row["file_id"],))
        upload = cursor.fetchone()
        if not upload:
            conn.close()
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Section '{row['section_name']}' is approved but its file "
                    f"(file_id={row['file_id']}) no longer exists in uploads. "
                    "Cannot build final analysis until this is resolved."
                ),
            )
        row["file_type"] = upload.get("semantic_file_type") or upload.get("file_type") or "other"

    conn.close()
    return rows

def get_engagement_combined_dataset(engagement_id: int) -> dict:
    """
    THE single source of truth for "what data does this engagement's
    approved work actually consist of". Both build_engagement_final_analysis
    (AI Insights) and the report generator call this — never independently —
    so a report and the final analysis can never show different numbers for
    the same engagement.

    Gathers every in-scope, approved section's cleaned data and concatenates
    it into one combined trial balance. Sections that share the same file_id
    are deduplicated automatically.

    Returns a dict: { combined_df, mapping, account_mapping, file_type,
    client_id, included_sections }.

    Raises HTTPException(409) if any in-scope section is not yet approved.
    Raises HTTPException(422) if approved sections' files use inconsistent
    file_types.
    """
    completion = check_all_sections_completed(engagement_id)
    if not completion["all_completed"]:
        raise HTTPException(
            status_code=409,
            detail={
                "message": "Not all in-scope sections are approved yet.",
                "pending_sections": completion["pending_sections"],
                "completed_sections": completion["completed_sections"],
                "excluded_sections": completion["excluded_sections"],
            },
        )

    approved_files = get_engagement_approved_files(engagement_id)
    if not approved_files:
        raise HTTPException(status_code=400, detail="No approved, in-scope sections with data found for this engagement.")

    file_types = {f["file_type"] for f in approved_files}
    if len(file_types) > 1:
        raise HTTPException(
            status_code=422,
            detail=(
                f"Approved sections use inconsistent file types ({', '.join(sorted(file_types))}). "
                "Final analysis requires all approved section files to be the same type."
            ),
        )
    file_type = file_types.pop()

    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT client_id FROM engagements WHERE engagement_id = %s", (engagement_id,))
    engagement = cursor.fetchone()
    if not engagement:
        conn.close()
        raise HTTPException(status_code=404, detail=f"Engagement {engagement_id} not found.")
    client_id = str(engagement["client_id"])
    conn.close()
    mapping = get_mapping(client_id, file_type)
    if not mapping:
        raise HTTPException(status_code=400, detail=f"No saved column mapping found for client {client_id} / {file_type}.")

    # Dedupe by (file_id, sheet_name) — a shared workbook referenced by
    # multiple sections is only read and cleaned once.
    distinct_sources = {}
    for f in approved_files:
        key = (f["file_id"], f["sheet_name"])
        distinct_sources.setdefault(key, []).append(f["section_name"])

    cleaned_frames = []
    included_sections = []
    for (file_id, sheet_name), section_names in distinct_sources.items():
        cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping, sheet_name)
        cleaned_frames.append(cleaned_df)
        included_sections.append({
            "sections": section_names,
            "file_id": file_id,
            "sheet_name": sheet_name,
            "total_issues": report.get("total_issues"),
        })

    combined_df = pd.concat(cleaned_frames, ignore_index=True)
    account_mapping = get_account_mapping(client_id, file_type)

    return {
        "combined_df": combined_df,
        "mapping": mapping,
        "account_mapping": account_mapping,
        "file_type": file_type,
        "client_id": client_id,
        "included_sections": included_sections,
    }


def build_engagement_final_analysis(engagement_id: int) -> dict:
    """
    Runs the existing (untouched) generate_financial_statements()/
    generate_financial_ai_insights() pipeline once on this engagement's
    combined approved dataset (see get_engagement_combined_dataset) — this
    is the AI Insights consumer of that shared dataset; build_report_payload
    (below) is the other.
    """
    dataset = get_engagement_combined_dataset(engagement_id)
    combined_df = dataset["combined_df"]
    mapping = dataset["mapping"]

    financial_statements = generate_financial_statements(combined_df, mapping, dataset["account_mapping"])
    breakdowns = calculate_breakdowns(combined_df, mapping)
    monthly_trend = calculate_monthly_trend(combined_df, mapping)
    anomalies = detect_anomalies(monthly_trend)
    ai_insights = generate_financial_ai_insights(financial_statements, breakdowns, monthly_trend, anomalies)

    return {
        "engagement_id": engagement_id,
        "included_sections": dataset["included_sections"],
        "financial_statements": financial_statements,
        "breakdowns": breakdowns,
        "monthly_trend": monthly_trend,
        "anomalies": anomalies,
        "ai_insights": ai_insights,
    }

# --- Report Generator helpers -----------------------------------------------
#
# The Report Generator (Month 3) reuses the same cleaned dataframe and mapping
# that the Financial Engine and AI Insights Engine already work with (Month 2),
# it just scopes them to a specific period first.

# Find the original column name in `mapping` that was mapped to a given
# standard field (e.g. "date"). Returns None if nothing is mapped to it.
def resolve_mapped_column(mapping: dict, standard_field: str) -> Optional[str]:
    for original_col, info in mapping.items():
        if isinstance(info, dict) and info.get("mapped_to") == standard_field:
            return original_col
    return None

# Parse a single date value the same flexible way validate_field_types_against_data
# does elsewhere in this file, so period filtering agrees with the cleaning report.
def _parse_flexible_date(value):
    val_str = str(value).strip()
    if not val_str:
        return pd.NaT
    dayfirst = not (len(val_str) >= 10 and val_str[4] == '-' and val_str[7] == '-')
    return pd.to_datetime(val_str, dayfirst=dayfirst, errors='coerce')

# Work out the [period_start, period_end] date bounds and a human-readable
# label from a ReportGenerateRequest's report_type and the fields relevant to it.
def resolve_report_period(req: "ReportGenerateRequest"):
    if req.report_type == "monthly":
        if not req.year or not req.month:
            raise HTTPException(status_code=400, detail="Monthly reports require both 'year' and 'month'.")
        if not (1 <= req.month <= 12):
            raise HTTPException(status_code=400, detail="'month' must be between 1 and 12.")
        period_start = pd.Timestamp(year=req.year, month=req.month, day=1)
        period_end = period_start + pd.offsets.MonthEnd(0)
        period_label = period_start.strftime("%B %Y")

    elif req.report_type == "yearly":
        if not req.year:
            raise HTTPException(status_code=400, detail="Yearly reports require 'year'.")
        period_start = pd.Timestamp(year=req.year, month=1, day=1)
        period_end = pd.Timestamp(year=req.year, month=12, day=31)
        period_label = str(req.year)

    else:  # custom
        if not req.start_date or not req.end_date:
            raise HTTPException(status_code=400, detail="Custom reports require both 'start_date' and 'end_date'.")
        
        # Validate dates using centralized date validator
        is_valid, error_msg = DateValidator.validate_date_range(
            req.start_date, req.end_date, input_format="storage"
        )
        if not is_valid:
            raise HTTPException(status_code=400, detail=error_msg)
        
        # Additional validation for individual dates
        is_valid, error_msg = DateValidator.is_valid_calendar_date(req.start_date, input_format="storage")
        if not is_valid:
            raise HTTPException(status_code=400, detail=f"Start date: {error_msg}")
        
        is_valid, error_msg = DateValidator.is_valid_calendar_date(req.end_date, input_format="storage")
        if not is_valid:
            raise HTTPException(status_code=400, detail=f"End date: {error_msg}")
        
        period_start = pd.to_datetime(req.start_date, errors="coerce")
        period_end = pd.to_datetime(req.end_date, errors="coerce")
        period_label = f"{period_start.strftime('%d %b %Y')} \u2013 {period_end.strftime('%d %b %Y')}"

    return period_start, period_end, period_label

# Filter a cleaned dataframe down to rows whose mapped date column falls
# within [period_start, period_end]. Raises if there's no date column mapped,
# or if the period contains no rows at all, since a report with silently zero
# data is worse than an explicit error.
def filter_dataframe_by_period(df: pd.DataFrame, mapping: dict, period_start, period_end) -> pd.DataFrame:
    date_col = resolve_mapped_column(mapping, "date")
    if not date_col or date_col not in df.columns:
        raise HTTPException(
            status_code=400,
            detail="No column is mapped to 'date' for this file, so reports cannot be scoped to a period. "
                   "Please map a date column first."
        )
    parsed_dates = df[date_col].apply(_parse_flexible_date)
    mask = (parsed_dates >= period_start) & (parsed_dates <= period_end)
    filtered = df.loc[mask]
    if filtered.empty:
        raise HTTPException(
            status_code=400,
            detail="No transactions fall within the selected period. Try a different date range."
        )
    return filtered

# Describe the charts the frontend dashboard (Month 2) should render for this
# report, without pre-rendering images here — actual PDF/image chart embedding
# belongs to the Export System, built separately.
def build_chart_specs(breakdowns: dict, monthly_trend: dict) -> list:
    """
    Dynamically generates chart specifications based on available data.
    Only includes chart types that have corresponding data available,
    avoiding errors when columns are empty or data is missing.
    """
    chart_specs = []
    
    # Only add monthly trend charts if we have trend data
    if monthly_trend and len(monthly_trend) > 0:
        # Find the first revenue-like trend if available
        revenue_trend_key = None
        for key in monthly_trend.keys():
            if 'revenue' in key.lower() or 'income' in key.lower():
                revenue_trend_key = key
                break
        
        if revenue_trend_key:
            chart_specs.append({
                "type": "revenue_line", 
                "title": "Revenue Trend", 
                "source": f"monthly_trend.{revenue_trend_key}"
            })
        else:
            # Fall back to any available trend
            first_trend_key = list(monthly_trend.keys())[0]
            chart_specs.append({
                "type": "trend_line", 
                "title": f"{first_trend_key.replace('_by_', ' ').replace('_month', '').title()} Trend", 
                "source": f"monthly_trend.{first_trend_key}"
            })
    
    # Only add breakdown charts if we have breakdown data
    if breakdowns and len(breakdowns) > 0:
        # Look for expense-like breakdowns
        expense_breakdown_key = None
        for key in breakdowns.keys():
            if 'expense' in key.lower() or 'cost' in key.lower():
                expense_breakdown_key = key
                break
        
        if expense_breakdown_key:
            chart_specs.append({
                "type": "expense_pie", 
                "title": "Expense Breakdown", 
                "source": f"breakdowns.{expense_breakdown_key}"
            })
        
        # Look for revenue/income breakdowns
        revenue_breakdown_key = None
        for key in breakdowns.keys():
            if 'revenue' in key.lower() or 'income' in key.lower():
                revenue_breakdown_key = key
                break
        
        if revenue_breakdown_key:
            chart_specs.append({
                "type": "revenue_bar", 
                "title": "Revenue Breakdown", 
                "source": f"breakdowns.{revenue_breakdown_key}"
            })
        
        # If no specific financial breakdowns found, add the first available breakdown
        if not expense_breakdown_key and not revenue_breakdown_key:
            first_breakdown_key = list(breakdowns.keys())[0]
            chart_specs.append({
                "type": "category_bar", 
                "title": f"{first_breakdown_key.replace('_by_', ' ').title()}", 
                "source": f"breakdowns.{first_breakdown_key}"
            })
    
    return chart_specs

# Pydantic model for a client record
class Client(BaseModel):
    company_name: str
    contact_person: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    industry: Optional[str] = None
    address: Optional[str] = None
    status: Optional[str] = "Active"
    kra_pin: Literal[True, False] = False
    kra_pin_number: Optional[str] = None

    @field_validator("phone")
    @classmethod
    def phone_must_be_valid_kenyan(cls, v):
        # Optional field — skip validation entirely if not provided
        if v is None or not v.strip():
            return v
        result = validate_kenyan_phone(v)
        if not result.valid:
            raise ValueError(f"Invalid Kenyan phone number: {result.reason}")
        # Store the normalized +254 format so downstream code never has to
        # deal with mixed formats (0712..., 254712..., +254712... etc.)
        return result.e164

    @field_validator("email")
    @classmethod
    def email_must_be_real(cls, v):
        if v is None or not v.strip():
            return v
        result = validate_real_email(v, check_deliverability=False)
        if not result.valid:
            raise ValueError(f"Invalid email address: {result.reason}")
        return result.normalized

# Pydantic model for creating a user
class User(BaseModel):
    full_name: str
    email: str
    password: str
    phone: Optional[str] = None
    role: Literal["Admin", "Accountant", "Auditor", "Senior Auditor", "Assistant Manager", "Audit Manager", "Engagement Partner", "Quality Reviewer"]
    assigned_client_id: Optional[int] = None
    status: Optional[str] = "Active"

    @field_validator("phone")
    @classmethod
    def phone_must_be_valid_kenyan(cls, v):
        if v is None or not v.strip():
            return v
        result = validate_kenyan_phone(v)
        if not result.valid:
            raise ValueError(f"Invalid Kenyan phone number: {result.reason}")
        return result.e164

    @field_validator("email")
    @classmethod
    def email_must_be_real(cls, v):
        result = validate_real_email(v, check_deliverability=False)
        if not result.valid:
            raise ValueError(f"Invalid email address: {result.reason}")
        return result.normalized

    @field_validator("password")
    @classmethod
    def password_must_be_strong(cls, v):
        result = validate_password_strength(v)
        if not result.valid:
            raise ValueError("Weak password: " + " ".join(result.reasons))
        return v

# Pydantic model for updating an existing user (no password field)
class UserUpdate(BaseModel):
    full_name: str
    email: str
    phone: Optional[str] = None
    role: Literal["Admin", "Accountant", "Auditor", "Senior Auditor", "Assistant Manager", "Audit Manager", "Engagement Partner", "Quality Reviewer"]
    assigned_client_id: Optional[int] = None
    status: Optional[str] = "Active"

    @field_validator("phone")
    @classmethod
    def phone_must_be_valid_kenyan(cls, v):
        if v is None or not v.strip():
            return v
        result = validate_kenyan_phone(v)
        if not result.valid:
            raise ValueError(f"Invalid Kenyan phone number: {result.reason}")
        return result.e164

    @field_validator("email")
    @classmethod
    def email_must_be_real(cls, v):
        result = validate_real_email(v, check_deliverability=False)
        if not result.valid:
            raise ValueError(f"Invalid email address: {result.reason}")
        return result.normalized

# Pydantic model for a login request
class LoginRequest(BaseModel):
    email: str
    password: str

# Pydantic model for requesting a password reset token
class PasswordResetRequest(BaseModel):
    email: str

# Pydantic model for confirming a password reset with a token and new password
class PasswordResetConfirm(BaseModel):
    token: str
    new_password: str

# Pydantic model for a single column mapping record
class ColumnMapping(BaseModel):
    client_id: str
    file_type: Optional[str] = "general"
    original_column: str
    mapped_to: str
    confirmed_by: Optional[str] = None

# Pydantic model for an audit engagement
class Engagement(BaseModel):
    client_id: int
    engagement_name: str
    financial_year: str
    status: Optional[str] = "Planning"
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    sections: List[str]

# Pydantic model for assigning a user to an engagement team
class EngagementTeam(BaseModel):
    engagement_id: int
    user_id: int
    role: str

# Pydantic model for an audit section within an engagement
class AuditSection(BaseModel):
    engagement_id: int
    section_name: str
    status: Optional[str] = "Pending"
    assigned_to: Optional[int] = None
class SectionScopeUpdate(BaseModel):
    in_scope: bool
    scope_reason: Optional[str] = None
    set_by: int          # user_id of whoever is making the change
    confirm_override: bool = False   # must be True if section already has an Approved submissions    

# Pydantic model for a submission of work for review
class Submission(BaseModel):
    engagement_id: int
    section_id: int
    submitted_by: int
    file_id: str                          # required — every submission must reference an actual uploaded file
    sheet_name: Optional[str] = None      # set when file_id points to a multi-section workbook
    status: Optional[str] = "Draft"
    current_stage: Optional[str] = "Accountant"
    notes: Optional[str] = None
class SaveFinalAnalysis(BaseModel):
    saved_by: int
# Pydantic model for updating a submission's status and workflow stage
class SubmissionStatus(BaseModel):
    status: Literal["Draft", "Submitted", "Under Review", "Changes Requested", "Approved", "Cancelled"]
    current_stage: Optional[str] = None
    notes: Optional[str] = None
    updated_by: Optional[int] = None

# Pydantic model for an in-app notification
class Notification(BaseModel):
    user_id: int
    message: str
    type: Optional[str] = "engagement_alert"

# Pydantic model for requesting a generated report. Exactly which of
# year/month/start_date/end_date is required depends on report_type:
#   monthly -> year + month
#   yearly  -> year
#   custom  -> start_date + end_date (YYYY-MM-DD)
class ReportGenerateRequest(BaseModel):
    client_id: int
    engagement_id: int
    # Deprecated: reports now derive their dataset from every approved,
    # in-scope section (via get_engagement_combined_dataset), the same
    # source AI Insights uses — not from a single uploaded file. Kept
    # optional so older frontend calls that still send these don't break;
    # the values are no longer read by generate_report.
    file_id: Optional[str] = None
    file_type: Optional[str] = "general"
    report_type: Literal["monthly", "yearly", "custom"]
    year: Optional[int] = None
    month: Optional[int] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    commentary: Optional[str] = ""
    generated_by: Optional[int] = None

# Simple health check endpoint
@app.get("/")
def root():
    return {"message": "Audit AI API is running"}

# Upload a file for the AI pipeline (CSV, Excel, PDF, or DOCX), extract its
# contents into a dataframe, and return a preview along with column info
# Inspect Excel sheets for multi-sheet file upload
@app.post("/upload/inspect-sheets")
async def inspect_sheets(file: UploadFile = File(...)):
    ext = get_extension(file.filename)
    if ext not in ["xlsx", "xls"]:
        return {"filename": file.filename, "sheets": []}
    
    file_bytes = await file.read()
    try:
        excel_file = pd.ExcelFile(io.BytesIO(file_bytes))
        sheets = []
        for sheet_name in excel_file.sheet_names:
            try:
                df = pd.read_excel(excel_file, sheet_name=sheet_name, dtype=str)
                df = df.dropna(axis=1, how='all')
                df = df.loc[:, ~(df == '').all()]
                sheets.append({
                    "name": sheet_name,
                    "rows": len(df),
                    "cols": len(df.columns)
                })
            except Exception:
                sheets.append({
                    "name": sheet_name,
                    "rows": 0,
                    "cols": 0
                })
        return {"filename": file.filename, "sheets": sheets}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not inspect Excel sheets: {str(e)}")


# Get preview data for a specific Excel sheet
@app.post("/upload/sheet-preview")
async def get_sheet_preview(
    file: UploadFile = File(...),
    sheet_name: str = Form(...)
):
    ext = get_extension(file.filename)
    if ext not in ["xlsx", "xls"]:
        raise HTTPException(status_code=400, detail="File is not an Excel document.")
    
    file_bytes = await file.read()
    temp_path = os.path.join(UPLOAD_DIR, f"preview_{uuid.uuid4()}.xlsx")
    try:
        with open(temp_path, "wb") as buffer:
            buffer.write(file_bytes)
        excel_file = pd.ExcelFile(io.BytesIO(file_bytes))
        if sheet_name not in excel_file.sheet_names:
            raise HTTPException(status_code=400, detail=f"Sheet '{sheet_name}' not found in Excel file.")

        _, df, _ = detect_excel_header(temp_path, sheet_name, file_type=None)
        if df is None:
            raise HTTPException(status_code=400, detail=f"Could not read sheet '{sheet_name}'.")
        df = df.map(lambda x: x.strip() if isinstance(x, str) else x)
        df = df.dropna(axis=1, how='all')
        df = df.loc[:, ~(df == '').all()]
        df = df.drop(columns=[c for c in df.columns if c in RESERVED_INTERNAL_COLUMNS], errors='ignore')
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read sheet '{sheet_name}': {str(e)}")
    finally:
        # Robust file cleanup with retry mechanism for Windows file locking
        if os.path.exists(temp_path):
            for attempt in range(5):  # Try up to 5 times
                try:
                    os.remove(temp_path)
                    break
                except PermissionError:
                    if attempt < 4:  # Don't sleep on the last attempt
                        time.sleep(0.1 * (attempt + 1))  # Exponential backoff: 0.1s, 0.2s, 0.3s, 0.4s
                    else:
                        # Log warning but don't fail the request if cleanup fails
                        warnings.warn(f"Could not delete temporary file {temp_path} after 5 attempts due to file locking")
    
    return {
        "sheet_name": sheet_name,
        "rows": len(df),
        "columns": list(df.columns),
        "preview": df.head(5).fillna("").to_dict(orient="records")
    }


# Upload a file for the AI pipeline (CSV, Excel, PDF, or DOCX), extract its
# contents into a dataframe, and return a preview along with column info
@app.post("/upload")
async def upload_file_ai(
    file: UploadFile = File(...),
    client_id: str = Form(...),
    section_id: int = Form(None),
    engagement_id: int = Form(None),
    uploaded_by: int = Form(None),
    sheet_name: str = Form(None),
    selected_sheets: str = Form(None)
):
    ext = get_extension(file.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=415, detail=f"File type .{ext} not supported. Upload Excel, CSV, PDF or DOCX file only.")
    if ext in ["xlsx", "xls"]:
        try:
            sheets_list = json.loads(selected_sheets) if selected_sheets else []
        except json.JSONDecodeError:
            sheets_list = []
        if not sheets_list:
            raise HTTPException(status_code=400, detail="Select at least one sheet before uploading.")
    MAX_FILE_SIZE = 50
    file_bytes = await file.read()
    file_size_mb = len(file_bytes) / (1024 * 1024)
    if file_size_mb > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail=f"File size exceeds the maximum limit of {MAX_FILE_SIZE} MB. Uploaded file size: {file_size_mb:.2f} MB.")
    file.file.seek(0)
    file_id = str(uuid.uuid4())
    save_path = os.path.join(UPLOAD_DIR, f"{file_id}.{ext}")
    with open(save_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # Handle PDF uploads separately since extraction works differently
    if ext == "pdf":
        df, source = extract_pdf(save_path)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not extract any content from PDF.")
        save_upload(file_id, client_id, file.filename, ext, len(df), section_id, engagement_id=engagement_id, uploaded_by=uploaded_by)
        fill_rates = calculate_fill_rates(df)
        return {"file_id": file_id, "client_id": client_id, "filename": file.filename, "source": source,
                "rows": len(df), "columns": list(df.columns), "fill_rates": fill_rates,
                "preview": df.head(5).fillna("").to_dict(orient="records"), "message": f"PDF uploaded — extracted via {source}"}

    # Handle DOCX uploads separately since extraction works differently
    if ext == "docx":
        df, source = extract_docx(save_path)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not extract any content from DOCX.")
        save_upload(file_id, client_id, file.filename, ext, len(df), section_id, engagement_id=engagement_id, uploaded_by=uploaded_by)
        fill_rates = calculate_fill_rates(df)
        return {"file_id": file_id, "client_id": client_id, "filename": file.filename, "source": source,
                "rows": len(df), "columns": list(df.columns), "fill_rates": fill_rates,
                "preview": df.head(5).fillna("").to_dict(orient="records"), "message": f"DOCX uploaded — extracted via {source}"}

    # Handle CSV and Excel uploads with smart header detection
    # Use generic structural detection (file_type=None) for all uploads initially
    header_row_index = None
    try:
        if ext == "csv":
            header_row_index, df, metadata = detect_csv_header(save_path, file_type=None)
        elif ext in ["xlsx", "xls"]:
            header_row_index, df, metadata = detect_excel_header(save_path, sheet_name, file_type=None)
        else:
            df, metadata = None, {}
            
        if df is None:
            raise HTTPException(status_code=400, detail="Could not read file.")
            
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")
    
    # Clean the DataFrame
    if df is not None:
        df = df.dropna(axis=1, how='all')
        df = df.loc[:, ~(df == '').all()]
        df = df.map(lambda x: x.strip() if isinstance(x, str) else x)
        df = df.drop(columns=[c for c in df.columns if c in RESERVED_INTERNAL_COLUMNS], errors='ignore')
    
    # Store header_row_index and sheet_name in database
    save_upload(file_id, client_id, file.filename, ext, len(df), section_id, header_row_index, sheet_name, engagement_id, uploaded_by, selected_sheets)
    fill_rates = calculate_fill_rates(df)
    
    response = {
        "file_id": file_id, 
        "client_id": client_id, 
        "filename": file.filename, 
        "source": "table",
        "rows": len(df), 
        "columns": list(df.columns), 
        "fill_rates": fill_rates,
        "fingerprint": compute_schema_fingerprint(list(df.columns)),
        "preview": df.head(5).fillna("").to_dict(orient="records"), 
        "message": "File uploaded and processed successfully"
    }
    
    # Include extracted metadata if any was found
    if metadata:
        response["metadata"] = metadata
    
    return response


# Re-upload a corrected Trial Balance against the same file_id.
@app.post("/trial-balance/upload-corrected")
async def upload_corrected_trial_balance(
    file: UploadFile = File(...),
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("trial_balance"),
    corrected_by: str = Form(None),
    sheet_name: str = Form(None),
):
    ext = get_extension(file.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=415, detail=f"File type .{ext} not supported. Upload Excel, CSV, PDF or DOCX file only.")

    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM uploads WHERE file_id = %s AND client_id = %s", (file_id, client_id))
    existing_upload = cursor.fetchone()
    conn.close()
    if not existing_upload:
        raise HTTPException(status_code=404, detail="No existing Trial Balance upload was found for this file.")

    old_save_path, old_ext = locate_uploaded_file(file_id, existing_upload.get("file_type"))
    previous_fingerprint = None
    if old_save_path:
        try:
            old_df, _ = read_file_to_df(old_save_path, old_ext)
            if old_df is not None:
                previous_fingerprint = compute_schema_fingerprint(list(old_df.columns))
        except Exception:
            previous_fingerprint = None

    file_bytes = await file.read()
    if len(file_bytes) / (1024 * 1024) > 50:
        raise HTTPException(status_code=413, detail="Uploaded file exceeds the 50 MB limit.")
    file.file.seek(0)

    save_path = os.path.join(UPLOAD_DIR, f"{file_id}.{ext}")
    with open(save_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    remove_uploaded_file_variants(file_id, keep_ext=ext)

    try:
        if ext == "pdf":
            df, _ = extract_pdf(save_path)
            if df is None:
                raise HTTPException(status_code=400, detail="Could not extract any content from PDF.")
        elif ext == "docx":
            df, _ = extract_docx(save_path)
            if df is None:
                raise HTTPException(status_code=400, detail="Could not extract any content from DOCX.")
        else:
            stored_header_row = existing_upload.get("header_row_index")
            if ext == "csv":
                df = pd.read_csv(save_path, header=stored_header_row if stored_header_row is not None else 0, dtype=str)
            else:
                df = pd.read_excel(save_path, dtype=str, sheet_name=(sheet_name if sheet_name else 0), header=stored_header_row if stored_header_row is not None else 0)
            df = df.map(lambda x: x.strip() if isinstance(x, str) else x)
            df = df.dropna(axis=1, how='all')
            df = df.loc[:, ~(df == '').all()]
            df = df.drop(columns=[c for c in df.columns if c in RESERVED_INTERNAL_COLUMNS], errors='ignore')
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read corrected file: {str(e)}")

    new_columns = list(df.columns)
    new_fingerprint = compute_schema_fingerprint(new_columns)
    same_structure = bool(previous_fingerprint and previous_fingerprint == new_fingerprint)

    save_upload(file_id, client_id, file.filename, ext, len(df), existing_upload.get("section_id"))

    response = {
        "file_id": file_id,
        "client_id": client_id,
        "file_type": file_type,
        "filename": file.filename,
        "rows": len(df),
        "columns": new_columns,
        "fingerprint": new_fingerprint,
        "same_structure": same_structure,
        "preview": df.head(5).fillna("").to_dict(orient="records"),
        "message": "Corrected Trial Balance uploaded successfully.",
    }

    if not same_structure:
        clear_tb_validation_result(file_id, client_id, file_type)
        update_workflow_stage(file_id, client_id, file_type, "uploaded")
        response["next_step"] = "mapping"
        response["detail"] = "The corrected file structure is different from the original TB. Please confirm mapping before validation."
        return response

    mapping = get_file_specific_mapping(file_id, client_id, file_type)

    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    save_cleaning_snapshot(file_id, client_id, file_type, cleaned_df)

    if not report.get("can_proceed", False):
        clear_tb_validation_result(file_id, client_id, file_type)
        update_workflow_stage(file_id, client_id, file_type, "cleaning_in_progress")
        response["next_step"] = "cleaning"
        response["detail"] = "The corrected Trial Balance still has cleaning issues. Please review the cleaning result before validating again."
        response["cleaning_report"] = report
        response["cleaned_data"] = cleaned_df.fillna("").astype(str).to_dict(orient="records")
        return response

    validation_result = validate_trial_balance(cleaned_df, mapping)
    save_tb_validation_result(file_id, client_id, file_type, validation_result)
    update_workflow_stage(file_id, client_id, file_type, "tb_validation")
    response["next_step"] = "tb_validation"
    response["cleaning_report"] = report
    response["cleaned_data"] = cleaned_df.fillna("").astype(str).to_dict(orient="records")
    response["validation_result"] = validation_result
    response["detail"] = (
        "Trial balance is balanced." if validation_result.get("is_balanced") else
        "Trial balance still does not balance. Review the difference before proceeding."
    )
    return response




# Debug endpoint to check notifications for a specific user
@app.get("/test/user/{user_id}/notifications")
def check_user_notifications(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT * FROM notifications 
        WHERE user_id = %s 
        ORDER BY created_at DESC
    """, (user_id,))
    notifications = cursor.fetchall()
    return {
        "user_id": user_id,
        "notifications": notifications,
        "count": len(notifications)
    }


# Get file preview by file_id for notifications
@app.get("/file-preview/{file_id}")
def get_file_preview(file_id: str, client_id: str = Query(...), db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    
    cursor.execute("""
        SELECT file_id, filename, row_count, file_type, semantic_file_type, header_row_index, sheet_name
        FROM uploads
        WHERE file_id = %s AND client_id = %s
    """, (file_id, client_id))
    
    upload = cursor.fetchone()
    
    if not upload:
        # Return a graceful response instead of 404 for deleted/non-existent files
        return {
            "file_id": file_id,
            "filename": None,
            "rows": 0,
            "columns": [],
            "preview": [],
            "file_type": None,
            "semantic_file_type": None,
            "error": "File not found in uploads table. It may have been deleted."
        }
    
    # Try to load the file using stored canonical header configuration
    try:
        file_path = os.path.join(UPLOAD_DIR, f"{file_id}.{upload['file_type']}")
        stored_header_row = upload.get('header_row_index')
        stored_sheet_name = upload.get('sheet_name')
        
        # Use read_file_to_df with stored header configuration to ensure canonical consistency
        df, _ = read_file_to_df(file_path, upload['file_type'], stored_sheet_name, file_type=None, header_row_index=stored_header_row)
        
        if df is None:
            raise Exception("Could not read file using stored header configuration")
        
        result = {
            "file_id": upload['file_id'],
            "filename": upload['filename'],
            "rows": len(df),
            "columns": list(df.columns),
            "preview": df.head(5).fillna("").to_dict(orient="records"),
            "file_type": upload['file_type'],
            "semantic_file_type": upload.get('semantic_file_type')
        }
        return result
    except Exception as e:
        return {
            "file_id": upload['file_id'],
            "filename": upload['filename'],
            "rows": upload['row_count'],
            "columns": [],
            "preview": [],
            "file_type": upload['file_type'],
            "semantic_file_type": upload.get('semantic_file_type'),
            "error": str(e)
        }

# Simple endpoint to list all clients with their IDs
@app.get("/test/clients")
def list_all_clients(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT client_id, company_name FROM clients")
    clients = cursor.fetchall()
    return {"clients": clients}

# Test endpoint to check engagement and team setup for a client
@app.get("/test/client/{client_id}/engagement")
def test_client_engagement(client_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    
    # Get engagement and team in a single query
    cursor.execute("""
        SELECT 
            e.engagement_id, 
            e.engagement_name,
            u.user_id,
            u.full_name,
            u.role
        FROM engagements e
        LEFT JOIN engagement_team et ON e.engagement_id = et.engagement_id
        LEFT JOIN users u ON et.user_id = u.user_id
        WHERE e.client_id = %s
    """, (client_id,))
    
    results = cursor.fetchall()
    
    if not results or not results[0]['engagement_id']:
        return {"error": "No engagement found for this client", "client_id": client_id}
    
    engagement = {
        "engagement_id": results[0]['engagement_id'],
        "engagement_name": results[0]['engagement_name']
    }
    
    team = []
    for row in results:
        if row['user_id']:
            team.append({
                "user_id": row['user_id'],
                "full_name": row['full_name'],
                "role": row['role']
            })
    
    return {
        "engagement": engagement,
        "team": team,
        "team_count": len(team)
    }

# Submit an uploaded file for auditor review - creates notification for assigned auditor
@app.post("/upload/submit")
async def submit_uploaded_file(
    file_id: str = Form(...),
    client_id: str = Form(...),
    submitted_by: int = Form(...),
    db=Depends(get_db)
):
    cursor = db.cursor(dictionary=True)
    
    # Find the engagement for this client
    cursor.execute("""
        SELECT engagement_id, engagement_name 
        FROM engagements 
        WHERE client_id = %s 
        LIMIT 1
    """, (int(client_id),))
    
    engagement = cursor.fetchone()
    print(f"DEBUG: Engagement found: {engagement}")
    
    if engagement:
        # Find auditors in the engagement team (same approach as engagement notifications)
        cursor.execute("""
            SELECT u.user_id, u.full_name 
            FROM users u
            INNER JOIN engagement_team et ON u.user_id = et.user_id
            WHERE et.engagement_id = %s 
            AND u.role = 'Auditor'
        """, (engagement['engagement_id'],))
        
        auditors = cursor.fetchall()
        print(f"DEBUG: Auditors found: {auditors}")
        
        if auditors:
            # Create notification for each auditor in the team with file details
            message = f"New file submitted. Please review and then proceed to the Detect phase."
            for auditor in auditors:
                print(f"DEBUG: Creating notification for user_id={auditor['user_id']}")
                cursor.execute("""
                    INSERT INTO notifications (user_id, message, type, file_id, client_id) 
                    VALUES (%s, %s, %s, %s, %s)
                """, (auditor['user_id'], message, 'file_submission', file_id, client_id))
            
            # Update upload status
            cursor.execute("""
                UPDATE uploads SET status = 'Submitted' 
                WHERE file_id = %s AND client_id = %s
            """, (file_id, client_id))
            
            db.commit()
            print(f"DEBUG: Successfully submitted and notified {len(auditors)} auditors")
            return {"message": "File submitted successfully", "notified_auditors": len(auditors)}
    
    print(f"DEBUG: No engagement or auditors found")
    return {"message": "File submitted but no engagement or auditor found for this client"}


# Detect what each column in an uploaded file means, reusing a saved mapping or
# fingerprint cache when possible, and falling back to LLM detection otherwise
@app.post("/detect-columns")
async def detect_columns_endpoint(
    client_id: str = Form(...),
    file_id: str = Form(...),
    columns: str = Form(...),
    file_type: str = Form("general"),
    fill_rates: str = Form("{}"),
    fingerprint: str = Form("")
):
    try:
        columns_list = json.loads(columns)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid columns format.")

    try:
        fill_rates_dict = json.loads(fill_rates)
    except json.JSONDecodeError:
        fill_rates_dict = {}

    # Locate the previously uploaded file on disk
    save_path = None
    file_ext = None
    for extension in ALLOWED_EXTENSIONS:
        path = os.path.join(UPLOAD_DIR, f"{file_id}.{extension}")
        if os.path.exists(path):
            save_path = path
            file_ext = extension
            break
    if not save_path:
        raise HTTPException(status_code=404, detail="File not found. Please upload the file first.")

    # Get stored header configuration from uploads table
    upload_record = get_upload(file_id)
    stored_header_row = upload_record.get('header_row_index') if upload_record else None
    stored_sheet_name = upload_record.get('sheet_name') if upload_record else None

    # Read the file using stored header configuration (not re-detecting)
    try:
        df, metadata = read_file_to_df(save_path, file_ext, stored_sheet_name, file_type=None, header_row_index=stored_header_row)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not read file.")
        
        sample_values = {}
        for col in columns_list:
            if col in df.columns:
                non_empty = df[col].dropna().replace("", float("nan")).dropna()
                sample_values[col] = str(non_empty.iloc[0]) if len(non_empty) > 0 else ""
            else:
                sample_values[col] = ""
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")

    # Suggest a file type category based on the columns and sample values
    try:
        file_type_suggestion = suggest_file_type(columns_list, sample_values)
    except Exception:
        file_type_suggestion = {"file_type": "other", "file_type_label": "Other"}
    effective_file_type = file_type_suggestion["file_type"]

    # Save this schema's fingerprint so future uploads with the same columns can be recognized
    computed_fingerprint = compute_schema_fingerprint(columns_list)
    save_fingerprint(client_id, computed_fingerprint, effective_file_type, columns_list)
    
    # First check for file-specific mapping (file_id-based isolation)
    file_specific_mapping = get_mapping(client_id, effective_file_type, None, file_id)
    if file_specific_mapping:
        # File-specific mapping found - use it
        filtered_mapping = {
            col: file_specific_mapping.get(col)
            for col in columns_list
            if col in file_specific_mapping
        }
        if filtered_mapping and len(filtered_mapping) == len(columns_list):
            result = build_detection_result(columns_list, filtered_mapping, sample_values, fill_rates_dict)
            result["file_id"] = file_id
            result["source"] = "file_specific"
            result["message"] = "Loaded mapping for this specific file."
            result["suggested_file_type"] = file_type_suggestion["file_type"]
            result["suggested_file_type_label"] = file_type_suggestion["file_type_label"]
            # Include metadata if it was extracted during header detection
            if metadata:
                result["metadata"] = metadata
            return result
    
    # Fall back to fingerprint-based cache (kept for future explicit reuse feature)
    saved_mapping = get_mapping(client_id, effective_file_type, computed_fingerprint)
    existing_fp = get_fingerprint(client_id, computed_fingerprint, effective_file_type)

    # Fingerprint cache hit: only trust this if every uploaded column actually
    # has a matching entry in the saved mapping, otherwise fall through to detection
    if existing_fp and saved_mapping:
        filtered_mapping = {
            col: saved_mapping.get(col)
            for col in columns_list
            if col in saved_mapping
        }
        if filtered_mapping and len(filtered_mapping) == len(columns_list):
            result = build_detection_result(columns_list, filtered_mapping, sample_values, fill_rates_dict)
            result["file_id"] = file_id
            result["source"] = "fingerprint_cache"
            result["message"] = "Mapping reused from cache."
            result["suggested_file_type"] = file_type_suggestion["file_type"]
            result["suggested_file_type_label"] = file_type_suggestion["file_type_label"]
            # Include metadata if it was extracted during header detection
            if metadata:
                result["metadata"] = metadata
            return result

    # Saved mapping hit: only trust this if every uploaded column is covered by the saved mapping
    if saved_mapping:
        all_mapped = all(col in saved_mapping for col in columns_list)
        if all_mapped:
            filtered_mapping = {col: saved_mapping[col] for col in columns_list}
            result = build_detection_result(columns_list, filtered_mapping, sample_values, fill_rates_dict)
            result["file_id"] = file_id
            result["source"] = "saved_mapping"
            result["message"] = "This mapping was loaded from this client's saved profile. Review, change and confirm as needed."
            # Return the actual saved file_type instead of suggestion
            result["suggested_file_type"] = effective_file_type
            result["suggested_file_type_label"] = FILE_TYPE_CATEGORIES.get(effective_file_type, effective_file_type)
            # Include metadata if it was extracted during header detection
            if metadata:
                result["metadata"] = metadata
            return result

    # No usable cache found: run LLM detection on the columns
    try:
        mapping = detect_columns_with_llm(columns_list, sample_values, fill_rates_dict)
        if not mapping:
            raise HTTPException(status_code=500, detail="LLM returned empty mapping.")

        result = build_detection_result(columns_list, mapping, sample_values, fill_rates_dict)

        # Demote any column that was mapped to the same target as another column,
        # since two columns sharing one mapped name would crash cleaning later
        dedup_warnings = []
        seen_targets = set()
        for original_col, info in list(result["mapping"].items()):
            if not isinstance(info, dict):
                continue
            target = str(info.get("mapped_to", "")).strip()
            if target in ("", "unknown"):
                continue
            if target in seen_targets:
                suggestion = target
                result["mapping"][original_col] = {
                    "mapped_to": "unknown",
                    "field_type": "unknown",
                    "suggestion": suggestion,
                    "sample_value": info.get("sample_value", ""),
                    "fill_rate": info.get("fill_rate", 1.0),
                }
                dedup_warnings.append(
                    f"'{original_col}' was also mapped to '{target}' — demoted to unknown. "
                    f"Please give it a unique 'Mapped To' name."
                )
            else:
                seen_targets.add(target)

        # Count how many columns still need manual review
        unknown_columns = [
            col for col, info in result["mapping"].items()
            if info.get("mapped_to") == "unknown" or info.get("field_type") == "unknown"
        ]
        result["unknown_count"] = len(unknown_columns)
        result["requires_manual_mapping"] = len(unknown_columns) > 0
        if dedup_warnings:
            if "warnings" in result and isinstance(result["warnings"], list):
                result["warnings"].extend(
                    {"type": "dedup_mapping", "message": w, "action": "Please update the mapped-to field to a unique name."}
                    for w in dedup_warnings
                )
            else:
                result["warnings"] = [
                    {"type": "dedup_mapping", "message": w, "action": "Please update the mapped-to field to a unique name."}
                    for w in dedup_warnings
                ]
        if result.get("warnings"):
            result["warning_count"] = len(result["warnings"])

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Detection failed: {str(e)}")
    result["file_id"] = file_id
    result["source"] = "llm_detection"
    result["suggested_file_type"] = file_type_suggestion["file_type"]
    result["suggested_file_type_label"] = file_type_suggestion["file_type_label"]
    
    # Include metadata if it was extracted during header detection
    if metadata:
        result["metadata"] = metadata
    
    return result

# Save a confirmed column mapping for a client and file type
@app.post("/save-mapping")
async def save_mapping_endpoint(
    client_id: str = Form(...),
    file_type: str = Form(...),
    mapping: str = Form(...),
    file_id: str = Form(None),
    confirmed_by: str = Form(None),
    fingerprint: str = Form(None)
):
    try:
        mapping_dict = json.loads(mapping)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid mapping format.")
    if not mapping_dict:
        raise HTTPException(status_code=400, detail="Mapping cannot be empty. Please provide a valid mapping.")
    
    # Validate for duplicate target fields before saving
    is_valid, validation_error = validate_mapping_before_save(mapping_dict)
    if not is_valid:
        raise HTTPException(status_code=400, detail=validation_error)

    # Accounting-aware header validation: if file_type is trial_balance or general_ledger,
    # run accounting-aware detection to validate the stored header row
    parsing_warnings = []
    if file_id and file_type in ["trial_balance", "general_ledger"]:
        save_path, file_ext = locate_uploaded_file(file_id)
        if save_path and file_ext in ["xlsx", "xls"]:
            upload_record = get_upload(file_id)
            stored_header_row = upload_record.get('header_row_index') if upload_record else None
            stored_sheet_name = upload_record.get('sheet_name') if upload_record else None
            
            # Run accounting-aware detection to see if it suggests a different header
            try:
                accounting_header_row, _, _ = detect_excel_header(save_path, stored_sheet_name, file_type=file_type)
                if accounting_header_row != stored_header_row:
                    parsing_warnings.append(
                        f"Accounting-aware detection suggests header row {accounting_header_row} "
                        f"but stored configuration uses row {stored_header_row}. "
                        f"Review the table structure for correctness."
                    )
            except Exception:
                pass  # Don't fail save if validation detection fails

    # If we have access to the original file, validate that each column's
    # declared field_type is consistent with its actual data before saving
    # Use stored header configuration for consistency
    if file_id:
        save_path, file_ext = locate_uploaded_file(file_id)
        if save_path:
            upload_record = get_upload(file_id)
            stored_header_row = upload_record.get('header_row_index') if upload_record else None
            stored_sheet_name = upload_record.get('sheet_name') if upload_record else None
            
            try:
                df, _ = read_file_to_df(save_path, file_ext, stored_sheet_name, file_type=None, header_row_index=stored_header_row)
                if df is not None:
                    problems = validate_field_types_against_data(df, mapping_dict)
                    if problems:
                        raise HTTPException(
                            status_code=400,
                            detail="Field type mismatch detected: " + " | ".join(problems)
                        )
            except HTTPException:
                raise
            except Exception:
                # If validation itself fails for an unexpected reason, don't block
                # saving entirely — this check is a safety net, not the source of truth
                pass

    save_mapping(client_id, file_type, mapping_dict, confirmed_by, fingerprint, file_id)
    
    # Also update the uploads table with the semantic file_type for future reference
    if file_id:
        db = get_connection()
        cursor = db.cursor()
        cursor.execute("UPDATE uploads SET semantic_file_type = %s WHERE file_id = %s", (file_type, file_id))
        db.commit()
        cursor.close()
        db.close()
    
    # Update workflow stage to 'mapped' when mapping is saved
    if file_id:
        update_workflow_stage(file_id, client_id, file_type, "mapped")
    
    response = {"client_id": client_id, "file_type": file_type, "columns_saved": len(mapping_dict),
                "message": f"Mapping saved successfully for client {client_id} and file type {file_type}."}
    
    if parsing_warnings:
        response["parsing_warnings"] = parsing_warnings
    
    return response

# Retrieve a previously saved column mapping for a client and file type
@app.get("/get-mapping/{client_id}")
async def get_mapping_endpoint(client_id: str, file_type: str = "general"):
    mapping = get_mapping(client_id, file_type)
    if not mapping:
        return {"client_id": client_id, "file_type": file_type, "mapping": {},
                "message": "No saved mapping found for this client."}
    return {"client_id": client_id, "file_type": file_type, "mapping": mapping,
            "columns_mapped": len(mapping), "message": "Saved mapping retrieved successfully."}

# List all uploads recorded for a given client
@app.get("/uploads/{client_id}")
async def get_uploads_endpoint(client_id: str):
    uploads = get_uploads(client_id)
    return {"client_id": client_id, "total_uploads": len(uploads), "uploads": uploads}


# Run the cleaning engine on a file and return the cleaned data along with the validation report
@app.post("/clean")
async def clean_file(
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general")
) -> dict:
    # Get the fingerprint for this specific file to ensure file-specific mapping retrieval
    save_path, file_ext = locate_uploaded_file(file_id)
    if not save_path:
        raise HTTPException(status_code=404, detail="File not found. Please upload the file first.")
    
    # Get stored header configuration from uploads table
    upload_record = get_upload(file_id)
    stored_header_row = upload_record.get('header_row_index') if upload_record else None
    stored_sheet_name = upload_record.get('sheet_name') if upload_record else None
    
    try:
        df, _ = read_file_to_df(save_path, file_ext, stored_sheet_name, file_type=None, header_row_index=stored_header_row)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not read file.")
        fingerprint = compute_schema_fingerprint(list(df.columns))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not compute file fingerprint: {str(e)}")
    
    # Use file_id for file-specific mapping retrieval
    mapping = get_mapping(client_id, file_type, None, file_id)
    if not mapping:
        raise HTTPException(status_code=400, detail="No saved mapping found for this specific file. Please detect the columns and confirm the mapping first.")
    
    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    return {"file_id": file_id, "client_id": client_id, "file_type": file_type,
            "cleaned_data": cleaned_df.fillna("").astype(str).map(lambda x: x.strip()).to_dict(orient="records"),
            "validation_report": report,
            "can_proceed": report.get("can_proceed", False),
            "message": "File cleaned successfully."}

# Export the cleaned data as a downloadable Excel workbook, and save a snapshot
# of it so a later corrected upload can be compared against this exact state
@app.get("/clean/export-cleaned/{file_id}")
async def export_cleaned_workbook(file_id: str, client_id: str, file_type: str = "general"):
    mapping = get_file_specific_mapping(file_id, client_id, file_type)
    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    save_cleaning_snapshot(file_id, client_id, file_type, cleaned_df)
    workbook_buffer = build_cleaning_workbook(cleaned_df, report, mapping)
    download_filename = f"{file_id}_cleaned_data.xlsx"
    return StreamingResponse(
        workbook_buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{download_filename}"'}
    )

# Accept an auditor-edited Excel file, diff it against the last downloaded
# snapshot, and apply any corrections, deletions, or renames that were found
@app.post("/clean/submit-corrected-excel")
async def submit_corrected_excel(
    file: UploadFile = File(...),
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general"),
    corrected_by: str = Form(None),
):
    mapping = get_file_specific_mapping(file_id, client_id, file_type)

    snapshot_rows = get_cleaning_snapshot(file_id, client_id, file_type)
    if not snapshot_rows:
        raise HTTPException(
            status_code=400,
            detail="No downloaded snapshot found for this file. Please download the cleaned Excel first, edit it, then upload it back."
        )

    # Save the uploaded file to a temporary path so it can be read and diffed
    temp_path = os.path.join(UPLOAD_DIR, f"corrected_{uuid.uuid4()}.xlsx")
    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    try:
        diff_result = diff_uploaded_against_snapshot(temp_path, snapshot_rows)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    finally:
        # Clean up the temp file, but never let a cleanup failure crash the request
        try:
            if os.path.exists(temp_path):
                os.remove(temp_path)
        except (PermissionError, OSError):
            pass
    corrections = diff_result["corrections"]
    deleted_row_ids = diff_result["deleted_row_ids"]
    deleted_columns = diff_result["deleted_columns"]
    renamed_columns = diff_result["renamed_columns"]
    unresolved_columns = diff_result.get("unresolved_columns", [])

    # Build a warning for any new column the diff couldn't confidently explain,
    # without blocking the rest of the upload from being applied
    unresolved_warning = None
    if unresolved_columns:
        names = ", ".join(f"'{c}'" for c in unresolved_columns)
        unresolved_warning = (
            f"Could not confidently determine what happened to {names} — this usually "
            f"means either it's a genuinely new column we don't recognize, or its rename "
            f"couldn't be told apart from another change in this upload with confidence. "
            f"Everything else in this upload was applied. If {names} was meant to be a "
            f"rename, please try renaming just that one column by itself in a follow-up "
            f"upload."
        )

    # Check that the proposed renames don't collide with each other, or with a
    # different column's existing mapped name, before applying anything
    if renamed_columns:
        new_name_targets = {}
        for old_name, new_name in renamed_columns.items():
            new_name_targets.setdefault(new_name, []).append(old_name)
        mutual_collisions = {n: olds for n, olds in new_name_targets.items() if len(olds) > 1}
        if mutual_collisions:
            descriptions = "; ".join(
                f"{', '.join(repr(o) for o in olds)} were all renamed to '{n}'"
                for n, olds in mutual_collisions.items()
            )
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Cannot apply these renames: {descriptions}. Each column needs a "
                    f"unique name. Please give each one a different name and re-upload."
                )
            )

        for old_name, new_name in renamed_columns.items():
            colliding_column = next(
                (
                    original_col for original_col, info in mapping.items()
                    if isinstance(info, dict)
                    and info.get("mapped_to") == new_name
                    and info.get("mapped_to") != old_name
                ),
                None,
            )
            if colliding_column:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        f"Cannot rename '{old_name}' to '{new_name}': that name is already used "
                        f"by column '{colliding_column}'. Please choose a different name for "
                        f"one of these two columns and re-upload."
                    )
                )

    # If nothing changed at all in the uploaded file, there's nothing to apply
    if not corrections and not deleted_row_ids and not deleted_columns and not renamed_columns:
        raise HTTPException(
            status_code=400,
            detail="No changes were detected in the uploaded file compared to what was downloaded."
        )

    # Validate each cell correction against its column's declared field_type
    # before saving anything, same rule as manual inline corrections: reject
    # bad numeric/date input instead of silently writing it through.
    if corrections:
        validation_errors = []
        for correction in corrections:
            column = correction.get("column")
            value = correction.get("corrected_value", "")
            col_info = mapping.get(column)
            if not isinstance(col_info, dict):
                continue
            field_type = col_info.get("field_type")

            if value is None or str(value).strip() == "":
                continue

            if field_type == "numeric":
                try:
                    FieldValidator.validate_financial_amount(value, field_name=column)
                except FieldValidationError as e:
                    validation_errors.append(f"Row {correction.get('row_index')}, '{column}': {e.message}")

            elif field_type == "date":
                if pd.isna(_parse_flexible_date(str(value))):
                    validation_errors.append(
                        f"Row {correction.get('row_index')}, '{column}': "
                        f"'{value}' is not a valid date."
                    )

        if validation_errors:
            raise HTTPException(
                status_code=422,
                detail={
                    "message": "The uploaded file contains invalid values in one or more cells.",
                    "errors": validation_errors,
                },
            )

    # Save plain cell value corrections
    if corrections:
        save_cleaning_corrections(file_id, client_id, file_type, corrections, corrected_by)

    # Save row deletions
    if deleted_row_ids:
        deletion_records = [
            {"row_index": row_id, "column": "_row_deleted",
             "original_value": "present", "corrected_value": "deleted_by_auditor"}
            for row_id in deleted_row_ids
        ]
        save_cleaning_corrections(file_id, client_id, file_type, deletion_records, corrected_by)

    # Save column deletions
    if deleted_columns:
        column_deletion_records = [
            {"row_index": -1, "column": f"_column_deleted:{col}",
             "original_value": "present", "corrected_value": "deleted_by_auditor"}
            for col in deleted_columns
        ]
        save_cleaning_corrections(file_id, client_id, file_type, column_deletion_records, corrected_by)

    # Apply every confidently-resolved rename to the saved mapping. old_name from
    # the diff may be either the original raw column key (if it was unresolved)
    # or the mapped_to value (if it was already resolved), so the actual mapping key has to be looked up before being updated
    if renamed_columns:
        updated_mapping = dict(mapping)
        for old_name, new_name in renamed_columns.items():
            original_key = next(
                (oc for oc, info in updated_mapping.items()
                 if isinstance(info, dict) and info.get("mapped_to") == old_name),
                None,
            )
            if original_key is None:
                original_key = old_name

            existing_info = updated_mapping.get(original_key, {})
            updated_mapping[original_key] = {
                **(existing_info if isinstance(existing_info, dict) else {}),
                "mapped_to": new_name,
                "field_type": (
                    existing_info.get("field_type")
                    if isinstance(existing_info, dict) and existing_info.get("field_type") not in (None, "unknown")
                    else "text"
                ),
            }
        save_mapping(client_id, file_type, updated_mapping, corrected_by, None)
        mapping = updated_mapping

    # Re-clean the file with all corrections applied, then refresh the snapshot
    # so the next diff round compares against this latest state
    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    save_cleaning_snapshot(file_id, client_id, file_type, cleaned_df)

    if renamed_columns:
        rename_descriptions = "; ".join(f"'{o}' renamed to '{n}'" for o, n in renamed_columns.items())
        rename_summary = f", {rename_descriptions} in the mapping"
    else:
        rename_summary = ""
    response = {
        "file_id": file_id,
        "client_id": client_id,
        "file_type": file_type,
        "cleaned_data": cleaned_df.fillna("").astype(str).map(lambda x: x.strip()).to_dict(orient="records"),
        "validation_report": report,
        "can_proceed": report.get("can_proceed", False),
        "corrections_applied": len(corrections),
        "rows_deleted": deleted_row_ids,
        "columns_deleted": deleted_columns,
        "columns_renamed": renamed_columns,
        "unresolved_columns": unresolved_columns,
        "message": (
            f"{len(corrections)} correction(s), {len(deleted_row_ids)} row deletion(s), "
            f"and {len(deleted_columns)} column deletion(s) applied from the uploaded file"
            f"{rename_summary}. File re-cleaned successfully."
        )
    }
    if unresolved_warning:
        response["warning"] = unresolved_warning
    return response

# Record that an auditor has accepted an issue as valid as-is, then re-clean the file
@app.post("/clean/acknowledge-issue")
async def acknowledge_issue_endpoint(
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general"),
    issue: str = Form(...),
    acknowledged_by: str = Form(None),
):
    mapping = get_file_specific_mapping(file_id, client_id, file_type)
    try:
        issue_dict = json.loads(issue)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid issue format.")
    issue_id = issue_dict.get("issue_id") or issue_fingerprint(file_id, client_id, file_type, issue_dict)
    save_cleaning_acknowledgment(issue_id, file_id, client_id, file_type, issue_dict, acknowledged_by)
    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    return {
        "file_id": file_id, "client_id": client_id, "file_type": file_type,
        "cleaned_data": cleaned_df.fillna("").astype(str).map(lambda x: x.strip()).to_dict(orient="records"),
        "validation_report": report,
        "can_proceed": report.get("can_proceed", False),
        "message": "Issue acknowledged."
    }

# Save inline cell edits made directly in the web app (not via Excel re-upload), then re-clean
@app.post("/clean/submit-inline-corrections")
async def submit_inline_corrections_endpoint(
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general"),
    corrections: str = Form(...),
    corrected_by: str = Form(None),
):
    mapping = get_file_specific_mapping(file_id, client_id, file_type)
    try:
        corrections_list = json.loads(corrections)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid corrections format.")
    if not corrections_list:
        raise HTTPException(status_code=400, detail="No corrections provided.")

    # Validate each correction against its column's declared field_type
    # before saving — reject bad numeric/date input instead of silently
    # writing it to the cleaned data.
    validation_errors = []
    for correction in corrections_list:
        column = correction.get("column")
        value = correction.get("corrected_value", "")
        col_info = mapping.get(column)
        if not isinstance(col_info, dict):
            continue  # unmapped column — nothing to validate against
        field_type = col_info.get("field_type")

        if value is None or str(value).strip() == "":
            continue  # clearing a cell is always allowed

        if field_type == "numeric":
            try:
                FieldValidator.validate_financial_amount(value, field_name=column)
            except FieldValidationError as e:
                validation_errors.append(f"Row {correction.get('row_index')}, '{column}': {e.message}")

        elif field_type == "date":
            if pd.isna(_parse_flexible_date(str(value))):
                validation_errors.append(
                    f"Row {correction.get('row_index')}, '{column}': "
                    f"'{value}' is not a valid date."
                )

    if validation_errors:
        raise HTTPException(
            status_code=422,
            detail={"message": "One or more corrections are invalid.", "errors": validation_errors},
        )

    save_cleaning_corrections(file_id, client_id, file_type, corrections_list, corrected_by)
    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    return {
        "file_id": file_id, "client_id": client_id, "file_type": file_type,
        "cleaned_data": cleaned_df.fillna("").astype(str).map(lambda x: x.strip()).to_dict(orient="records"),
        "validation_report": report,
        "can_proceed": report.get("can_proceed", False),
        "message": f"{len(corrections_list)} correction(s) saved and re-cleaned."
    }
# List every file for a client that has a cleaned version saved, most recently updated first. Lets the auditor come back later and see what's ready.
@app.get("/clients/{client_id}/cleaned-files")
async def list_cleaned_files(client_id: str):
    files = get_cleaned_files_for_client(client_id)
    return {"client_id": client_id, "total": len(files), "files": files}

# Download the current cleaned version of a file directly, plain and clean.
@app.get("/cleaned-files/{file_id}/download")
async def download_cleaned_file(file_id: str, client_id: str, file_type: str = "general"):
    record = get_cleaned_file_data(file_id, client_id, file_type)
    if not record:
        raise HTTPException(
            status_code=404,
            detail="No cleaned version found for this file. Please run cleaning first."
        )

    df = pd.DataFrame(record["cleaned_data"])
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Cleaned Data")
    output.seek(0)
    # Use the original uploaded filename if available, otherwise default to a generic name
    download_filename = record.get("filename") or f"{file_id}_cleaned.xlsx"
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{download_filename}"'}
    )

# Return the cleaned data as JSON for viewing inline in the browser
@app.get("/cleaned-files/{file_id}/view")
async def view_cleaned_file(file_id: str, client_id: str, file_type: str = ""):
    record = get_cleaned_file_data(file_id, client_id, file_type)
    if not record:
        raise HTTPException(
            status_code=404,
            detail="No cleaned version found for this file. Please run cleaning first."
        )
    return {
        "file_id": file_id,
        "client_id": client_id,
        "file_type": file_type,
        "filename": record.get("filename"),
        "cleaned_data": record["cleaned_data"],
        "total_issues": record.get("total_issues", 0),
        "can_proceed": bool(record.get("can_proceed")),
        "updated_at": str(record.get("updated_at")),
    }

# Standardize a near-duplicate value across an entire column in one action.
# Finds every row where the column matches from_value and replaces it with
# to_value, saving one correction record per affected row, then re-cleans.
@app.post("/clean/standardize-value")
async def standardize_value_endpoint(
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general"),
    column: str = Form(...),
    from_value: str = Form(...),
    to_value: str = Form(...),
    corrected_by: str = Form(None),
):
    mapping = get_file_specific_mapping(file_id, client_id, file_type)

    save_path, file_ext = locate_uploaded_file(file_id)
    if not save_path:
        raise HTTPException(status_code=404, detail="File not found. Please upload the file first.")
    
    # Get stored header configuration from uploads table
    upload_record = get_upload(file_id)
    stored_header_row = upload_record.get('header_row_index') if upload_record else None
    stored_sheet_name = upload_record.get('sheet_name') if upload_record else None

    try:
        df, _ = read_file_to_df(save_path, file_ext, stored_sheet_name, file_type=None, header_row_index=stored_header_row)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not read file.")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")

    # Apply any previously saved corrections first, so we're matching against
    # the current state of the data, not the raw original file
    existing_corrections = get_cleaning_corrections(file_id, client_id, file_type)
    df = apply_saved_corrections(df, mapping, existing_corrections)

    # Translate the standard column name back to the original column name in df
    standard_to_original = {
        info.get("mapped_to"): original_col
        for original_col, info in mapping.items()
        if isinstance(info, dict) and info.get("mapped_to") not in ("", "unknown", None)
    }
    source_col = standard_to_original.get(column, column)

    if source_col not in df.columns:
        raise HTTPException(status_code=400, detail=f"Column '{column}' not found in file.")

    # Find every row where this column currently matches from_value
    matching_rows = df.index[df[source_col].astype(str).str.strip() == from_value].tolist()

    if not matching_rows:
        raise HTTPException(
            status_code=400,
            detail=f"No rows found with value '{from_value}' in column '{column}'. It may have already been corrected."
        )

    # Validate the replacement value against this column's declared field_type
    # before applying it to every matching row.
    field_type = mapping.get(source_col, {}).get("field_type") if isinstance(mapping.get(source_col), dict) else None
    if to_value and str(to_value).strip():
        if field_type == "numeric":
            try:
                FieldValidator.validate_financial_amount(to_value, field_name=column)
            except FieldValidationError as e:
                raise HTTPException(status_code=422, detail=e.message)
        elif field_type == "date":
            if pd.isna(_parse_flexible_date(str(to_value))):
                raise HTTPException(
                    status_code=422,
                    detail=f"'{to_value}' is not a valid date for column '{column}'.",
                )

    # Create one correction record per affected row
    standardize_records = [
        {
            "row_index": int(row_idx),
            "column": column,
            "original_value": from_value,
            "corrected_value": to_value,
        }
        for row_idx in matching_rows
    ]
    save_cleaning_corrections(file_id, client_id, file_type, standardize_records, corrected_by)

    # Re-clean with the standardization applied
    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    save_cleaning_snapshot(file_id, client_id, file_type, cleaned_df)

    return {
        "file_id": file_id,
        "client_id": client_id,
        "file_type": file_type,
        "cleaned_data": cleaned_df.fillna("").astype(str).map(lambda x: x.strip()).to_dict(orient="records"),
        "validation_report": report,
        "can_proceed": report.get("can_proceed", False),
        "rows_updated": len(matching_rows),
        "message": f"Standardized {len(matching_rows)} row(s) in '{column}' from '{from_value}' to '{to_value}'."
    }

# Run the Financial Engine on a client's cleaned data using pandas calculations
@app.post("/analyze/{client_id}")
async def analyze_financials(
    client_id: str,
    file_id: str = Form(...),
    file_type: str = Form("general")
):
    mapping = get_file_specific_mapping(file_id, client_id, file_type)
    # Run the cleaning cycle to ensure the data is up-to-date and valid before analysis
    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    if not report.get("can_proceed", False):
        raise HTTPException(
            status_code=400,
            detail="This file still has unresolved issues. Please finish cleaning before running analysis."
        )
    analysis_context = build_financial_analysis_context(cleaned_df, mapping, client_id, file_type)

    return {
        "client_id": client_id,
        "file_id": file_id,
        "file_type": file_type,
        **analysis_context,
        "message": "Financial analysis completed successfully."
    }

# Run the AI Insights Engine on already-calculated financial data.
@app.post("/analyze/{client_id}/insights")
async def analyze_insights(
    client_id: str,
    file_id: str = Form(...),
    file_type: str = Form("general")
):
    mapping = get_file_specific_mapping(file_id, client_id, file_type)
    # Run the cleaning cycle to ensure the data is up-to-date and valid before generating insights
    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    # Check if the cleaned data is valid and ready for analysis
    if not report.get("can_proceed", False):
        raise HTTPException(
            status_code=400,
            detail="This file still has unresolved issues. Please finish cleaning before generating insights."
        )
    analysis_context = build_financial_analysis_context(cleaned_df, mapping, client_id, file_type)
    # Generate AI insights based on the calculated financial data
    try:
        if analysis_context.get("financial_analytics"):
            insights = generate_financial_ai_insights(
                analysis_context.get("financial_statements"),
                analysis_context.get("financial_ratios"),
                analysis_context.get("financial_analytics"),
                analysis_context.get("comparative_analytics"),
                analysis_context.get("generic_analysis"),
            )
        else:
            insights = generate_ai_insights(
                analysis_context.get("breakdowns", {}),
                analysis_context.get("monthly_trend", {}),
                analysis_context.get("anomalies", []),
            )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Insight generation failed: {str(e)}")
    # Return the AI insights along with a success message
    return {
        "client_id": client_id,
        "file_id": file_id,
        "file_type": file_type,
        "analysis_scope": analysis_context.get("analysis_scope"),
        "analysis_basis": analysis_context.get("analysis_basis"),
        "ai_insights": insights,
        "message": "AI insights generated successfully."
    }


# Validates a trial balance / general ledger file after cleaning.
@app.post("/validate-trial-balance")
async def validate_trial_balance_endpoint(
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general")
):
    mapping = get_file_specific_mapping(file_id, client_id, file_type)
    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    if not report.get("can_proceed", False):
        raise HTTPException(
            status_code=400,
            detail="This file still has unresolved cleaning issues. Please finish cleaning before validating the trial balance."
        )
    validation_result = validate_trial_balance(cleaned_df, mapping)
    save_tb_validation_result(file_id, client_id, file_type, validation_result)
    update_workflow_stage(file_id, client_id, file_type, "tb_validation")
    return {
        "client_id": client_id,
        "file_id": file_id,
        "file_type": file_type,
        "trial_balance_validation": validation_result,
    }

# Suggest standard category classifications
@app.post("/detect-account-mapping")
async def detect_account_mapping_endpoint(
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general")
):
    mapping = get_file_specific_mapping(file_id, client_id, file_type)

    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    if not report.get("can_proceed", False):
        raise HTTPException(
            status_code=400,
            detail="This file still has unresolved cleaning issues. Please finish cleaning first."
        )
    saved_account_mapping = get_account_mapping(client_id, file_type)
    result = build_account_mapping_result(cleaned_df, mapping, saved_account_mapping)
    return {
        "client_id": client_id,
        "file_id": file_id,
        "file_type": file_type,
        "account_mapping": result,
    }

# Saves the auditor's confirmed account category classifications.
@app.post("/save-account-mapping")
async def save_account_mapping_endpoint(
    client_id: str = Form(...),
    file_type: str = Form(...),
    accounts: str = Form(...),
    confirmed_by: str = Form(None)
):
    try:
        accounts_list = json.loads(accounts)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid accounts format.")
    if not accounts_list:
        raise HTTPException(status_code=400, detail="No accounts provided.")

    save_account_mapping(client_id, file_type, accounts_list, confirmed_by)
    return {
        "client_id": client_id,
        "file_type": file_type,
        "accounts_saved": len(accounts_list),
        "message": "Account mapping saved successfully."
    }

# Generates the Income Statement and Balance Sheet from confirmed account classifications. Requires account mapping to have been
# saved first — pure aggregation, no AI involved at this stage.
@app.post("/generate-financial-statements")
async def generate_financial_statements_endpoint(
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general")
):
    mapping = get_file_specific_mapping(file_id, client_id, file_type)

    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)

    if not report.get("can_proceed", False):
        raise HTTPException(
            status_code=400,
            detail="This file still has unresolved cleaning issues. Please finish cleaning first."
        )

    account_mapping = get_account_mapping(client_id, file_type)
    if not account_mapping:
        raise HTTPException(
            status_code=400,
            detail="No account mapping found for this client. Please complete account mapping first."
        )

    result = generate_financial_statements(cleaned_df, mapping, account_mapping)
    if result.get("applicable"):
        result["financial_ratios"] = calculate_financial_ratios(
            result["income_statement"],
            result["balance_sheet"],
        )

    return {
        "client_id": client_id,
        "file_id": file_id,
        "file_type": file_type,
        "financial_statements": result,
        "financial_ratios": result.get("financial_ratios"),
    }
    
# =============================================================================
# REFACTORED REPORT GENERATION LOGIC
# =============================================================================
# --- Section completion validation ------------------------------------------

class SectionValidationResult:
    """
    Small, explicit result object for a section-completion check, so callers
    never have to remember the shape of the raw dict from
    check_all_sections_completed() (e.g. which key means what).
    """
    def __init__(self, all_completed: bool, pending_sections: list, sections_data: dict):
        self.all_completed = all_completed
        self.pending_sections = pending_sections
        self.sections_data = sections_data


def validate_sections_completed(engagement_id: int) -> SectionValidationResult:
    """
    The single source of truth for "is this engagement ready for a report?".

    Uses the existing check_all_sections_completed() DB helper to find out
    whether every audit section for this engagement has reached a completed
    state, and get_all_sections_data() to pull the actual section content
    (real data entered by accountants/auditors — not mock data) so it can be
    folded into the report later.

    Any endpoint that wants to gate on "all sections done" should call this
    function rather than querying check_all_sections_completed() directly,
    so the gating rule lives in exactly one place.
    """
    if not engagement_id:
        # No engagement context at all means there's nothing to validate
        # against — treat as "not completed" rather than silently passing.
        return SectionValidationResult(
            all_completed=False,
            pending_sections=["No engagement_id provided"],
            sections_data={},
        )

    status = check_all_sections_completed(engagement_id)
    all_completed = bool(status.get("all_completed"))
    pending_sections = status.get("pending_sections", [])

    # Only pull the actual section data once we know it's worth pulling —
    # no point querying every section's content for a report we're about
    # to reject anyway.
    sections_data = get_all_sections_data(engagement_id) if all_completed else {}

    return SectionValidationResult(
        all_completed=all_completed,
        pending_sections=pending_sections,
        sections_data=sections_data,
    )


# --- Single combined reporting function -------------------------------------

def build_report_payload(
    req: "ReportGenerateRequest",
    period_start,
    period_end,
    period_label: str,
    cleaned_df: "pd.DataFrame",
    mapping: dict,
    sections_data: dict,
    included_sections: list = None,
    account_mapping: dict = None,
) -> dict:
    """
    The one function that turns "everything we know" into "the report".

    Inputs are the combined, real data from every approved, in-scope
    section (via get_engagement_combined_dataset — the same source AI
    Insights uses, so a report and the final analysis can't diverge):
      - cleaned_df / mapping  -> the combined approved-sections dataset
      - sections_data         -> the actual completed audit sections for
                                  this engagement (from get_all_sections_data)
      - included_sections     -> which sections/files fed this report,
                                  for the same auditability final analysis has
      - req                   -> the request parameters (period, commentary, etc.)

    Returns a single dict containing every piece the caller needs both to
    persist the report (financial_summary / ai_insights / commentary /
    chart_refs map straight onto report_versions columns) and to return to
    the client. No other function should independently compute breakdowns,
    insights, or chart specs for a report — they all happen here so there's
    exactly one code path producing a report's content.
    """
    # Scope the cleaned data to the requested period. Trial balances often
    # have no date column, so fall back to the full cleaned dataset in that
    # case — same behavior as before, just kept inside the single builder now.
    date_col = resolve_mapped_column(mapping, "date")
    if date_col and date_col in cleaned_df.columns:
        period_df = filter_dataframe_by_period(cleaned_df, mapping, period_start, period_end)
    else:
        period_df = cleaned_df

    breakdowns = calculate_breakdowns(period_df, mapping)
    monthly_trend = calculate_monthly_trend(period_df, mapping)
    anomalies = detect_anomalies(monthly_trend)
    financial_statements = None
    if account_mapping:
        financial_statements = generate_financial_statements(period_df, mapping, account_mapping)
        if financial_statements.get("applicable"):
            financial_statements["financial_ratios"] = calculate_financial_ratios(
                financial_statements["income_statement"],
                financial_statements["balance_sheet"],
            )

    # For ledger-style data (trial balance / general ledger), generic
    # breakdowns are usually empty — fall back to the account-mapped
    # statements, same numbers "Generate Analysis" already computed.
    report_financial_summary = breakdowns
    if financial_statements and financial_statements.get("applicable"):
        report_financial_summary = {
            "financial_statements": financial_statements,
            "breakdowns": breakdowns,
        }

    try:
        ai_insights = generate_ai_insights(breakdowns, monthly_trend, anomalies)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Insight generation failed: {str(e)}")

    chart_specs = build_chart_specs(breakdowns, monthly_trend)

    return {
        "client_id": req.client_id,
        "engagement_id": req.engagement_id,
        # file_id deprecated (a report can now be fed by multiple approved
        # files) — kept as None for the reports table's existing column,
        # included_sections below is the real record of what fed this report.
        "file_id": None,
        "included_sections": included_sections or [],
        "report_type": req.report_type,
        "period_start": period_start,
        "period_end": period_end,
        "period_label": period_label,
        "financial_summary": report_financial_summary,
        "monthly_trend": monthly_trend,
        "anomalies": anomalies,
        "ai_insights": ai_insights,
        "chart_specs": chart_specs,
        # Real, completed-section data — this is what makes the report
        # reflect the actual engagement rather than just the uploaded file.
        "sections_data": sections_data,
        "commentary": req.commentary or "",
    }


# --- Endpoint -----------------------------------------------------------------

@app.post("/api/reports/generate")
def generate_report(req: ReportGenerateRequest, db=Depends(get_db)):
    # 1. GATHER — the same combined, approved-sections dataset AI Insights
    #    uses (get_engagement_combined_dataset). This single call replaces
    #    the old "validate sections, then separately clean one file" flow —
    #    it raises 409 itself if any in-scope section isn't approved yet,
    #    and it's the same function final analysis calls, so a report and
    #    the final analysis can never show different numbers for the same
    #    engagement.
    dataset = get_engagement_combined_dataset(req.engagement_id)
    combined_df = dataset["combined_df"]
    mapping = dataset["mapping"]

    # Real completed-section content (accountant/auditor notes, statuses)
    # for display in the report — separate from the financial numbers above.
    sections_data = get_all_sections_data(req.engagement_id)

    period_start, period_end, period_label = resolve_report_period(req)

    # 2. BUILD — single function combines the combined dataset with the
    #    completed sections data into the final report content.
    payload = build_report_payload(
        req, period_start, period_end, period_label,
        combined_df, mapping, sections_data, dataset["included_sections"],
        account_mapping=dataset.get("account_mapping"),
    )

    # 3. PERSIST — unchanged from before: write reports + report_versions.
    report_id = str(uuid.uuid4())
    version_id = str(uuid.uuid4())

    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO reports
           (id, client_id, engagement_id, file_id, type, period_start, period_end, status, current_version_id, created_by, created_at)
           VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
        (
            report_id, payload["client_id"], payload["engagement_id"], payload["file_id"], payload["report_type"],
            payload["period_start"].date(), payload["period_end"].date(),
            "draft", None, req.generated_by, datetime.utcnow(),
        )
    )
    cursor.execute(
        """INSERT INTO report_versions
           (id, report_id, version_number, financial_summary, ai_insights,
            commentary, chart_refs, generated_by, status, created_at)
           VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
        (
            version_id, report_id, 1,
            json.dumps(payload["financial_summary"], default=str),
            json.dumps(payload["ai_insights"], default=str),
            payload["commentary"],
            json.dumps({
                "chart_specs": payload["chart_specs"],
                "monthly_trend": payload["monthly_trend"],
                "anomalies": payload["anomalies"],
                "sections_data": payload["sections_data"],
            }, default=str),
            "ai", "draft", datetime.utcnow(),
        )
    )
    cursor.execute(
        "UPDATE reports SET current_version_id = %s WHERE id = %s",
        (version_id, report_id)
    )
    db.commit()

    return {
        "report_id": report_id,
        "client_id": payload["client_id"],
        "engagement_id": payload["engagement_id"],
        "report_type": payload["report_type"],
        "period_label": payload["period_label"],
        "period_start": payload["period_start"].date().isoformat(),
        "period_end": payload["period_end"].date().isoformat(),
        "financial_summary": payload["financial_summary"],
        "monthly_trend": payload["monthly_trend"],
        "ai_insights": payload["ai_insights"],
        "anomalies": payload["anomalies"],
        "chart_specs": payload["chart_specs"],
        "sections_data": payload["sections_data"],
        "included_sections": payload["included_sections"],
        "commentary": payload["commentary"],
        "message": "Report generated successfully.",
    }


# NOTE: GET /reports/{id}, GET /clients/{id}/reports, PUT /reports/{id}/commentary,
# and DELETE /reports/{id} moved to report_routes.py under /api/reports —
# see GET/PATCH /api/reports/{report_id}[/commentary] and DELETE /api/reports/{report_id}.


# --- Engagement status auto-progression helpers -----------------------------
#
# Status is DERIVED at read-time from real workflow data instead of being set
# directly, so it can never drift out of sync (e.g. if a section gets
# un-approved after review started, status correctly drops back down):
#
#   Planning      -> default, nothing has happened yet
#   In Progress   -> at least one submission has left Draft / the Accountant stage
#   Under Review  -> every section's LATEST submission is "Approved"
#   Completed     -> engagement.sent_to_client_at has been set
#
# Requires a migration (run once):
#   ALTER TABLE engagements ADD COLUMN sent_to_client_at DATETIME NULL;

def fetch_engagement_progress(db, engagement_ids: list) -> dict:
    """Batch-fetch the section/submission counts needed to derive status for
    a list of engagement ids, in 3 queries total (not one per engagement)."""
    if not engagement_ids:
        return {}
    placeholders = ",".join(["%s"] * len(engagement_ids))
    progress = {
        eid: {"total_sections": 0, "approved_sections": 0, "forwarded": False}
        for eid in engagement_ids
    }
    cursor = db.cursor(dictionary=True)

    # Total sections per engagement
    cursor.execute(
        f"SELECT engagement_id, COUNT(*) AS total FROM audit_sections "
        f"WHERE engagement_id IN ({placeholders}) AND in_scope = 1 "
        f"GROUP BY engagement_id",
        tuple(engagement_ids)
    )
    for row in cursor.fetchall():
        progress[row["engagement_id"]]["total_sections"] = row["total"]

    # Sections whose most recent submission is "Approved" (MySQL 8+ window function)
    cursor.execute(
        f"""
        SELECT engagement_id, COUNT(*) AS approved
        FROM (
            SELECT section_id, engagement_id, status,
                   ROW_NUMBER() OVER (PARTITION BY section_id ORDER BY created_at DESC) AS rn
            FROM submissions
            WHERE engagement_id IN ({placeholders})
        ) latest
        WHERE rn = 1 AND status = 'Approved'
        GROUP BY engagement_id
        """,
        tuple(engagement_ids)
    )
    for row in cursor.fetchall():
        progress[row["engagement_id"]]["approved_sections"] = row["approved"]

    # Engagements where at least one submission has moved past Draft/Accountant
    cursor.execute(
        f"""
        SELECT DISTINCT engagement_id
        FROM submissions
        WHERE engagement_id IN ({placeholders})
          AND (status != 'Draft' OR current_stage != 'Accountant')
        """,
        tuple(engagement_ids)
    )
    for row in cursor.fetchall():
        progress[row["engagement_id"]]["forwarded"] = True

    return progress


def apply_display_status(engagement: dict, progress: dict) -> None:
    """Mutates `engagement` in place, adding a `display_status` field derived
    from its section/submission progress. `engagement` must include
    `sent_to_client_at` (from a plain SELECT * on engagements)."""
    total = progress.get("total_sections", 0)
    approved = progress.get("approved_sections", 0)
    forwarded = progress.get("forwarded", False)

    if engagement.get("sent_to_client_at"):
        status = "Completed"
    elif total > 0 and approved == total:
        status = "Under Review"
    elif forwarded:
        status = "In Progress"
    else:
        status = "Planning"

    engagement["display_status"] = status
    engagement["total_sections"] = total
    engagement["approved_sections"] = approved


# List all clients
@app.get("/clients")
def get_clients(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM clients")
    return cursor.fetchall()

# Get a single client by id
@app.get("/clients/{client_id}")
def get_client(client_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM clients WHERE client_id = %s", (client_id,))
    client = cursor.fetchone()
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    return client

# Create a new client
 
@app.post("/clients")
def create_client(c: Client, db=Depends(get_db)):
    try:
        email = FieldValidator.validate_email(c.email) if c.email else None
        phone = FieldValidator.validate_phone(c.phone) if c.phone else None
        kra_pin_number = (
            FieldValidator.validate_kra_pin(c.kra_pin_number) if c.kra_pin else None
        )
    except FieldValidationError as e:
        raise HTTPException(status_code=422, detail=e.message)
 
    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO clients (company_name, contact_person, email, phone, industry, address, status, kra_pin, kra_pin_number)
           VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)""",
        (c.company_name, c.contact_person, email, phone, c.industry, c.address, c.status, c.kra_pin, kra_pin_number)
    )
    db.commit()
    return {"client_id": cursor.lastrowid, "message": "Client created"}
# Update an existing client

@app.put("/clients/{client_id}")
def update_client(client_id: int, c: Client, db=Depends(get_db)):
    try:
        email = FieldValidator.validate_email(c.email) if c.email else None
        phone = FieldValidator.validate_phone(c.phone) if c.phone else None
        kra_pin_number = (
            FieldValidator.validate_kra_pin(c.kra_pin_number) if c.kra_pin else None
        )
    except FieldValidationError as e:
        raise HTTPException(status_code=422, detail=e.message)
 
    cursor = db.cursor()
    cursor.execute(
        """UPDATE clients SET company_name=%s, contact_person=%s, email=%s,
           phone=%s, industry=%s, address=%s, status=%s, kra_pin=%s, kra_pin_number=%s WHERE client_id=%s""",
        (c.company_name, c.contact_person, email, phone, c.industry, c.address, c.status, c.kra_pin, kra_pin_number, client_id)
    )
    db.commit()
    return {"message": "Client updated"}

# Delete a client along with every record that depends on it (engagements,
# sections, team assignments, submissions, uploads), and unassign any users
@app.delete("/clients/{client_id}")
def delete_client(client_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    try:
        cursor.execute("SELECT engagement_id FROM engagements WHERE client_id = %s", (client_id,))
        rows = cursor.fetchall()
        engagement_ids = [r['engagement_id'] for r in rows] if rows else []
        if engagement_ids:
            placeholders = ','.join(['%s'] * len(engagement_ids))
            cursor.execute(f"DELETE FROM submissions WHERE engagement_id IN ({placeholders})", tuple(engagement_ids))
            cursor.execute(f"DELETE FROM audit_sections WHERE engagement_id IN ({placeholders})", tuple(engagement_ids))
            cursor.execute(f"DELETE FROM engagement_team WHERE engagement_id IN ({placeholders})", tuple(engagement_ids))
            cursor.execute(f"DELETE FROM engagements WHERE engagement_id IN ({placeholders})", tuple(engagement_ids))
        cursor.execute("DELETE FROM uploads WHERE client_id = %s", (str(client_id),))
        cursor.execute("UPDATE users SET assigned_client_id = NULL WHERE assigned_client_id = %s", (client_id,))
        cursor.execute("DELETE FROM clients WHERE client_id = %s", (client_id,))
        if cursor.rowcount == 0:
            db.rollback()
            raise HTTPException(status_code=404, detail="Client not found")
        db.commit()
        return {"message": "Client and dependent records deleted"}
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Could not delete client: {str(e)}")

# List all users along with their assigned client's company name
@app.get("/users")
def get_users(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT u.user_id, u.full_name, u.email, u.phone, u.role, u.status,
               u.assigned_client_id, u.created_at, c.company_name
        FROM users u LEFT JOIN clients c ON u.assigned_client_id = c.client_id
    """)
    return cursor.fetchall()

# Get a single user by id, with the password hash removed from the response
@app.get("/users/{user_id}")
def get_user(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE user_id = %s", (user_id,))
    user = cursor.fetchone()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    user.pop("password_hash", None)
    return user

# Create a new user with a hashed password
@app.post("/users")
def create_user(u: User, db=Depends(get_db)):
    # A bad assigned_client_id would otherwise surface as a raw MySQL
    # foreign-key error (1452) — check it up front so the person gets a
    # clear message instead (e.g. the client was deleted since the
    # dropdown was last loaded).
    if u.assigned_client_id is not None:
        validate_client_exists(db, u.assigned_client_id)

    hashed = hash_password(u.password)
    cursor = db.cursor()
    try:
        cursor.execute(
            """INSERT INTO users (full_name, email, password_hash, phone, role, assigned_client_id, status)
               VALUES (%s, %s, %s, %s, %s, %s, %s)""",
            (u.full_name, u.email, hashed, u.phone or None, u.role, u.assigned_client_id, u.status)
        )
        db.commit()
        return {"user_id": cursor.lastrowid, "message": "User created"}
    except HTTPException:
        raise
    except Exception as e:
        error_msg = str(e)
        if "Duplicate entry" in error_msg and "email" in error_msg.lower():
            raise HTTPException(
                status_code=400,
                detail=error_detail("Email already exists", error_code="CONFLICT", details={"field": "email"}),
            )
        else:
            raise HTTPException(
                status_code=400,
                detail=error_detail(f"Failed to create user: {error_msg}", error_code="VALIDATION_ERROR"),
            )

# Update an existing user's details (does not change the password)
@app.put("/users/{user_id}")
def update_user(user_id: int, u: UserUpdate, db=Depends(get_db)):
    # Same check as create_user — a bad assigned_client_id should come
    # back as a clear 400, not a raw MySQL foreign-key error.
    if u.assigned_client_id is not None:
        validate_client_exists(db, u.assigned_client_id)

    cursor = db.cursor()
    try:
        cursor.execute(
            """UPDATE users SET full_name=%s, email=%s, phone=%s,
               role=%s, assigned_client_id=%s, status=%s WHERE user_id=%s""",
            (u.full_name, u.email, u.phone, u.role, u.assigned_client_id, u.status, user_id)
        )
        db.commit()
        return {"message": "User updated"}
    except HTTPException:
        raise
    except Exception as e:
        error_msg = str(e)
        if "Duplicate entry" in error_msg and "email" in error_msg.lower():
            raise HTTPException(
                status_code=400,
                detail=error_detail("Email already exists", error_code="CONFLICT", details={"field": "email"}),
            )
        else:
            raise HTTPException(
                status_code=400,
                detail=error_detail(f"Failed to update user: {error_msg}", error_code="VALIDATION_ERROR"),
            )

# Reset a user's password
@app.put("/users/{user_id}/reset-password")
def reset_user_password(user_id: int, payload: dict, db=Depends(get_db)):
    new_password = payload.get("new_password")
    try:
        new_password = FieldValidator.validate_password(new_password)
    except FieldValidationError as e:
        raise HTTPException(status_code=422, detail=e.message)
    hashed = hash_password(new_password)
    cursor = db.cursor()
    cursor.execute("UPDATE users SET password_hash = %s WHERE user_id = %s", (hashed, user_id))
    db.commit()
    return {"message": "Password reset successful"}

# Lock or unlock a user's account
@app.put("/users/{user_id}/login-lock")
def set_user_login_lock(user_id: int, payload: dict, db=Depends(get_db)):
    locked = payload.get("locked")
    if locked is None:
        raise HTTPException(status_code=400, detail="Missing locked flag")
    cursor = db.cursor()
    cursor.execute("UPDATE users SET login_locked = %s WHERE user_id = %s", (1 if locked else 0, user_id))
    db.commit()
    return {"message": "Login access updated", "login_locked": bool(locked)}

# Get a user's login history
@app.get("/users/{user_id}/login-history")
def get_user_login_history(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute(
        "SELECT id, timestamp, ip_address, device, status FROM login_history WHERE user_id = %s ORDER BY timestamp DESC",
        (user_id,)
    )
    return cursor.fetchall()

# Delete a user
@app.delete("/users/{user_id}")
def delete_user(user_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM users WHERE user_id = %s", (user_id,))
    db.commit()
    return {"message": "User deleted"}

# Assign a user to a client
@app.put("/users/{user_id}/assign/{client_id}")
def assign_user_to_client(user_id: int, client_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("UPDATE users SET assigned_client_id=%s WHERE user_id=%s", (client_id, user_id))
    db.commit()
    return {"message": "User assigned to client"}

# Log in a user with email and password, returning a JWT access token on success
@app.post("/auth/login")
def login(req: LoginRequest, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE email = %s", (req.email,))
    user = cursor.fetchone()
    if not user or not verify_password(req.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    if user["status"] != "Active":
        raise HTTPException(status_code=403, detail="Account is inactive")
    token = create_token({"user_id": user["user_id"], "email": user["email"], "role": user["role"]})
    return {"access_token": token, "token_type": "bearer",
            "user": {"user_id": user["user_id"], "full_name": user["full_name"],
                     "email": user["email"], "role": user["role"]}}

GMAIL_USER = os.getenv("GMAIL_USER")
GMAIL_APP_PASSWORD = os.getenv("app_password")
FRONTEND_URL = os.getenv("FRONTEND_URL", "http://localhost:5173")

def send_reset_email(to_email: str, token: str):
    reset_link = f"{FRONTEND_URL}/password-reset/confirm/{token}"
    msg = MIMEMultipart()
    msg["From"] = GMAIL_USER
    msg["To"] = to_email
    msg["Subject"] = "Password Reset Request - Audit AI"
    body = f"""Hello,

You requested a password reset for your Audit AI account.

Click the link below to reset your password:
{reset_link}

This link will expire in 1 hour.

If you did not request this, please ignore this email.

Best regards,
The Audit AI Team"""
    msg.attach(MIMEText(body, "plain"))
    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(GMAIL_USER, GMAIL_APP_PASSWORD)
        server.sendmail(GMAIL_USER, to_email, msg.as_string())

# Return everything needed to resume work on a file from the Client Details page.
# Reads the actual file from disk to get columns and fill_rates so MappingPage
# can run detection correctly on resume, exactly as a fresh upload would.
# Mark workflow steps as completed
@app.post("/workflow/complete-step")
async def complete_workflow_step(
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form(...),
    step: str = Form(...),
    next_stage: str = Form(None)
):
    """
    Mark a workflow step as completed and optionally move to the next stage.
    step can be: 'tb_validation', 'account_mapping', 'financial_analysis'
    """
    try:
        mark_workflow_step_completed(file_id, client_id, file_type, step)
        
        if next_stage:
            update_workflow_stage(file_id, client_id, file_type, next_stage)
        
        return {"message": f"Step {step} marked as completed", "next_stage": next_stage}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/workflow/stage/{file_id}")
async def get_workflow_stage_endpoint(file_id: str, client_id: str, file_type: str = "general"):
    """
    Get the current workflow stage for a file.
    """
    try:
        workflow = get_workflow_stage(file_id, client_id, file_type)
        if not workflow:
            return {"current_stage": None, "tb_validation_completed": False, 
                    "account_mapping_completed": False, "financial_analysis_completed": False}
        return workflow
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/files/{file_id}/resume-state")
def get_resume_state(file_id: str, client_id: str, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)

    # Get the upload record
    cursor.execute("SELECT * FROM uploads WHERE file_id = %s AND client_id = %s", (file_id, client_id))    
    upload_rows = cursor.fetchall()
    upload = upload_rows[0] if upload_rows else None
    if not upload:
        # Return a graceful response instead of 404 for deleted/non-existent files
        return {
            "file_id": file_id,
            "client_id": client_id,
            "filename": None,
            "file_type": None,
            "row_count": 0,
            "columns": [],
            "fill_rates": {},
            "fingerprint": None,
            "upload_time": None,
            "stage": "file_not_found",
            "has_mapping": False,
            "has_corrections": False,
            "total_issues": None,
            "can_proceed": False,
            "last_cleaned_at": None,
            "workflow": None,
            "error": "File not found in uploads table. It may have been deleted."
        }

    # Use semantic_file_type if available, otherwise fall back to file_type
    file_type = upload.get("semantic_file_type") or upload.get("file_type") or "other"
    filename = upload.get("filename") or upload.get("file_name") or ""

    # Read the actual file from disk to get columns and fill_rates
    columns = []
    fill_rates = {}
    fingerprint = ""
    row_count = upload.get("row_count") or 0
    # Use the original file extension from upload record if available
    original_ext = upload.get("file_type") if upload.get("file_type") in ALLOWED_EXTENSIONS else None
    save_path, file_ext = locate_uploaded_file(file_id, original_ext)
    
    # Get stored header configuration from uploads table
    stored_header_row = upload.get('header_row_index')
    stored_sheet_name = upload.get('sheet_name')
    
    if save_path:
        try:
            df, _ = read_file_to_df(save_path, file_ext, stored_sheet_name, file_type=None, header_row_index=stored_header_row)
            if df is not None:
                columns = list(df.columns)
                fill_rates = calculate_fill_rates(df)
                fingerprint = compute_schema_fingerprint(columns)
                row_count = len(df)
        except Exception:
            pass

    # The uploads table stores the file extension (e.g. 'xlsx'), not the mapped
    # category (e.g. 'general_ledger'). Resolve the real file_type by finding
    # which column_mappings entry covers the most columns from this file.
    # First check if this file already has a workflow record — its file_type is authoritative
    cursor.execute(
        "SELECT file_type FROM workflow_stages WHERE file_id = %s AND client_id = %s",
        (file_id, client_id)
    )
    workflow_rows = cursor.fetchall()
    existing_workflow_row = cursor.fetchone()
    if existing_workflow_row:
        file_type = existing_workflow_row["file_type"]
    elif columns:
        # No workflow record yet — guess file_type by column overlap
        placeholders = ",".join(["%s"] * len(columns))
        cursor.execute(
            f"SELECT file_type, COUNT(*) AS cnt FROM column_mappings "
            f"WHERE client_id = %s AND original_column IN ({placeholders}) "
            f"GROUP BY file_type ORDER BY cnt DESC LIMIT 1",
            (client_id, *columns)
        )
        ft_row = cursor.fetchone()
        if ft_row:
            file_type = ft_row["file_type"]

    # Check if a mapping exists for this specific file (file-specific mapping isolation)
    # This prevents cross-file mapping reuse when files have identical schemas
    # If a workflow row already exists, mapping obviously happened — trust that over file_id check
    if existing_workflow_row:
        has_mapping = True
    else:
        # Check for file-specific mapping using file_id
        cursor.execute(
            "SELECT COUNT(*) AS cnt FROM column_mappings WHERE client_id = %s AND file_id = %s",
            (client_id, file_id)
        )
        has_mapping = cursor.fetchone()["cnt"] > 0

    # Check cleaned registry
    cursor.execute(
         "SELECT total_issues, can_proceed, updated_at FROM cleaned_files_registry WHERE file_id = %s AND client_id = %s AND file_type = %s",
        (file_id, client_id, file_type)
    )
    cleaned_rows = cursor.fetchall()
    cleaned = cleaned_rows[0] if cleaned_rows else None
    

    # Check if any corrections exist — tells us the auditor was on CorrectedResultsPage
    cursor.execute(
        "SELECT COUNT(*) AS cnt FROM cleaning_corrections WHERE file_id = %s AND client_id = %s AND file_type = %s",
        (file_id, client_id, file_type)
    )
    has_corrections = cursor.fetchone()["cnt"] > 0

    # Initialize workflow variable
    workflow = None

    # Determine stage - mapping is now file-specific, not shared across files
    if not has_mapping:
        stage = "uploaded"
    elif cleaned:
        # This specific file has been through cleaning
        # Always check workflow stage if file has been cleaned
        workflow = get_workflow_stage(file_id, client_id, file_type)
        if workflow and workflow["current_stage"]:
            # Use the workflow stage if it exists
            stage = workflow["current_stage"]
        elif cleaned["can_proceed"]:
            # If no workflow stage but can proceed, initialize it
            initialize_workflow_stage(file_id, client_id, file_type)
            workflow = get_workflow_stage(file_id, client_id, file_type)
            if workflow and workflow["current_stage"]:
                stage = workflow["current_stage"]
            else:
                # Fallback to clean if initialization failed
                stage = "clean"
        else:
            # File has issues and no workflow stage
            stage = "cleaning_in_progress"
    else:
        # Mapping exists but this specific file hasn't been cleaned yet
        stage = "mapped"

    # Get workflow stage data if file is cleaned or in workflow
    if not workflow and stage in ["clean", "tb_validation", "account_mapping", "financial_analysis", "analysis"]:
        workflow_data = get_workflow_stage(file_id, client_id, file_type)
        if workflow_data:
            workflow = workflow_data
    
    return {
        "file_id": file_id,
        "client_id": client_id,
        "filename": filename,
        "file_type": file_type,
        "row_count": row_count,
        "columns": columns,
        "fill_rates": fill_rates,
        "fingerprint": fingerprint,
        "upload_time": str(upload.get("upload_time") or ""),
        "stage": stage,
        "has_mapping": has_mapping,
        "has_corrections": has_corrections,
        "total_issues": cleaned["total_issues"] if cleaned else None,
        "can_proceed": bool(cleaned["can_proceed"]) if cleaned else False,
        "last_cleaned_at": str(cleaned["updated_at"]) if cleaned else None,
        "workflow": workflow,
    }

# Generate a password reset token for a user's email and store it with a one hour expiry
@app.post("/auth/password-reset-request")
def password_reset_request(req: PasswordResetRequest, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE email = %s", (req.email,))
    user = cursor.fetchone()
    if not user:
        raise HTTPException(status_code=404, detail="Email not found")
    token = secrets.token_urlsafe(32)
    expires_at = datetime.utcnow() + timedelta(hours=1)
    cursor2 = db.cursor()
    cursor2.execute("INSERT INTO password_resets (user_id, token, expires_at) VALUES (%s, %s, %s)",
                    (user["user_id"], token, expires_at))
    db.commit()
    try:
        send_reset_email(user["email"], token)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Token generated but email could not be sent: {str(e)}")
    return {"message": "Password reset link sent to your email"}

# Confirm a password reset using a valid, unexpired token and set the new password
@app.post("/auth/password-reset-confirm")
def password_reset_confirm(req: PasswordResetConfirm, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM password_resets WHERE token = %s AND expires_at > UTC_TIMESTAMP()", (req.token,))
    reset = cursor.fetchone()
    if not reset:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    hashed = hash_password(req.new_password)
    cursor2 = db.cursor()
    cursor2.execute("UPDATE users SET password_hash = %s WHERE user_id = %s", (hashed, reset["user_id"]))
    cursor2.execute("DELETE FROM password_resets WHERE token = %s", (req.token,))
    db.commit()
    return {"message": "Password reset successful"}

# List every column mapping record in the system
@app.get("/column-mappings")
def get_all_mappings(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM column_mappings")
    return cursor.fetchall()

# List column mapping records for a single client
@app.get("/column-mappings/{client_id}")
def get_client_mappings(client_id: str, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM column_mappings WHERE client_id = %s", (client_id,))
    return cursor.fetchall()

# Create a single column mapping record directly (separate from the AI mapping save flow)
@app.post("/column-mappings")
def create_mapping(m: ColumnMapping, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO column_mappings (client_id, file_type, original_column, mapped_to, confirmed_by)
           VALUES (%s, %s, %s, %s, %s)""",
        (m.client_id, m.file_type, m.original_column, m.mapped_to, m.confirmed_by)
    )
    db.commit()
    return {"id": cursor.lastrowid, "message": "Column mapping created"}

# Update a single column mapping record by id
@app.put("/column-mappings/{mapping_id}")
def update_mapping(mapping_id: int, m: ColumnMapping, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """UPDATE column_mappings SET client_id=%s, file_type=%s, original_column=%s,
           mapped_to=%s, confirmed_by=%s WHERE id=%s""",
        (m.client_id, m.file_type, m.original_column, m.mapped_to, m.confirmed_by, mapping_id)
    )
    db.commit()
    return {"message": "Column mapping updated"}

# Delete a single column mapping record by id
@app.delete("/column-mappings/{mapping_id}")
def delete_mapping(mapping_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM column_mappings WHERE id = %s", (mapping_id,))
    db.commit()
    return {"message": "Column mapping deleted"}

# List all engagements with their client's company name, most recent first,
# along with an auto-derived workflow status (see compute_engagement_status_fields)
@app.get("/engagements")
def get_engagements(db=Depends(get_db), current_user: dict = Depends(get_current_user)):
    cursor = db.cursor(dictionary=True)
    if current_user["role"] == "Accountant":
        cursor.execute("""
            SELECT e.*, c.company_name FROM engagements e
            LEFT JOIN clients c ON e.client_id = c.client_id
            INNER JOIN engagement_team et ON et.engagement_id = e.engagement_id
            WHERE et.user_id = %s
            ORDER BY e.created_at DESC
        """, (current_user["user_id"],))
    else:
        cursor.execute("""
            SELECT e.*, c.company_name FROM engagements e
            LEFT JOIN clients c ON e.client_id = c.client_id
            ORDER BY e.created_at DESC
        """)
    engagements = cursor.fetchall()
    if not engagements:
        return engagements
    engagement_ids = [row["engagement_id"] for row in engagements]
    progress_by_id = fetch_engagement_progress(db, engagement_ids)
    for row in engagements:
        apply_display_status(row, progress_by_id.get(row["engagement_id"], {}))
    return engagements

# Get a single engagement by id, with its client's company name and derived status
@app.get("/engagements/{engagement_id}")
def get_engagement(engagement_id: int, db=Depends(get_db), current_user: dict = Depends(get_current_user)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT e.*, c.company_name FROM engagements e
        LEFT JOIN clients c ON e.client_id = c.client_id
        WHERE e.engagement_id = %s
    """, (engagement_id,))
    engagement = cursor.fetchone()
    if not engagement:
        raise HTTPException(status_code=404, detail="Engagement not found")
    if current_user["role"] == "Accountant":
        cursor.execute(
            "SELECT 1 FROM engagement_team WHERE engagement_id = %s AND user_id = %s",
            (engagement_id, current_user["user_id"])
        )
        if not cursor.fetchone():
            raise HTTPException(status_code=403, detail="You are not assigned to this engagement.")
    progress_by_id = fetch_engagement_progress(db, [engagement_id])
    apply_display_status(engagement, progress_by_id.get(engagement_id, {}))
    return engagement


# Create a new engagement and automatically create its four default audit sections
# Checks the audit period is real and start comes before end
def validate_engagement_dates(start_date: Optional[str], end_date: Optional[str]):
    if not start_date or not end_date:
        return
    start = pd.to_datetime(start_date, errors="coerce")
    end = pd.to_datetime(end_date, errors="coerce")
    if pd.isna(start) or pd.isna(end):
        raise HTTPException(status_code=400, detail="Start date and end date must be valid dates.")
    if start > end:
        raise HTTPException(status_code=400, detail="Start date must be on or before end date.")

# Create a new engagement and its selected audit sections
@app.post("/engagements")
def create_engagement(e: Engagement, db=Depends(get_db)):

    validate_engagement_dates(e.start_date, e.end_date)
    if not e.sections:
        raise HTTPException(status_code=400, detail="Select at least one audit section for this engagement.")

    validate_engagement_creation(
        db, e.client_id, e.engagement_name, e.financial_year, e.start_date, e.end_date
    )

    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO engagements (client_id, engagement_name, financial_year, status, start_date, end_date)
           VALUES (%s, %s, %s, %s, %s, %s)""",
        (e.client_id, e.engagement_name, e.financial_year, e.status, e.start_date, e.end_date)
    )
    engagement_id = cursor.lastrowid
    for section in e.sections:
        cursor.execute("INSERT INTO audit_sections (engagement_id, section_name) VALUES (%s, %s)",
                        (engagement_id, section))
    db.commit()

    return {"engagement_id": engagement_id, "message": "Engagement created with selected audit sections"}

# Update an existing engagement
@app.put("/engagements/{engagement_id}")
def update_engagement(engagement_id: int, e: Engagement, db=Depends(get_db)):
    validate_engagement_dates(e.start_date, e.end_date)
    cursor = db.cursor()
    cursor.execute(
        """UPDATE engagements SET client_id=%s, engagement_name=%s, financial_year=%s,
           status=%s, start_date=%s, end_date=%s WHERE engagement_id=%s""",
        (e.client_id, e.engagement_name, e.financial_year, e.status, e.start_date, e.end_date, engagement_id)
    )
    db.commit()
    return {"message": "Engagement updated"}

    return {"engagement_id": engagement_id, "message": "Engagement created with default audit sections"}
# Update an existing engagement
@app.put("/engagements/{engagement_id}")
def update_engagement(engagement_id: int, e: Engagement, db=Depends(get_db)):
       validate_engagement_update(
           db, engagement_id, e.client_id, e.engagement_name, e.financial_year, e.start_date, e.end_date
       )
       cursor = db.cursor()
       cursor.execute(
           """UPDATE engagements SET client_id=%s, engagement_name=%s, financial_year=%s,
              status=%s, start_date=%s, end_date=%s WHERE engagement_id=%s""",
           (e.client_id, e.engagement_name, e.financial_year, e.status, e.start_date, e.end_date, engagement_id)
       )
       db.commit()
       return {"message": "Engagement updated"}


# Mark an engagement as sent to the client. Only allowed once every section's
# latest submission is Approved — i.e. the derived status is "Under Review".
@app.put("/engagements/{engagement_id}/send-to-client")
def send_engagement_to_client(
    engagement_id: int,
    db=Depends(get_db),
    current_user: dict = Depends(require_role("Engagement Partner")),
):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT e.*, c.company_name, c.email as client_email, c.contact_person 
        FROM engagements e
        LEFT JOIN clients c ON e.client_id = c.client_id
        WHERE e.engagement_id = %s
    """, (engagement_id,))
    engagement = cursor.fetchone()
    if not engagement:
        raise HTTPException(status_code=404, detail="Engagement not found")
    if not engagement.get("client_email"):
        raise HTTPException(status_code=400, detail="Client has no email address on record")

    progress_by_id = fetch_engagement_progress(db, [engagement_id])
    apply_display_status(engagement, progress_by_id.get(engagement_id, {}))
    if engagement["display_status"] not in {"Under Review", "Completed"}:
        raise HTTPException(
            status_code=400,
            detail=(
                "Cannot send to client until every section's latest submission is "
                f"Approved (current status: {engagement['display_status']})."
            )
        )

    # Get approved sections for the email
    cursor.execute("""
        SELECT sec.section_name, s.status, s.notes
        FROM submissions s
        LEFT JOIN audit_sections sec ON s.section_id = sec.section_id
        WHERE s.engagement_id = %s AND s.status = 'Approved'
        ORDER BY sec.section_name
    """, (engagement_id,))
    approved_sections = cursor.fetchall()

    if not approved_sections:
        raise HTTPException(status_code=400, detail="No approved sections found for this engagement")

    sections_html = "".join([
        f"<tr><td style='padding:8px;border:1px solid #ddd'>{s['section_name']}</td>"
        f"<td style='padding:8px;border:1px solid #ddd;color:green'>Approved</td></tr>"
        for s in approved_sections
    ])

    html_body = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <h2 style="color: #1E3A5F;">Audit Report — {engagement['engagement_name']}</h2>
        <p>Dear {engagement.get('contact_person') or engagement.get('company_name')},</p>
        <p>We are pleased to inform you that the following audit sections for <strong>{engagement['engagement_name']}</strong>
        (Financial Year {engagement['financial_year']}) have been completed and approved:</p>
        <table style="width:100%;border-collapse:collapse;margin:16px 0">
            <thead>
                <tr style="background:#1E3A5F;color:white">
                    <th style="padding:10px;text-align:left">Section</th>
                    <th style="padding:10px;text-align:left">Status</th>
                </tr>
            </thead>
            <tbody>{sections_html}</tbody>
        </table>
        <p>Please contact us if you have any questions regarding this audit.</p>
        <br>
        <p style="color:#7f8c8d;font-size:12px">This is an automated message from Audit AI.</p>
    </div>
    """

    # Send the email
    msg = MIMEMultipart()
    msg["From"] = GMAIL_USER
    msg["To"] = engagement["client_email"]
    msg["Subject"] = f"Audit Report — {engagement['engagement_name']} (FY {engagement['financial_year']})"
    msg.attach(MIMEText(html_body, "html"))

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(GMAIL_USER, GMAIL_APP_PASSWORD)
            server.sendmail(GMAIL_USER, engagement["client_email"], msg.as_string())
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to send email: {str(e)}")

    # Add notifications for the engagement team
    cursor2 = db.cursor()
    cursor.execute("""
        SELECT u.user_id FROM users u
        INNER JOIN engagement_team et ON u.user_id = et.user_id
        WHERE et.engagement_id = %s
    """, (engagement_id,))
    for row in cursor.fetchall():
        cursor2.execute(
            "INSERT INTO notifications (user_id, message, type) VALUES (%s, %s, %s)",
            (row['user_id'],
             f"Audit report for {engagement['engagement_name']} has been sent to {engagement.get('company_name')}",
             "engagement_alert")
        )

    # Mark engagement as sent to client
    cursor2.execute(
        "UPDATE engagements SET sent_to_client_at = NOW(), status = 'Completed' WHERE engagement_id = %s",
        (engagement_id,)
    )
    db.commit()
    
    return {"message": "Engagement marked as sent to client and email sent", "display_status": "Completed"}

# Delete an engagement along with its audit sections and team assignments
@app.delete("/engagements/{engagement_id}")
def delete_engagement(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM audit_sections WHERE engagement_id = %s", (engagement_id,))
    cursor.execute("DELETE FROM engagement_team WHERE engagement_id = %s", (engagement_id,))
    cursor.execute("DELETE FROM engagements WHERE engagement_id = %s", (engagement_id,))
    db.commit()
    return {"message": "Engagement deleted"}

# List the team members assigned to an engagement
@app.get("/engagements/{engagement_id}/team")
def get_engagement_team(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT et.*, u.full_name, u.email, u.role FROM engagement_team et
        LEFT JOIN users u ON et.user_id = u.user_id
        WHERE et.engagement_id = %s
    """, (engagement_id,))
    return cursor.fetchall()

# Add a team member to an engagement
@app.post("/engagements/{engagement_id}/team")
def add_team_member(engagement_id: int, t: EngagementTeam, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("INSERT INTO engagement_team (engagement_id, user_id, role) VALUES (%s, %s, %s)",
                   (engagement_id, t.user_id, t.role))
    if t.role == "Accountant":
        info_cursor = db.cursor(dictionary=True)
        info_cursor.execute("SELECT engagement_name FROM engagements WHERE engagement_id = %s", (engagement_id,))
        info = info_cursor.fetchone()
        if info:
            cursor.execute(
                "INSERT INTO notifications (user_id, message, type, engagement_id) VALUES (%s, %s, %s, %s)",
                (t.user_id, f"You have been assigned to {info['engagement_name']}", "engagement_assigned", engagement_id)
            )
    db.commit()
    return {"team_id": cursor.lastrowid, "message": "Team member added"}

# Remove a team member from an engagement
@app.delete("/engagements/{engagement_id}/team/{user_id}")
def remove_team_member(engagement_id: int, user_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM engagement_team WHERE engagement_id=%s AND user_id=%s",
                   (engagement_id, user_id))
    db.commit()
    return {"message": "Team member removed"}

# List the audit sections for an engagement, with the assigned user's name
@app.get("/engagements/{engagement_id}/sections")
def get_audit_sections(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT s.*, u.full_name as assigned_to_name FROM audit_sections s
        LEFT JOIN users u ON s.assigned_to = u.user_id
        WHERE s.engagement_id = %s
    """, (engagement_id,))
    return cursor.fetchall()

# Add a new audit section to an engagement
@app.post("/engagements/{engagement_id}/sections")
def add_audit_section(engagement_id: int, s: AuditSection, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        "INSERT INTO audit_sections (engagement_id, section_name, status, assigned_to) VALUES (%s, %s, %s, %s)",
        (engagement_id, s.section_name, s.status, s.assigned_to)
    )
    db.commit()
    return {"section_id": cursor.lastrowid, "message": "Audit section added"}
@app.put("/audit-sections/{section_id}/scope")
def update_section_scope(section_id: int, body: SectionScopeUpdate, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
 
    # Check whether this section already has an approved submission.
    cursor.execute("""
        SELECT status FROM submissions
        WHERE section_id = %s
        ORDER BY created_at DESC
        LIMIT 1
    """, (section_id,))
    latest = cursor.fetchone()
 
    if latest and latest["status"] == "Approved" and not body.confirm_override:
        raise HTTPException(
            status_code=409,
            detail=(
                "This section already has an approved submission. "
                "Changing its scope will exclude already-reviewed work from "
                "final analysis and reports. Pass confirm_override=true to proceed."
            ),
        )
 
    cursor.execute(
        """
        UPDATE audit_sections
        SET in_scope = %s, scope_reason = %s, scope_set_by = %s, scope_set_at = NOW()
        WHERE section_id = %s
        """,
        (body.in_scope, body.scope_reason, body.set_by, section_id),
    )
    db.commit()
    return {"message": "Section scope updated", "in_scope": body.in_scope}
# Update an audit section
@app.put("/audit-sections/{section_id}")
def update_audit_section(section_id: int, s: AuditSection, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        "UPDATE audit_sections SET section_name=%s, status=%s, assigned_to=%s WHERE section_id=%s",
        (s.section_name, s.status, s.assigned_to, section_id)
    )
    db.commit()
    return {"message": "Audit section updated"}

# Delete an audit section
@app.delete("/audit-sections/{section_id}")
def delete_audit_section(section_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM audit_sections WHERE section_id = %s", (section_id,))
    db.commit()
    return {"message": "Audit section deleted"}

# Get the most recent submission for a single audit section, or none if it has never been submitted
@app.get("/audit-sections/{section_id}/latest-submission")
def get_section_latest_submission(section_id: int, file_id: str = Query(None), db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    if file_id:
        cursor.execute(
            "SELECT s.*, u.full_name as submitted_by_name "
            "FROM submissions s "
            "LEFT JOIN users u ON s.submitted_by = u.user_id "
            "WHERE s.section_id = %s AND s.file_id = %s "
            "ORDER BY s.created_at DESC LIMIT 1",
            (section_id, file_id)
        )
    else:
        cursor.execute(
            "SELECT s.*, u.full_name as submitted_by_name "
            "FROM submissions s "
            "LEFT JOIN users u ON s.submitted_by = u.user_id "
            "WHERE s.section_id = %s "
            "ORDER BY s.created_at DESC LIMIT 1",
            (section_id,)
        )
    row = cursor.fetchone()
    return row if row else None

# List every submission across all engagements, most recent first
@app.get("/submissions")
def get_all_submissions(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT s.*, u.full_name as submitted_by_name, e.engagement_name, sec.section_name
        FROM submissions s
        LEFT JOIN users u ON s.submitted_by = u.user_id
        LEFT JOIN engagements e ON s.engagement_id = e.engagement_id
        LEFT JOIN audit_sections sec ON s.section_id = sec.section_id
        ORDER BY s.created_at DESC
    """)
    return cursor.fetchall()

# Get a single submission by id
@app.get("/submissions/{submission_id}")
def get_submission(submission_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT s.*, u.full_name as submitted_by_name, e.engagement_name, sec.section_name
        FROM submissions s
        LEFT JOIN users u ON s.submitted_by = u.user_id
        LEFT JOIN engagements e ON s.engagement_id = e.engagement_id
        LEFT JOIN audit_sections sec ON s.section_id = sec.section_id
        WHERE s.submission_id = %s
    """, (submission_id,))
    submission = cursor.fetchone()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")
    return submission

# Get comprehensive review data for a submission
@app.get("/submissions/{submission_id}/review-data")
def get_submission_review_data(submission_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT s.*, u.full_name as submitted_by_name, e.engagement_name, e.client_id,
               sec.section_name
        FROM submissions s
        LEFT JOIN users u ON s.submitted_by = u.user_id
        LEFT JOIN engagements e ON s.engagement_id = e.engagement_id
        LEFT JOIN audit_sections sec ON s.section_id = sec.section_id
        WHERE s.submission_id = %s
    """, (submission_id,))
    submission = cursor.fetchone()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")

    file_id = submission.get("file_id")
    client_id = str(submission.get("client_id"))

    result = {
        "submission": submission,
        "file": None,
        "cleaning_summary": None,
        "trial_balance_validation": None,
        "account_mapping": None,
        "saved_analysis": None,
    }

    if not file_id:
        return result

    cursor.execute("SELECT filename, semantic_file_type, file_type FROM uploads WHERE file_id = %s", (file_id,))
    upload = cursor.fetchone()
    if not upload:
        return result

    file_type = upload.get("semantic_file_type") or upload.get("file_type") or "general"
    result["file"] = {"file_id": file_id, "filename": upload.get("filename"), "file_type": file_type}

    # Try to get file-specific mapping first, fall back to file_type if not found
    try:
        mapping = get_file_specific_mapping(file_id, client_id, file_type)
    except HTTPException:
        # No file-specific mapping exists, try file_type fallback
        mapping = get_mapping(client_id, file_type)
        if not mapping:
            return result

    # Cleaning summary + TB validation are recomputed live from the current cleaned
    # state, same pattern used everywhere else in this codebase (never trust a stale snapshot)
    try:
        cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
        result["cleaning_summary"] = {
            "total_issues": report.get("total_issues", 0),
            "can_proceed": report.get("can_proceed", False),
            "flagged_rows": report.get("flagged_rows", 0),
            "clean_rows": report.get("clean_rows", 0),
        }
        is_ledger = file_type in ("trial_balance", "general_ledger")
        if is_ledger and report.get("can_proceed"):
            result["trial_balance_validation"] = validate_trial_balance(cleaned_df, mapping)
            account_mapping = get_account_mapping(client_id, file_type)
            if account_mapping:
                result["account_mapping"] = account_mapping
    except HTTPException:
        pass

    analyses = get_saved_analyses_for_file(submission["engagement_id"], file_id)
    if analyses:
        result["saved_analysis"] = analyses[0]  # most recent, per existing ORDER BY created_at DESC

    return result
    return submission

# Create a new submission, and notify the relevant users if it has moved past the first workflow stage

@app.post("/submissions")
def create_submission(s: Submission, db=Depends(get_db)):
    # Engagement must actually exist before anything is submitted against it
    # (doc #4: "Does engagement exist? ✓").
    validate_engagement_exists(db, s.engagement_id)

    # A submission with no real file behind it can never be analyzed later —
    # confirm the file_id actually exists before accepting the submission,
    # rather than storing a reference that turns out to be dead.
    verify_cursor = db.cursor(dictionary=True)
    verify_cursor.execute("SELECT file_id, file_type FROM uploads WHERE file_id = %s", (s.file_id,))
    upload_row = verify_cursor.fetchone()
    if not upload_row:
        raise HTTPException(
            status_code=400,
            detail=error_detail(
                f"No uploaded file found with file_id '{s.file_id}'. Upload the file before submitting.",
                error_code="NOT_FOUND",
                details={"file_id": s.file_id},
            ),
        )

    # Section must belong to this engagement AND currently be in scope
    # (doc #4: "Is section in scope? ✓") — reject rather than silently
    # expanding the engagement's scope to cover it.
    validate_section_in_engagement_scope(db, s.engagement_id, s.section_id)

    # If a specific sheet was named, it must actually exist inside the
    # uploaded workbook (doc #3: "Verify selected sheets exist in
    # workbook") — never trust a client-supplied sheet name.
    if s.sheet_name:
        file_path, ext = locate_uploaded_file(s.file_id, upload_row.get("file_type"))
        if not file_path:
            raise HTTPException(
                status_code=400,
                detail=error_detail(
                    f"Uploaded file '{s.file_id}' could not be located on disk.",
                    error_code="NOT_FOUND",
                    details={"file_id": s.file_id},
                ),
            )
        validate_selected_sheets(file_path, ext, [s.sheet_name])

    insert_cursor = db.cursor()
    insert_cursor.execute(
        """INSERT INTO submissions (engagement_id, section_id, submitted_by, file_id, sheet_name, status, current_stage, notes)
           VALUES (%s, %s, %s, %s, %s, %s, %s, %s)""",
        (s.engagement_id, s.section_id, s.submitted_by, s.file_id, s.sheet_name, s.status, s.current_stage, s.notes)
    )
    submission_id = insert_cursor.lastrowid
    cursor = db.cursor(dictionary=True)
    if s.current_stage and s.current_stage != "Accountant":
        cursor.execute("""
            SELECT e.engagement_name, sec.section_name FROM engagements e
            LEFT JOIN audit_sections sec ON sec.engagement_id = e.engagement_id
            WHERE sec.section_id = %s
        """, (s.section_id,))
        info = cursor.fetchone()
        if info:
            message = f"{info['section_name']} for {info['engagement_name']} is now {s.status}"
            cursor.execute("""
                SELECT DISTINCT u.user_id FROM users u
                INNER JOIN engagement_team et ON u.user_id = et.user_id
                WHERE et.engagement_id = %s AND COALESCE(NULLIF(et.role, ''), u.role) = %s
            """, (s.engagement_id, s.current_stage))
            for auditor in cursor.fetchall():
                cursor.execute(
                    "INSERT INTO notifications (user_id, message, type, engagement_id) VALUES (%s, %s, %s, %s)",
                    (auditor['user_id'], message, "submission_review", s.engagement_id)
                )
    db.commit()
    return {"submission_id": submission_id, "message": "Submission created"}
 
# Update a submission's status and workflow stage, and notify the relevant users of the change
@app.put("/submissions/{submission_id}/status")
def update_submission_status(submission_id: int, s: SubmissionStatus, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT sub.*, e.engagement_name, e.engagement_id, sec.section_name
        FROM submissions sub
        LEFT JOIN engagements e ON sub.engagement_id = e.engagement_id
        LEFT JOIN audit_sections sec ON sub.section_id = sec.section_id
        WHERE sub.submission_id = %s
    """, (submission_id,))
    sub = cursor.fetchone()
    if not sub:
        raise HTTPException(status_code=404, detail="Submission not found")

    cursor2 = db.cursor(dictionary=True)
    if s.updated_by:
        cursor2.execute(
            "UPDATE submissions SET status=%s, current_stage=%s, notes=%s, submitted_by=%s WHERE submission_id=%s",
            (s.status, s.current_stage, s.notes, s.updated_by, submission_id)
        )
    else:
        cursor2.execute(
            "UPDATE submissions SET status=%s, current_stage=%s, notes=%s WHERE submission_id=%s",
            (s.status, s.current_stage, s.notes, submission_id)
        )

    # Decide who should be notified: the new current stage's role, or every workflow role if the submission was just approved or cancelled
    target_roles = []
    if s.current_stage:
        target_roles = [s.current_stage]
    elif s.status in ("Approved", "Cancelled"):
        target_roles = ["Accountant", "Auditor", "Senior Auditor", "Assistant Manager",
                        "Audit Manager", "Engagement Partner", "Quality Reviewer"]

    if target_roles:
        message = f"{sub['section_name']} for {sub['engagement_name']} is now {s.status}"
        format_strings = ','.join(['%s'] * len(target_roles))
        cursor2.execute(f"""
            SELECT DISTINCT u.user_id FROM users u
            INNER JOIN engagement_team et ON u.user_id = et.user_id
            WHERE et.engagement_id = %s AND COALESCE(NULLIF(et.role, ''), u.role) IN ({format_strings})
        """, (sub['engagement_id'], *target_roles))
        for row in cursor2.fetchall():
            cursor2.execute(
                "INSERT INTO notifications (user_id, message, type, engagement_id) VALUES (%s, %s, %s, %s)",
                (row['user_id'], message, "submission_review", sub['engagement_id'])
            )

    db.commit()
    if s.status == "Approved":
        try:
            progress_by_id = fetch_engagement_progress(db, [sub['engagement_id']])
            progress = progress_by_id.get(sub['engagement_id'], {})
            if progress.get("total_sections", 0) > 0 and progress.get("approved_sections", 0) == progress.get("total_sections", 0):
                complete_cursor = db.cursor(dictionary=True)
                complete_cursor.execute("""
                    SELECT DISTINCT u.user_id FROM users u
                    INNER JOIN engagement_team et ON u.user_id = et.user_id
                    WHERE et.engagement_id = %s AND COALESCE(NULLIF(et.role, ''), u.role) IN
                        ('Auditor', 'Senior Auditor', 'Assistant Manager', 'Audit Manager', 'Engagement Partner', 'Quality Reviewer')
                """, (sub['engagement_id'],))
                complete_message = f"All in-scope sections for {sub['engagement_name']} are now approved. Ready for final review."
                write_cursor = db.cursor()
                for row in complete_cursor.fetchall():
                    write_cursor.execute(
                        "INSERT INTO notifications (user_id, message, type, engagement_id) VALUES (%s, %s, %s, %s)",
                        (row['user_id'], complete_message, "engagement_ready", sub['engagement_id'])
                    )
                db.commit()
            notify_if_ready_for_final_analysis(db, sub["engagement_id"])
        except Exception as e:
            # The status update itself already committed successfully above 
            print(f"WARNING: post-approval notification step failed for submission {submission_id}: {e}")
    return {"message": f"Submission status updated to {s.status}"}

# Delete a submission
@app.delete("/submissions/{submission_id}")
def delete_submission(submission_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM submissions WHERE submission_id = %s", (submission_id,))
    db.commit()
    return {"message": "Submission deleted"}

# List all notifications for a user, most recent first
@app.get("/notifications/{user_id}")
def get_user_notifications(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM notifications WHERE user_id = %s ORDER BY created_at DESC", (user_id,))
    notifications = cursor.fetchall()
    return notifications

# List only the unread notifications for a user, most recent first
@app.get("/notifications/{user_id}/unread")
def get_unread_notifications(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM notifications WHERE user_id = %s AND is_read = FALSE ORDER BY created_at DESC",
                   (user_id,))
    return cursor.fetchall()

# Mark a single notification as read
@app.put("/notifications/{notification_id}/read")
def mark_notification_read(notification_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("UPDATE notifications SET is_read = TRUE WHERE notification_id = %s", (notification_id,))
    db.commit()
    return {"message": "Notification marked as read"}

# Mark every notification for a user as read
@app.put("/notifications/{user_id}/read-all")
def mark_all_read(user_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("UPDATE notifications SET is_read = TRUE WHERE user_id = %s", (user_id,))
    db.commit()
    return {"message": "All notifications marked as read"}

# Upload a general client document (separate from the AI pipeline upload), storing it on disk and recording it in the uploads table
@app.post("/clients/{client_id}/upload")
def upload_client_file(client_id: int, file: UploadFile = File(...), db=Depends(get_db)):
    allowed_types = ["xlsx", "xls", "csv", "pdf", "tiff", "tif", "jpg", "jpeg", "png", "xml", "json", "txt"]
    file_ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if file_ext not in allowed_types:
        raise HTTPException(status_code=400, detail="File format not allowed.")
    file_path = f"{UPLOAD_DIR}/{client_id}_{file.filename}"
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    cursor = db.cursor()
    cursor.execute("INSERT INTO uploads (client_id, file_name, file_type, file_path) VALUES (%s, %s, %s, %s)",
                   (client_id, file.filename, file_ext.upper(), file_path))
    db.commit()
    return {"file_id": cursor.lastrowid, "filename": file.filename, "type": file_ext.upper(),
            "message": "File uploaded successfully"}

# List all files uploaded for a single client
@app.get("/clients/{client_id}/files")
def get_client_files(client_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM uploads WHERE client_id = %s", (client_id,))
    return cursor.fetchall()

# List every file uploaded across all clients, with the client's company name
@app.get("/files")
def get_all_files(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT f.*, c.company_name FROM uploads f
        LEFT JOIN clients c ON f.client_id = c.client_id
        ORDER BY f.upload_date DESC
    """)
    return cursor.fetchall()

# Save a financial analysis result
@app.post("/saved-analyses")
async def save_analysis_endpoint(
    user_id: int = Form(...),
    client_id: str = Form(...),
    engagement_id: int = Form(None),
    file_id: str = Form(...),
    file_type: str = Form("general"),
    analysis_data: str = Form(...),
    insights_data: str = Form(None)
):
    import json
    analysis_dict = json.loads(analysis_data) if analysis_data else None
    insights_list = json.loads(insights_data) if insights_data else None
    
    analysis_id = save_analysis(user_id, client_id, engagement_id, file_id, file_type, analysis_dict, insights_list)
    return {"analysis_id": analysis_id, "message": "Analysis saved successfully"}

# Get all saved analyses for a user
@app.get("/saved-analyses/{user_id}")
def get_user_saved_analyses(user_id: int):
    analyses = get_saved_analyses(user_id)
    return analyses

# Get all saved analyses for an engagement (team-scoped — every snapshot
# saved by anyone on this engagement, not just the requesting user), most
# recent first per file, so the frontend can group by file_id and show a
# history with who saved each one and when.
@app.get("/engagements/{engagement_id}/saved-analyses")
def get_engagement_saved_analyses(engagement_id: int):
    analyses = get_saved_analyses_for_engagement(engagement_id)
    return analyses

# Get saved analyses for a specific file within an engagement — this is
# what an auditor needs when working on a particular file. Returns all
# analysis snapshots for that specific file, with attribution.
@app.get("/engagements/{engagement_id}/saved-analyses/{file_id}")
def get_file_saved_analyses(engagement_id: int, file_id: str):
    analyses = get_saved_analyses_for_file(engagement_id, file_id)
    return analyses

# ── Engagement-level final analysis: on-demand, never persists ──────────────
@app.get("/engagements/{engagement_id}/final-analysis")
def get_engagement_final_analysis(engagement_id: int, db=Depends(get_db)):
    return build_engagement_final_analysis(engagement_id)


# ── Engagement-level final analysis: explicit save, locks the snapshot ──────
@app.post("/engagements/{engagement_id}/final-analysis/save")
def save_engagement_final_analysis(engagement_id: int, body: SaveFinalAnalysis, db=Depends(get_db)):
    # Recompute rather than trust a client-supplied payload — the saved
    # snapshot must reflect what's actually approved right now, not
    # whatever the caller happened to send.
    result = build_engagement_final_analysis(engagement_id)

    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO engagement_final_analysis
           (engagement_id, saved_by, analysis_data, insights_data, included_sections)
           VALUES (%s, %s, %s, %s, %s)""",
        (
            engagement_id,
            body.saved_by,
            json.dumps({
                "financial_statements": result["financial_statements"],
                "breakdowns": result["breakdowns"],
                "monthly_trend": result["monthly_trend"],
                "anomalies": result["anomalies"],
            }),
            json.dumps(result["ai_insights"]),
            json.dumps(result["included_sections"]),
        )
    )
    db.commit()
    return {"analysis_id": cursor.lastrowid, "message": "Final analysis saved.", **result}

# Get a specific saved analysis
@app.get("/saved-analyses/{analysis_id}/view")
def get_analysis_by_id(analysis_id: int):
    analysis = get_saved_analysis(analysis_id)
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return analysis

# Delete a saved analysis
@app.delete("/saved-analyses/{analysis_id}")
def delete_analysis_endpoint(analysis_id: int):
    delete_saved_analysis(analysis_id)
    return {"message": "Analysis deleted successfully"}

# --- Auditor Workspace Pydantic models and API endpoints ---

class WorkspaceOpenRequest(BaseModel):
    user_id: int
    engagement_id: Optional[int] = None
    section_id: Optional[int] = None
    file_id: Optional[str] = None
    client_id: Optional[str] = None

class WorkspaceUpdateRequest(BaseModel):
    status: Optional[str] = None
    notes: Optional[str] = None
    progress_data: Optional[dict] = None
    file_id: Optional[str] = None

class WorkspaceSubmitRequest(BaseModel):
    submitted_by: int
    notes: Optional[str] = None

@app.post("/workspaces/open")
def open_workspace_endpoint(req: WorkspaceOpenRequest, db=Depends(get_db)):
    engagement_id = req.engagement_id
    cursor = db.cursor(dictionary=True)
    if not engagement_id:
        if req.file_id:
            # Resolve engagement through the file's tagged section — reliable,
            # unlike guessing by client_id alone which breaks with multiple engagements
            cursor.execute("""
                SELECT sec.engagement_id FROM uploads u
                JOIN audit_sections sec ON u.section_id = sec.section_id
                WHERE u.file_id = %s LIMIT 1
            """, (req.file_id,))
            eng = cursor.fetchone()
            if eng:
                engagement_id = eng['engagement_id']
        if not engagement_id and req.client_id:
            cursor.execute("SELECT engagement_id FROM engagements WHERE CAST(client_id AS CHAR) = %s LIMIT 1", (str(req.client_id),))
            eng = cursor.fetchone()
            if eng:
                engagement_id = eng['engagement_id']
        if not engagement_id:
            cursor.execute("SELECT engagement_id FROM engagements ORDER BY created_at DESC LIMIT 1")
            eng = cursor.fetchone()
            if eng:
                engagement_id = eng['engagement_id']
            else:
                engagement_id = 1

    # NEW: resolve section_id — prefer the file's own tagged section (set by the
    # Accountant at upload time) over guessing from auditor assignment
    section_id = req.section_id
    cursor = db.cursor(dictionary=True)
    if not section_id and req.file_id:
        cursor.execute("SELECT section_id FROM uploads WHERE file_id = %s", (req.file_id,))
        row = cursor.fetchone()
        if row and row.get("section_id"):
            section_id = row["section_id"]

    if not section_id:
        cursor.execute(
            "SELECT section_id FROM audit_sections WHERE engagement_id = %s AND assigned_to = %s LIMIT 1",
            (engagement_id, req.user_id)
        )
        sec = cursor.fetchone()
        if sec:
            section_id = sec["section_id"]

    ws = get_or_create_workspace(
        user_id=req.user_id,
        engagement_id=engagement_id,
        section_id=section_id,
        file_id=req.file_id
    )
    if not ws:
        raise HTTPException(status_code=400, detail="Failed to open workspace")
    return ws

@app.get("/workspaces/{workspace_id}")
def get_workspace_endpoint(workspace_id: int):
    ws = get_workspace_by_id(workspace_id)
    if not ws:
        raise HTTPException(status_code=404, detail="Workspace not found")
    return ws

@app.put("/workspaces/{workspace_id}")
def update_workspace_endpoint(workspace_id: int, req: WorkspaceUpdateRequest):
    update_workspace_data(
        workspace_id=workspace_id,
        status=req.status,
        notes=req.notes,
        progress_data=req.progress_data,
        file_id=req.file_id
    )
    ws = get_workspace_by_id(workspace_id)
    return ws

@app.get("/engagements/{engagement_id}/workspaces")
def get_engagement_workspaces_endpoint(engagement_id: int):
    return get_engagement_workspaces(engagement_id)

@app.get("/users/{user_id}/workspaces")
def get_user_workspaces_endpoint(user_id: int):
    return get_user_workspaces(user_id)

@app.post("/workspaces/{workspace_id}/submit-for-review")
def submit_workspace_for_review(workspace_id: int, req: WorkspaceSubmitRequest, db=Depends(get_db)):
    ws = get_workspace_by_id(workspace_id)
    if not ws:
        raise HTTPException(status_code=404, detail="Workspace not found")

    section_id = ws.get("section_id")
    engagement_id = ws.get("engagement_id")
    cursor = db.cursor(dictionary=True)

    if not section_id and ws.get("file_id"):
        cursor.execute("SELECT section_id FROM uploads WHERE file_id = %s", (ws["file_id"],))
        row = cursor.fetchone()
        if row and row.get("section_id"):
            section_id = row["section_id"]

    if not section_id:
        cursor.execute(
            "SELECT section_id FROM audit_sections WHERE engagement_id = %s AND assigned_to = %s LIMIT 1",
            (engagement_id, req.submitted_by)
        )
        sec = cursor.fetchone()
        if not sec:
            raise HTTPException(
                status_code=400,
                detail="No audit section is assigned to you for this engagement. Please contact your Audit Manager."
            )
        section_id = sec["section_id"]
        fix_cursor = db.cursor()
        fix_cursor.execute("UPDATE auditor_workspaces SET section_id = %s WHERE workspace_id = %s", (section_id, workspace_id))
        db.commit()

    # Section must be in scope for this engagement before any submission is
    # written or updated against it (doc #4) — same rule as POST /submissions,
    # applied here since this endpoint is a second real path that writes to
    # the submissions table.
    if not section_id:
        raise HTTPException(
            status_code=400,
            detail=error_detail(
                "No audit section could be resolved for this workspace. "
                "Please contact your Audit Manager.",
                error_code="VALIDATION_ERROR",
                details={"workspace_id": workspace_id, "engagement_id": engagement_id},
            ),
        )
    validate_section_in_engagement_scope(db, engagement_id, section_id)

    file_id = ws.get("file_id")

    # Get sheet_name from uploads table for Excel files
    sheet_name = None
    if file_id:
        cursor.execute("SELECT sheet_name FROM uploads WHERE file_id = %s", (file_id,))
        upload_row = cursor.fetchone()
        if upload_row:
            sheet_name = upload_row.get("sheet_name")

    # Reuse the existing submission for this specific FILE (not just section) if one exists
    cursor.execute(
        "SELECT * FROM submissions WHERE section_id = %s AND file_id = %s ORDER BY created_at DESC LIMIT 1",
        (section_id, file_id)
    )
    existing = cursor.fetchone()

    # Block resubmission of a file that's already been submitted and hasn't
    # been sent back for corrections — only Draft/Cancelled/Changes Requested
    # (or no prior submission at all) may be (re)submitted
    if existing and existing["status"] in ("Submitted", "Under Review", "Approved"):
        raise HTTPException(
            status_code=400,
            detail=(
                f"This file has already been submitted and is currently "
                f"'{existing['status']}'. It can only be resubmitted after "
                f"changes are requested by a reviewer."
            )
        )

    write_cursor = db.cursor()
    if existing:
        write_cursor.execute(
            "UPDATE submissions SET status = %s, current_stage = %s, notes = %s, submitted_by = %s, sheet_name = %s WHERE submission_id = %s",
            ("Submitted", "Accountant", req.notes, req.submitted_by, sheet_name, existing["submission_id"])
        )
        submission_id = existing["submission_id"]
    else:
        write_cursor.execute(
            "INSERT INTO submissions (engagement_id, section_id, file_id, sheet_name, submitted_by, status, current_stage, notes) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)",
            (engagement_id, section_id, file_id, sheet_name, req.submitted_by, "Submitted", "Accountant", req.notes)
        )
        submission_id = write_cursor.lastrowid
    db.commit()

    # Notify the Accountant-stage reviewers, same pattern as update_submission_status
    cursor.execute("""
        SELECT e.engagement_name, sec.section_name FROM engagements e
        LEFT JOIN audit_sections sec ON sec.engagement_id = e.engagement_id
        WHERE sec.section_id = %s
    """, (section_id,))
    info = cursor.fetchone()
    if info:
        message = f"{info['section_name']} for {info['engagement_name']} has been submitted for review"
        cursor.execute("""
            SELECT DISTINCT u.user_id FROM users u
            INNER JOIN engagement_team et ON u.user_id = et.user_id
            WHERE et.engagement_id = %s AND COALESCE(NULLIF(et.role, ''), u.role) = 'Accountant'
        """, (engagement_id,))
        for row in cursor.fetchall():
            write_cursor.execute(
                "INSERT INTO notifications (user_id, message, type, engagement_id) VALUES (%s, %s, %s, %s)",
                (row["user_id"], message, "submission_review", engagement_id)
            )
        db.commit()
    return {"submission_id": submission_id, "section_id": section_id, "message": "Submitted for review successfully."}

# Run the app directly with uvicorn when this file is executed as a script
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
