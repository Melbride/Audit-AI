from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, Literal
from passlib.context import CryptContext
from jose import JWTError, jwt
from datetime import datetime, timedelta
import pdfplumber
from docx import Document
import pandas as pd
import json
import uuid
import shutil
import os
import sys
import secrets
import hashlib
from dotenv import load_dotenv
from fastapi_mail import FastMail, MessageSchema, ConnectionConfig, MessageType
from detector import detect_columns_with_llm, build_detection_result, suggest_file_type
from database import (
    get_db, init_db, save_mapping, get_mapping, save_upload, get_uploads,
    save_cleaning_acknowledgment, get_acknowledged_issue_ids,
    save_cleaning_corrections, get_cleaning_corrections,
    save_fingerprint, get_fingerprint,
    save_cleaning_snapshot, get_cleaning_snapshot,
)
from cleaner import clean_dataframe
from report_routes import router as report_router
load_dotenv()


mail_conf = ConnectionConfig(
    MAIL_USERNAME=os.getenv("MAIL_USERNAME"),
    MAIL_PASSWORD=os.getenv("MAIL_PASSWORD"),
    MAIL_FROM=os.getenv("MAIL_FROM"),
    MAIL_FROM_NAME=os.getenv("MAIL_FROM_NAME", "Audit AI"),
    MAIL_PORT=587,
    MAIL_SERVER="smtp.gmail.com",
    MAIL_STARTTLS=True,
    MAIL_SSL_TLS=False,
    USE_CREDENTIALS=True,
)

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BACKEND_DIR = os.path.join(BASE_DIR, "backend")
if BACKEND_DIR not in sys.path:
    sys.path.append(BACKEND_DIR)



app = FastAPI(title="AuditIQ API", debug=True)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.include_router(report_router)

SECRET_KEY = "auditiq-secret-key-change-in-production"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_HOURS = 8
pwd_context = CryptContext(schemes=["sha256_crypt"], deprecated="auto")

def hash_password(password: str):
    return pwd_context.hash(password)

def verify_password(plain: str, hashed: str):
    return pwd_context.verify(plain, hashed)

def create_token(data: dict):
    expire = datetime.utcnow() + timedelta(hours=ACCESS_TOKEN_EXPIRE_HOURS)
    data.update({"exp": expire})
    return jwt.encode(data, SECRET_KEY, algorithm=ALGORITHM)

@app.get("/")
def root():
    return {"message": "Audit AI API is running"}

@app.on_event("startup")
def startup_event():
    init_db()

# ── MODELS ────────────────────────────────────────────────────────────────────
class Client(BaseModel):
    company_name: str
    contact_person: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    industry: Optional[str] = None
    address: Optional[str] = None
    status: Optional[str] = "Active"
    kra_pin: Literal[True, False] = False
    kra_pin_number: Optional[str] = None

class User(BaseModel):
    full_name: str
    email: str
    password: str
    phone: Optional[str] = None
    role: Literal["Admin", "Accountant", "Auditor", "Senior Auditor", "Assistant Manager", "Audit Manager", "Engagement Partner", "Quality Reviewer"]
    assigned_client_id: Optional[int] = None
    status: Optional[str] = "Active"

class UserUpdate(BaseModel):
    full_name: str
    email: str
    phone: Optional[str] = None
    role: Literal["Admin", "Accountant", "Auditor", "Senior Auditor", "Assistant Manager", "Audit Manager", "Engagement Partner", "Quality Reviewer"]
    assigned_client_id: Optional[int] = None
    status: Optional[str] = "Active"

class LoginRequest(BaseModel):
    email: str
    password: str

class PasswordResetRequest(BaseModel):
    email: str

class PasswordResetConfirm(BaseModel):
    token: str
    new_password: str

class ColumnMapping(BaseModel):
    client_id: str
    file_type: Optional[str] = "general"
    original_column: str
    mapped_to: str
    confirmed_by: Optional[str] = None

class Engagement(BaseModel):
    client_id: int
    engagement_name: str
    financial_year: str
    status: Optional[str] = "Planning"
    start_date: Optional[str] = None
    end_date: Optional[str] = None

class EngagementTeam(BaseModel):
    engagement_id: int
    user_id: int
    role: str

class AuditSection(BaseModel):
    engagement_id: int
    section_name: str
    status: Optional[str] = "Pending"
    assigned_to: Optional[int] = None

class Submission(BaseModel):
    engagement_id: int
    section_id: int
    submitted_by: int
    status: Optional[str] = "Draft"
    current_stage: Optional[str] = "Accountant"
    notes: Optional[str] = None

class SubmissionStatus(BaseModel):
    status: Literal["Draft", "Submitted", "Under Review", "Changes Requested", "Approved", "Cancelled"]
    current_stage: Optional[str] = None
    notes: Optional[str] = None
    updated_by: Optional[int] = None

class Notification(BaseModel):
    user_id: int
    message: str
    type: Optional[str] = "engagement_alert"

# ── CLIENTS ───────────────────────────────────────────────────────────────────
@app.get("/clients")
def get_clients(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM clients")
    return cursor.fetchall()

@app.get("/clients/{client_id}")
def get_client(client_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM clients WHERE client_id = %s", (client_id,))
    client = cursor.fetchone()
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    return client

@app.post("/clients")
def create_client(c: Client, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO clients (company_name, contact_person, email, phone, industry, address, status, kra_pin, kra_pin_number)
           VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)""",
        (c.company_name, c.contact_person, c.email, c.phone, c.industry, c.address, c.status, c.kra_pin, c.kra_pin_number)
    )
    db.commit()
    return {"client_id": cursor.lastrowid, "message": "Client created"}
@app.put("/clients/{client_id}")
def update_client(client_id: int, c: Client, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """UPDATE clients SET company_name=%s, contact_person=%s, email=%s,
           phone=%s, industry=%s, address=%s, status=%s, kra_pin=%s, kra_pin_number=%s WHERE client_id=%s""",
        (c.company_name, c.contact_person, c.email, c.phone, c.industry, c.address, c.status, c.kra_pin, c.kra_pin_number, client_id)
    )
    db.commit()
    return {"message": "Client updated"}

@app.delete("/clients/{client_id}")
def delete_client(client_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    try:
        cursor.execute("SELECT engagement_id FROM engagements WHERE client_id = %s", (client_id,))
        rows = cursor.fetchall()
        engagement_ids = [r['engagement_id'] for r in rows] if rows else []
        if engagement_ids:
            placeholders = ','.join(['%s'] * len(engagement_ids))
            cursor.execute(f"DELETE FROM submissions WHERE engagement_id IN ({placeholders})", tuple(engagement_ids))
            cursor.execute(f"DELETE FROM audit_sections WHERE engagement_id IN ({placeholders})", tuple(engagement_ids))
            cursor.execute(f"DELETE FROM engagement_team WHERE engagement_id IN ({placeholders})", tuple(engagement_ids))
            cursor.execute(f"DELETE FROM engagements WHERE engagement_id IN ({placeholders})", tuple(engagement_ids))
        cursor.execute("DELETE FROM uploads WHERE client_id = %s", (str(client_id),))
        cursor.execute("UPDATE users SET assigned_client_id = NULL WHERE assigned_client_id = %s", (client_id,))
        cursor.execute("DELETE FROM clients WHERE client_id = %s", (client_id,))
        if cursor.rowcount == 0:
            db.rollback()
            raise HTTPException(status_code=404, detail="Client not found")
        db.commit()
        return {"message": "Client and dependent records deleted"}
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Could not delete client: {str(e)}")

# ── USERS ─────────────────────────────────────────────────────────────────────
@app.get("/users")
def get_users(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT u.user_id, u.full_name, u.email, u.phone, u.role, u.status,
               u.assigned_client_id, u.login_locked, u.last_login, u.created_at, c.company_name
        FROM users u
        LEFT JOIN clients c ON u.assigned_client_id = c.client_id
    """)
    return cursor.fetchall()

@app.put("/users/{user_id}/reset-password")
def reset_user_password(user_id: int, payload: dict, db=Depends(get_db)):
    new_password = payload.get("new_password")
    if not new_password or len(new_password) < 8:
        raise HTTPException(status_code=400, detail="New password must be at least 8 characters")
    cursor = db.cursor()
    hashed = hash_password(new_password)
    cursor.execute("UPDATE users SET password_hash = %s, failed_attempts = 0 WHERE user_id = %s", (hashed, user_id))
    db.commit()
    return {"message": "Password reset successful"}

@app.put("/users/{user_id}/login-lock")
def set_user_login_lock(user_id: int, payload: dict, db=Depends(get_db)):
    locked = payload.get("locked")
    if locked is None:
        raise HTTPException(status_code=400, detail="Missing locked flag")
    cursor = db.cursor()
    cursor.execute("UPDATE users SET login_locked = %s WHERE user_id = %s", (1 if locked else 0, user_id))
    db.commit()
    return {"message": "Login access updated", "login_locked": bool(locked)}

@app.get("/users/{user_id}/login-history")
def get_user_login_history(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute(
        "SELECT id, timestamp, ip_address, device, status FROM login_history WHERE user_id = %s ORDER BY timestamp DESC",
        (user_id,)
    )
    return cursor.fetchall()

@app.get("/users/{user_id}")
def get_user(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE user_id = %s", (user_id,))
    user = cursor.fetchone()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    user.pop("password_hash", None)
    return user

@app.post("/users")
def create_user(u: User, db=Depends(get_db)):
    hashed = hash_password(u.password)
    cursor = db.cursor()
    try:
        cursor.execute(
            """INSERT INTO users (full_name, email, password_hash, phone, role, assigned_client_id, status)
               VALUES (%s, %s, %s, %s, %s, %s, %s)""",
            (u.full_name, u.email, hashed, u.phone, u.role, u.assigned_client_id, u.status)
        )
        db.commit()
        return {"user_id": cursor.lastrowid, "message": "User created"}
    except Exception:
        raise HTTPException(status_code=400, detail="Email already exists")

@app.put("/users/{user_id}")
def update_user(user_id: int, u: UserUpdate, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """UPDATE users SET full_name=%s, email=%s, phone=%s,
           role=%s, assigned_client_id=%s, status=%s WHERE user_id=%s""",
        (u.full_name, u.email, u.phone, u.role, u.assigned_client_id, u.status, user_id)
    )
    db.commit()
    return {"message": "User updated"}

@app.delete("/users/{user_id}")
def delete_user(user_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM users WHERE user_id = %s", (user_id,))
    db.commit()
    return {"message": "User deleted"}

@app.put("/users/{user_id}/assign/{client_id}")
def assign_user_to_client(user_id: int, client_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("UPDATE users SET assigned_client_id=%s WHERE user_id=%s", (client_id, user_id))
    db.commit()
    return {"message": "User assigned to client"}

# ── AUTH ──────────────────────────────────────────────────────────────────────
@app.post("/auth/login")
def login(req: LoginRequest, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE email = %s", (req.email,))
    user = cursor.fetchone()
    if not user:
        raise HTTPException(status_code=401, detail="Invalid email or password")

    if user.get("login_locked"):
        cursor.execute(
            "INSERT INTO login_history (user_id, ip_address, device, status) VALUES (%s, %s, %s, %s)",
            (user["user_id"], None, None, "locked")
        )
        db.commit()
        raise HTTPException(status_code=403, detail="Account is locked")

    if not verify_password(req.password, user["password_hash"]):
        cursor.execute("UPDATE users SET failed_attempts = failed_attempts + 1 WHERE user_id = %s", (user["user_id"],))
        cursor.execute(
            "INSERT INTO login_history (user_id, ip_address, device, status) VALUES (%s, %s, %s, %s)",
            (user["user_id"], None, None, "failed")
        )
        db.commit()
        raise HTTPException(status_code=401, detail="Invalid email or password")

    if user["status"] != "Active":
        cursor.execute(
            "INSERT INTO login_history (user_id, ip_address, device, status) VALUES (%s, %s, %s, %s)",
            (user["user_id"], None, None, "inactive")
        )
        db.commit()
        raise HTTPException(status_code=403, detail="Account is inactive")

    token = create_token({"user_id": user["user_id"], "email": user["email"], "role": user["role"]})
    cursor.execute(
        "UPDATE users SET failed_attempts = 0, last_login = NOW() WHERE user_id = %s",
        (user["user_id"],)
    )
    cursor.execute(
        "INSERT INTO login_history (user_id, ip_address, device, status) VALUES (%s, %s, %s, %s)",
        (user["user_id"], None, None, "success")
    )
    db.commit()
    return {"access_token": token, "token_type": "bearer",
            "user": {"user_id": user["user_id"], "full_name": user["full_name"],
                     "email": user["email"], "role": user["role"]}}

@app.post("/auth/password-reset-request")
def password_reset_request(req: PasswordResetRequest, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE email = %s", (req.email,))
    user = cursor.fetchone()
    if not user:
        raise HTTPException(status_code=404, detail="Email not found")
    token = secrets.token_urlsafe(32)
    expires_at = datetime.utcnow() + timedelta(hours=1)
    cursor2 = db.cursor()
    cursor2.execute("INSERT INTO password_resets (user_id, token, expires_at) VALUES (%s, %s, %s)",
                    (user["user_id"], token, expires_at))
    db.commit()
    return {"message": "Password reset token generated", "token": token}

@app.post("/auth/password-reset-confirm")
def password_reset_confirm(req: PasswordResetConfirm, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM password_resets WHERE token = %s AND expires_at > NOW()", (req.token,))
    reset = cursor.fetchone()
    if not reset:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    hashed = hash_password(req.new_password)
    cursor2 = db.cursor()
    cursor2.execute("UPDATE users SET password_hash = %s WHERE user_id = %s", (hashed, reset["user_id"]))
    cursor2.execute("DELETE FROM password_resets WHERE token = %s", (req.token,))
    db.commit()
    return {"message": "Password reset successful"}

# ── COLUMN MAPPINGS ───────────────────────────────────────────────────────────
@app.get("/column-mappings")
def get_all_mappings(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM column_mappings")
    return cursor.fetchall()

@app.get("/column-mappings/{client_id}")
def get_client_mappings(client_id: str, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM column_mappings WHERE client_id = %s", (client_id,))
    return cursor.fetchall()

@app.post("/column-mappings")
def create_mapping(m: ColumnMapping, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO column_mappings (client_id, file_type, original_column, mapped_to, confirmed_by)
           VALUES (%s, %s, %s, %s, %s)""",
        (m.client_id, m.file_type, m.original_column, m.mapped_to, m.confirmed_by)
    )
    db.commit()
    return {"id": cursor.lastrowid, "message": "Column mapping created"}

@app.put("/column-mappings/{mapping_id}")
def update_mapping(mapping_id: int, m: ColumnMapping, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """UPDATE column_mappings SET client_id=%s, file_type=%s, original_column=%s,
           mapped_to=%s, confirmed_by=%s WHERE id=%s""",
        (m.client_id, m.file_type, m.original_column, m.mapped_to, m.confirmed_by, mapping_id)
    )
    db.commit()
    return {"message": "Column mapping updated"}

@app.delete("/column-mappings/{mapping_id}")
def delete_mapping(mapping_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM column_mappings WHERE id = %s", (mapping_id,))
    db.commit()
    return {"message": "Column mapping deleted"}

# ── ENGAGEMENTS ───────────────────────────────────────────────────────────────
@app.get("/engagements")
def get_engagements(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT e.*, c.company_name FROM engagements e
        LEFT JOIN clients c ON e.client_id = c.client_id
        ORDER BY e.created_at DESC
    """)
    return cursor.fetchall()

@app.get("/engagements/{engagement_id}")
def get_engagement(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT e.*, c.company_name FROM engagements e
        LEFT JOIN clients c ON e.client_id = c.client_id
        WHERE e.engagement_id = %s
    """, (engagement_id,))
    engagement = cursor.fetchone()
    if not engagement:
        raise HTTPException(status_code=404, detail="Engagement not found")
    return engagement

@app.post("/engagements")
def create_engagement(e: Engagement, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """INSERT INTO engagements (client_id, engagement_name, financial_year, status, start_date, end_date)
           VALUES (%s, %s, %s, %s, %s, %s)""",
        (e.client_id, e.engagement_name, e.financial_year, e.status, e.start_date, e.end_date)
    )
    engagement_id = cursor.lastrowid
    for section in ["Revenue", "Expenses", "Inventory", "Cash & Bank"]:
        cursor.execute("INSERT INTO audit_sections (engagement_id, section_name) VALUES (%s, %s)",
                       (engagement_id, section))
    db.commit()
    return {"engagement_id": engagement_id, "message": "Engagement created with default audit sections"}

@app.put("/engagements/{engagement_id}")
def update_engagement(engagement_id: int, e: Engagement, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        """UPDATE engagements SET client_id=%s, engagement_name=%s, financial_year=%s,
           status=%s, start_date=%s, end_date=%s WHERE engagement_id=%s""",
        (e.client_id, e.engagement_name, e.financial_year, e.status, e.start_date, e.end_date, engagement_id)
    )
    db.commit()
    return {"message": "Engagement updated"}
@app.delete("/engagements/{engagement_id}")
def delete_engagement(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    # Delete submissions first (they reference audit_sections)
    cursor.execute("DELETE FROM submissions WHERE engagement_id = %s", (engagement_id,))
    # Then delete audit sections
    cursor.execute("DELETE FROM audit_sections WHERE engagement_id = %s", (engagement_id,))
    # Then delete team members
    cursor.execute("DELETE FROM engagement_team WHERE engagement_id = %s", (engagement_id,))
    # Finally delete the engagement itself
    cursor.execute("DELETE FROM engagements WHERE engagement_id = %s", (engagement_id,))
    db.commit()
    return {"message": "Engagement deleted"}
# ── ENGAGEMENT TEAM ───────────────────────────────────────────────────────────
@app.get("/engagements/{engagement_id}/team")
def get_engagement_team(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT et.*, u.full_name, u.email, u.role FROM engagement_team et
        LEFT JOIN users u ON et.user_id = u.user_id
        WHERE et.engagement_id = %s
    """, (engagement_id,))
    return cursor.fetchall()

@app.post("/engagements/{engagement_id}/team")
def add_team_member(engagement_id: int, t: EngagementTeam, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("INSERT INTO engagement_team (engagement_id, user_id, role) VALUES (%s, %s, %s)",
                   (engagement_id, t.user_id, t.role))
    db.commit()
    return {"team_id": cursor.lastrowid, "message": "Team member added"}

@app.delete("/engagements/{engagement_id}/team/{user_id}")
def remove_team_member(engagement_id: int, user_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM engagement_team WHERE engagement_id=%s AND user_id=%s",
                   (engagement_id, user_id))
    db.commit()
    return {"message": "Team member removed"}

# ── AUDIT SECTIONS ────────────────────────────────────────────────────────────
@app.get("/engagements/{engagement_id}/sections")
def get_audit_sections(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT s.*, u.full_name as assigned_to_name FROM audit_sections s
        LEFT JOIN users u ON s.assigned_to = u.user_id
        WHERE s.engagement_id = %s
    """, (engagement_id,))
    return cursor.fetchall()

@app.post("/engagements/{engagement_id}/sections")
def add_audit_section(engagement_id: int, s: AuditSection, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        "INSERT INTO audit_sections (engagement_id, section_name, status, assigned_to) VALUES (%s, %s, %s, %s)",
        (engagement_id, s.section_name, s.status, s.assigned_to)
    )
    db.commit()
    return {"section_id": cursor.lastrowid, "message": "Audit section added"}

@app.put("/audit-sections/{section_id}")
def update_audit_section(section_id: int, s: AuditSection, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute(
        "UPDATE audit_sections SET section_name=%s, status=%s, assigned_to=%s WHERE section_id=%s",
        (s.section_name, s.status, s.assigned_to, section_id)
    )
    db.commit()
    return {"message": "Audit section updated"}

@app.delete("/audit-sections/{section_id}")
def delete_audit_section(section_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM audit_sections WHERE section_id = %s", (section_id,))
    db.commit()
    return {"message": "Audit section deleted"}

# ── SUBMISSIONS ───────────────────────────────────────────────────────────────
@app.get("/audit-sections/{section_id}/latest-submission")
def get_section_latest_submission(section_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute(
        "SELECT s.*, u.full_name as submitted_by_name "
        "FROM submissions s "
        "LEFT JOIN users u ON s.submitted_by = u.user_id "
        "WHERE s.section_id = %s "
        "ORDER BY s.created_at DESC "
        "LIMIT 1",
        (section_id,)
    )
    row = cursor.fetchone()
    return row if row else None

@app.get("/submissions")
def get_all_submissions(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT s.*, u.full_name as submitted_by_name, e.engagement_name, sec.section_name
        FROM submissions s
        LEFT JOIN users u ON s.submitted_by = u.user_id
        LEFT JOIN engagements e ON s.engagement_id = e.engagement_id
        LEFT JOIN audit_sections sec ON s.section_id = sec.section_id
        ORDER BY s.created_at DESC
    """)
    return cursor.fetchall()

@app.get("/submissions/{submission_id}")
def get_submission(submission_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT s.*, u.full_name as submitted_by_name, e.engagement_name, sec.section_name
        FROM submissions s
        LEFT JOIN users u ON s.submitted_by = u.user_id
        LEFT JOIN engagements e ON s.engagement_id = e.engagement_id
        LEFT JOIN audit_sections sec ON s.section_id = sec.section_id
        WHERE s.submission_id = %s
    """, (submission_id,))
    submission = cursor.fetchone()
    if not submission:
        raise HTTPException(status_code=404, detail="Submission not found")
    return submission

@app.post("/submissions")
def create_submission(s: Submission, db=Depends(get_db)):
    insert_cursor = db.cursor()
    insert_cursor.execute(
        """INSERT INTO submissions (engagement_id, section_id, submitted_by, status, current_stage, notes)
           VALUES (%s, %s, %s, %s, %s, %s)""",
        (s.engagement_id, s.section_id, s.submitted_by, s.status, s.current_stage, s.notes)
    )
    submission_id = insert_cursor.lastrowid
    cursor = db.cursor(dictionary=True)
    if s.current_stage and s.current_stage != "Accountant":
        cursor.execute("""
             SELECT e.engagement_name, sec.section_name FROM engagements e
             LEFT JOIN audit_sections sec ON sec.engagement_id = e.engagement_id
             WHERE sec.section_id = %s
       """, (s.section_id,))
        info = cursor.fetchone()
        if info:
            message = f"{info['section_name']} for {info['engagement_name']} is now {s.status}"
            cursor.execute("""
                SELECT DISTINCT u.user_id FROM users u
                INNER JOIN engagement_team et ON u.user_id = et.user_id
                WHERE et.engagement_id = %s AND COALESCE(NULLIF(et.role, ''), u.role) = %s
            """, (s.engagement_id, s.current_stage))
            for auditor in cursor.fetchall():
                cursor.execute(
                    "INSERT INTO notifications (user_id, message, type, engagement_id) VALUES (%s, %s, %s, %s)",
                    (auditor['user_id'], message, "submission_review", s.engagement_id)
                )
    db.commit()
    return {"submission_id": submission_id, "message": "Submission created"}
def auto_update_engagement_status(engagement_id: int, db):
    cursor = db.cursor(dictionary=True)

    # Get all sections for this engagement
    cursor.execute("SELECT section_id FROM audit_sections WHERE engagement_id = %s", (engagement_id,))
    sections = cursor.fetchall()
    if not sections:
        return

    section_ids = [s['section_id'] for s in sections]

    # Get latest submission status per section
    statuses = []
    for sid in section_ids:
        cursor.execute("""
            SELECT status FROM submissions
            WHERE section_id = %s
            ORDER BY created_at DESC LIMIT 1
        """, (sid,))
        row = cursor.fetchone()
        statuses.append(row['status'] if row else 'Draft')

    # Get current engagement status
    cursor.execute("SELECT status FROM engagements WHERE engagement_id = %s", (engagement_id,))
    eng = cursor.fetchone()
    current = eng['status'] if eng else 'Planning'

    # Determine new status based on submission statuses
    if all(s == 'Approved' for s in statuses):
        new_status = 'Review'
    elif any(s in ('Under Review', 'Changes Requested', 'Approved') for s in statuses):
        new_status = 'In Progress'
    else:
        new_status = current  # no change

    # Only update if status actually changed
    if new_status != current:
        cursor.execute(
            "UPDATE engagements SET status = %s WHERE engagement_id = %s",
            (new_status, engagement_id)
        )
@app.put("/submissions/{submission_id}/status")
def update_submission_status(submission_id: int, s: SubmissionStatus, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT sub.*, e.engagement_name, e.engagement_id, sec.section_name
        FROM submissions sub
        LEFT JOIN engagements e ON sub.engagement_id = e.engagement_id
        LEFT JOIN audit_sections sec ON sub.section_id = sec.section_id
        WHERE sub.submission_id = %s
    """, (submission_id,))
    sub = cursor.fetchone()
    if not sub:
        raise HTTPException(status_code=404, detail="Submission not found")

    cursor2 = db.cursor(dictionary=True)
    if s.updated_by:
        cursor2.execute(
            "UPDATE submissions SET status=%s, current_stage=%s, notes=%s, submitted_by=%s WHERE submission_id=%s",
            (s.status, s.current_stage, s.notes, s.updated_by, submission_id)
        )
    else:
        cursor2.execute(
            "UPDATE submissions SET status=%s, current_stage=%s, notes=%s WHERE submission_id=%s",
            (s.status, s.current_stage, s.notes, submission_id)
        )

    target_roles = []
    if s.current_stage:
        target_roles = [s.current_stage]
    elif s.status in ("Approved", "Cancelled"):
        target_roles = ["Accountant", "Auditor", "Senior Auditor", "Assistant Manager",
                        "Audit Manager", "Engagement Partner", "Quality Reviewer"]

    if target_roles:
        message = f"{sub['section_name']} for {sub['engagement_name']} is now {s.status}"
        format_strings = ','.join(['%s'] * len(target_roles))
        cursor2.execute(f"""
            SELECT DISTINCT u.user_id FROM users u
            INNER JOIN engagement_team et ON u.user_id = et.user_id
            WHERE et.engagement_id = %s AND COALESCE(NULLIF(et.role, ''), u.role) IN ({format_strings})
        """, (sub['engagement_id'], *target_roles))
        for row in cursor2.fetchall():
            cursor2.execute(
                "INSERT INTO notifications (user_id, message, type, engagement_id) VALUES (%s, %s, %s, %s)",
                (row['user_id'], message, "submission_review", sub['engagement_id'])
            )

    # Auto-update engagement status based on all section submission states
    auto_update_engagement_status(sub['engagement_id'], db)

    # Auto-update engagement status based on all section submission states
    auto_update_engagement_status(sub['engagement_id'], db)

    # Auto-update engagement status based on all section submission states
    auto_update_engagement_status(sub['engagement_id'], db)

    db.commit()
    return {"message": f"Submission status updated to {s.status}"}

@app.delete("/submissions/{submission_id}")
def delete_submission(submission_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("DELETE FROM submissions WHERE submission_id = %s", (submission_id,))
    db.commit()
    return {"message": "Submission deleted"}

# ── NOTIFICATIONS ─────────────────────────────────────────────────────────────
@app.get("/notifications/{user_id}")
def get_user_notifications(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM notifications WHERE user_id = %s ORDER BY created_at DESC", (user_id,))
    return cursor.fetchall()

@app.get("/notifications/{user_id}/unread")
def get_unread_notifications(user_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM notifications WHERE user_id = %s AND is_read = FALSE ORDER BY created_at DESC",
                   (user_id,))
    return cursor.fetchall()

@app.put("/notifications/{notification_id}/read")
def mark_notification_read(notification_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("UPDATE notifications SET is_read = TRUE WHERE notification_id = %s", (notification_id,))
    db.commit()
    return {"message": "Notification marked as read"}

@app.put("/notifications/{user_id}/read-all")
def mark_all_read(user_id: int, db=Depends(get_db)):
    cursor = db.cursor()
    cursor.execute("UPDATE notifications SET is_read = TRUE WHERE user_id = %s", (user_id,))
    db.commit()
    return {"message": "All notifications marked as read"}

# ── FILE UPLOAD ───────────────────────────────────────────────────────────────
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
ALLOWED_EXTENSIONS = {"csv", "xlsx", "xls", "pdf", "docx"}

# ── AI UPLOAD HELPERS ─────────────────────────────────────────────────────────
def get_extension(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

def extract_pdf(file_path: str):
    tables = []
    full_text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_tables = page.extract_tables()
            for table in page_tables:
                if table:
                    headers = table[0]
                    rows = table[1:]
                    df = pd.DataFrame(rows, columns=headers)
                    tables.append(df)
            full_text += page.extract_text() or ""
    if tables:
        return pd.concat(tables, ignore_index=True), "table"
    lines = [line.strip() for line in full_text.splitlines() if line.strip()]
    if lines:
        return pd.DataFrame({"raw_text": lines}), "text"
    return None, None

def extract_docx(file_path: str):
    doc = Document(file_path)
    tables = []
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = []
        for row in table.rows[1:]:
            rows.append([cell.text.strip() for cell in row.cells])
        df = pd.DataFrame(rows, columns=headers)
        tables.append(df)
    if tables:
        return pd.concat(tables, ignore_index=True), "table"
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if lines:
        return pd.DataFrame({"raw_text": lines}), "text"
    return None, None

def read_file_to_df(save_path: str, ext: str):
    if ext == "csv":
        df = pd.read_csv(save_path, dtype=str)
    elif ext in ["xlsx", "xls"]:
        df = pd.read_excel(save_path, dtype=str)
    elif ext == "pdf":
        df, _ = extract_pdf(save_path)
        return df
    elif ext == "docx":
        df, _ = extract_docx(save_path)
        return df
    else:
        return None
    if df is not None:
        df = df.dropna(axis=1, how='all')
        df = df.loc[:, ~(df == '').all()]
        df = df.map(lambda x: x.strip() if isinstance(x, str) else x)
    return df

def calculate_fill_rates(df: pd.DataFrame) -> dict:
    fill_rates = {}
    total = len(df)
    for col in df.columns:
        filled = df[col].replace("", float("nan")).dropna().count()
        fill_rates[col] = round(filled / total, 2) if total > 0 else 0.0
    return fill_rates

def compute_schema_fingerprint(columns: list) -> str:
    sorted_cols = sorted([col.lower().strip() for col in columns])
    return hashlib.md5(json.dumps(sorted_cols).encode()).hexdigest()

def locate_uploaded_file(file_id: str):
    for extension in ALLOWED_EXTENSIONS:
        path = os.path.join(UPLOAD_DIR, f"{file_id}.{extension}")
        if os.path.exists(path):
            return path, extension
    return None, None

def run_cleaning_cycle(file_id: str, client_id: str, file_type: str, mapping: dict):
    save_path, file_ext = locate_uploaded_file(file_id)
    if not save_path:
        raise HTTPException(status_code=404, detail="File not found. Please upload the file first.")
    try:
        df = read_file_to_df(save_path, file_ext)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not read file.")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")
    try:
        cleaned_df, report = clean_dataframe(df, mapping)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Cleaning failed: {str(e)}")
    return cleaned_df, report

# ── ROOT ──────────────────────────────────────────────────────────────────────
@app.get("/")
def root():
    return {"message": "AuditIQ API is running"}

# ── AI UPLOAD ROUTES ──────────────────────────────────────────────────────────
@app.post("/upload")
async def upload_file_ai(
    file: UploadFile = File(...),
    client_id: str = Form(...)
):
    ext = get_extension(file.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=415, detail=f"File type .{ext} not supported. Upload Excel, CSV, PDF or DOCX file only.")
    MAX_FILE_SIZE = 50
    file_bytes = await file.read()
    if len(file_bytes) / (1024 * 1024) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail=f"File exceeds {MAX_FILE_SIZE} MB limit.")
    file.file.seek(0)
    file_id = str(uuid.uuid4())
    save_path = os.path.join(UPLOAD_DIR, f"{file_id}.{ext}")
    with open(save_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    if ext == "pdf":
        df, source = extract_pdf(save_path)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not extract any content from PDF.")
        save_upload(file_id, client_id, file.filename, ext, len(df))
        fill_rates = calculate_fill_rates(df)
        return {"file_id": file_id, "client_id": client_id, "filename": file.filename, "source": source,
                "rows": len(df), "columns": list(df.columns), "fill_rates": fill_rates,
                "preview": df.head(5).fillna("").to_dict(orient="records"),
                "message": f"PDF uploaded — extracted via {source}"}
    if ext == "docx":
        df, source = extract_docx(save_path)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not extract any content from DOCX.")
        save_upload(file_id, client_id, file.filename, ext, len(df))
        fill_rates = calculate_fill_rates(df)
        return {"file_id": file_id, "client_id": client_id, "filename": file.filename, "source": source,
                "rows": len(df), "columns": list(df.columns), "fill_rates": fill_rates,
                "preview": df.head(5).fillna("").to_dict(orient="records"),
                "message": f"DOCX uploaded — extracted via {source}"}
    try:
        df = pd.read_csv(save_path, dtype=str) if ext == "csv" else pd.read_excel(save_path, dtype=str)
        df = df.map(lambda x: x.strip() if isinstance(x, str) else x)
        df = df.dropna(axis=1, how='all')
        df = df.loc[:, ~(df == '').all()]
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")
    save_upload(file_id, client_id, file.filename, ext, len(df))
    fill_rates = calculate_fill_rates(df)
    return {"file_id": file_id, "client_id": client_id, "filename": file.filename, "source": "table",
            "rows": len(df), "columns": list(df.columns), "fill_rates": fill_rates,
            "fingerprint": compute_schema_fingerprint(list(df.columns)),
            "preview": df.head(5).fillna("").to_dict(orient="records"),
            "message": "File uploaded and processed successfully"}

@app.post("/detect-columns")
async def detect_columns_endpoint(
    client_id: str = Form(...),
    file_id: str = Form(...),
    columns: str = Form(...),
    file_type: str = Form("general"),
):
    print(f"DEBUG: Detect columns requested - file_id={file_id}, client_id={client_id}, file_type={file_type}")
    try:
        columns_list = json.loads(columns)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid columns format.")
    save_path, file_ext = locate_uploaded_file(file_id)
    if not save_path:
        raise HTTPException(status_code=404, detail="File not found. Please upload the file first.")
    try:
        df = read_file_to_df(save_path, file_ext)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not read file.")
        sample_values = {}
        for col in columns_list:
            if col in df.columns:
                non_empty = df[col].dropna().replace("", float("nan")).dropna()
                sample_values[col] = str(non_empty.iloc[0]) if len(non_empty) > 0 else ""
            else:
                sample_values[col] = ""
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")
    saved_mapping = get_mapping(client_id, file_type)
    if saved_mapping and all(col in saved_mapping for col in columns_list):
        filtered_mapping = {col: saved_mapping[col] for col in columns_list}
        result = build_detection_result(columns_list, filtered_mapping)
        result.update({"file_id": file_id, "source": "saved_mapping",
                        "message": "Mapping loaded from saved client profile — LLM skipped."})
        return result
    try:
        mapping = detect_columns_with_llm(columns_list, sample_values)
        if not mapping:
            raise HTTPException(status_code=500, detail="LLM returned empty mapping.")
        result = build_detection_result(columns_list, mapping)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Detection failed: {str(e)}")
    result.update({"file_id": file_id, "source": "llm_detection"})
    return result

@app.post("/save-mapping")
async def save_mapping_endpoint(
    client_id: str = Form(...),
    file_type: str = Form(...),
    mapping: str = Form(...),
    confirmed_by: str = Form(None)
):
    try:
        mapping_dict = json.loads(mapping)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid mapping format.")
    if not mapping_dict:
        raise HTTPException(status_code=400, detail="Mapping cannot be empty.")
    save_mapping(client_id, file_type, mapping_dict, confirmed_by)
    return {"client_id": client_id, "file_type": file_type, "columns_saved": len(mapping_dict),
            "message": f"Mapping saved successfully for client {client_id} and file type {file_type}."}

@app.get("/get-mapping/{client_id}")
async def get_mapping_endpoint(client_id: str, file_type: str = "general"):
    mapping = get_mapping(client_id, file_type)
    if not mapping:
        return {"client_id": client_id, "file_type": file_type, "mapping": {},
                "message": "No saved mapping found for this client."}
    return {"client_id": client_id, "file_type": file_type, "mapping": mapping,
            "columns_mapped": len(mapping), "message": "Saved mapping retrieved successfully."}

@app.get("/uploads/{client_id}")
async def get_uploads_endpoint(client_id: str):
    uploads = get_uploads(client_id)
    return {"client_id": client_id, "total_uploads": len(uploads), "uploads": uploads}

@app.post("/clean")
async def clean_file(
    file_id: str = Form(...),
    client_id: str = Form(...),
    file_type: str = Form("general")
) -> dict:
    mapping = get_mapping(client_id, file_type)
    if not mapping:
        raise HTTPException(status_code=400, detail="No saved mapping found. Please detect and confirm the mapping first.")
    cleaned_df, report = run_cleaning_cycle(file_id, client_id, file_type, mapping)
    return {"file_id": file_id, "client_id": client_id, "file_type": file_type,
            "cleaned_data": cleaned_df.fillna("").astype(str).map(lambda x: x.strip()).to_dict(orient="records"),
            "validation_report": report,
            "can_proceed": report.get("total_issues", 0) == 0,
            "message": "File cleaned successfully."}

@app.post("/clients/{client_id}/upload")
def upload_file(client_id: int, file: UploadFile = File(...), db=Depends(get_db)):
    allowed_types = ["xlsx", "xls", "csv", "pdf", "tiff", "tif", "jpg", "jpeg", "png", "xml", "json", "txt"]
    file_ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if file_ext not in allowed_types:
        raise HTTPException(status_code=400, detail="File format not allowed.")
    file_path = f"{UPLOAD_DIR}/{client_id}_{file.filename}"
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    cursor = db.cursor()
    cursor.execute("INSERT INTO uploads (client_id, file_name, file_type, file_path) VALUES (%s, %s, %s, %s)",
                   (client_id, file.filename, file_ext.upper(), file_path))
    db.commit()
    return {"file_id": cursor.lastrowid, "filename": file.filename, "type": file_ext.upper(),
            "message": "File uploaded successfully"}

@app.get("/clients/{client_id}/files")
def get_client_files(client_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM uploads WHERE client_id = %s", (client_id,))
    return cursor.fetchall()

@app.get("/files")
def get_all_files(db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT f.*, c.company_name FROM uploads f
        LEFT JOIN clients c ON f.client_id = c.client_id
        ORDER BY f.upload_date DESC
    """)
    return cursor.fetchall()


# ── SEND TO CLIENT ────────────────────────────────────────────────────────────
@app.post("/engagements/{engagement_id}/send-to-client")
async def send_to_client(engagement_id: int, db=Depends(get_db)):
    cursor = db.cursor(dictionary=True)
    cursor.execute("""
        SELECT e.*, c.company_name, c.email as client_email, c.contact_person
        FROM engagements e
        LEFT JOIN clients c ON e.client_id = c.client_id
        WHERE e.engagement_id = %s
    """, (engagement_id,))
    engagement = cursor.fetchone()
    if not engagement:
        raise HTTPException(status_code=404, detail="Engagement not found")
    if not engagement.get("client_email"):
        raise HTTPException(status_code=400, detail="Client has no email address on record")

    cursor.execute("""
        SELECT sec.section_name, s.status, s.notes
        FROM submissions s
        LEFT JOIN audit_sections sec ON s.section_id = sec.section_id
        WHERE s.engagement_id = %s AND s.status = 'Approved'
        ORDER BY sec.section_name
    """, (engagement_id,))
    approved_sections = cursor.fetchall()

    if not approved_sections:
        raise HTTPException(status_code=400, detail="No approved sections found for this engagement")

    sections_html = "".join([
        f"<tr><td style='padding:8px;border:1px solid #ddd'>{s['section_name']}</td>"
        f"<td style='padding:8px;border:1px solid #ddd;color:green'>Approved</td></tr>"
        for s in approved_sections
    ])

    html_body = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <h2 style="color: #1E3A5F;">Audit Report — {engagement['engagement_name']}</h2>
        <p>Dear {engagement['contact_person'] or engagement['company_name']},</p>
        <p>We are pleased to inform you that the following audit sections for <strong>{engagement['engagement_name']}</strong>
        (Financial Year {engagement['financial_year']}) have been completed and approved:</p>
        <table style="width:100%;border-collapse:collapse;margin:16px 0">
            <thead>
                <tr style="background:#1E3A5F;color:white">
                    <th style="padding:10px;text-align:left">Section</th>
                    <th style="padding:10px;text-align:left">Status</th>
                </tr>
            </thead>
            <tbody>{sections_html}</tbody>
        </table>
        <p>Please contact us if you have any questions regarding this audit.</p>
        <br>
        <p style="color:#7f8c8d;font-size:12px">This is an automated message from Audit AI.</p>
    </div>
    """

    message = MessageSchema(
        subject=f"Audit Report — {engagement['engagement_name']} (FY {engagement['financial_year']})",
        recipients=[engagement["client_email"]],
        body=html_body,
        subtype=MessageType.html,
    )

    try:
        fm = FastMail(mail_conf)
        await fm.send_message(message)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to send email: {str(e)}")

    cursor.execute("""
        SELECT u.user_id FROM users u
        INNER JOIN engagement_team et ON u.user_id = et.user_id
        WHERE et.engagement_id = %s
    """, (engagement_id,))
    for row in cursor.fetchall():
        cursor.execute(
            "INSERT INTO notifications (user_id, message, type) VALUES (%s, %s, %s)",
            (row['user_id'],
             f"Audit report for {engagement['engagement_name']} has been sent to {engagement['company_name']}",
             "engagement_alert")
        )
        # Mark engagement as Completed when report is sent to client
    cursor.execute(
        "UPDATE engagements SET status = 'Completed' WHERE engagement_id = %s",
        (engagement_id,)
    )
    db.commit()
    return {"message": f"Audit report sent to {engagement['client_email']}"}

# ── RUN ───────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("Audit:app", host="0.0.0.0", port=8000, reload=True)